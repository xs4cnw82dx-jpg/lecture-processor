import os
import uuid
import threading
import time
import io
import json
import csv
import re
import shutil
import subprocess
import html
import sys
import warnings
import logging
import ipaddress
import statistics
from datetime import datetime, timedelta, timezone
from urllib.parse import urlparse
try:
    from zoneinfo import ZoneInfo
except Exception:
    ZoneInfo = None

# Keep startup clean in local dev environments.
warnings.filterwarnings(
    "ignore",
    message=r"urllib3 v2 only supports OpenSSL 1\.1\.1\+.*"
)

import stripe
from flask import Flask, request, jsonify, render_template, send_file, Response, stream_with_context, g, redirect, abort
from google import genai
from google.genai import types
from dotenv import load_dotenv
from werkzeug.utils import secure_filename
from werkzeug.exceptions import RequestEntityTooLarge
try:
    from flask_compress import Compress
except Exception:
    Compress = None
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
try:
    import sentry_sdk
    from sentry_sdk.integrations.flask import FlaskIntegration
except Exception:
    sentry_sdk = None
    FlaskIntegration = None
try:
    import imageio_ffmpeg
except Exception:
    imageio_ffmpeg = None

REPORTLAB_AVAILABLE = True
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, ListFlowable, ListItem, PageBreak
except Exception:
    REPORTLAB_AVAILABLE = False
import firebase_admin
from firebase_admin import credentials, auth, firestore
from lecture_processor.services import (
    analytics_service,
    auth_service,
    file_service,
    job_state_service,
    prompt_registry,
    rate_limit_service,
)
from lecture_processor.repositories import (
    admin_repo,
    job_logs_repo,
    purchases_repo,
    runtime_jobs_repo,
    study_repo,
    users_repo,
)
from lecture_processor.blueprints import auth_bp, account_bp, study_bp, upload_bp, admin_bp, payments_bp

LEGACY_MODULE_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT_DIR = os.path.dirname(LEGACY_MODULE_DIR)

load_dotenv()
app = Flask(
    __name__,
    template_folder=os.path.join(PROJECT_ROOT_DIR, 'templates'),
    static_folder=os.path.join(PROJECT_ROOT_DIR, 'static'),
)
app.secret_key = os.getenv('FLASK_SECRET_KEY', os.urandom(32).hex())
if Compress is not None:
    Compress(app)
LOG_LEVEL = (os.getenv('LOG_LEVEL', 'INFO') or 'INFO').strip().upper()
logging.basicConfig(
    level=getattr(logging, LOG_LEVEL, logging.INFO),
    format='%(asctime)s %(levelname)s %(name)s %(message)s',
)
logger = logging.getLogger('lecture_processor')


def log_event(level, event, **fields):
    payload = {'event': event}
    for key, value in fields.items():
        payload[str(key)] = value
    logger.log(level, json.dumps(payload, ensure_ascii=True))

UPLOAD_FOLDER = 'uploads'
STUDY_AUDIO_RELATIVE_DIR = 'study_audio'
STUDY_AUDIO_ROOT = os.path.abspath(os.path.join(UPLOAD_FOLDER, STUDY_AUDIO_RELATIVE_DIR))
ALLOWED_SLIDE_EXTENSIONS = {'pdf', 'pptx'}
ALLOWED_PDF_EXTENSIONS = ALLOWED_SLIDE_EXTENSIONS
ALLOWED_AUDIO_EXTENSIONS = {'mp3', 'm4a', 'wav', 'aac', 'ogg', 'flac'}
MAX_PDF_UPLOAD_BYTES = 50 * 1024 * 1024
MAX_AUDIO_UPLOAD_BYTES = 500 * 1024 * 1024
MAX_CONTENT_LENGTH = (MAX_PDF_UPLOAD_BYTES + MAX_AUDIO_UPLOAD_BYTES + (10 * 1024 * 1024))
ALLOWED_SLIDE_MIME_TYPES = {
    'application/pdf',
    'application/x-pdf',
    'application/vnd.openxmlformats-officedocument.presentationml.presentation',
    'application/vnd.ms-powerpoint',
}
ALLOWED_PDF_MIME_TYPES = ALLOWED_SLIDE_MIME_TYPES
ALLOWED_AUDIO_MIME_TYPES = {
    'audio/mpeg', 'audio/mp3', 'audio/mp4', 'audio/x-m4a', 'audio/wav', 'audio/x-wav',
    'audio/aac', 'audio/ogg', 'audio/flac',
}

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
GEMINI_API_KEY = (os.getenv('GEMINI_API_KEY', '') or '').strip()
if GEMINI_API_KEY:
    try:
        client = genai.Client(api_key=GEMINI_API_KEY)
    except Exception as e:
        client = None
        logger.warning(f"âš ï¸ Gemini client disabled: {e}")
else:
    client = None
    logger.warning("âš ï¸ GEMINI_API_KEY not set; AI processing features are disabled.")
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH

# --- Firebase Setup ---
db = None
firebase_init_error = ''
try:
    firebase_creds_raw = (os.getenv('FIREBASE_CREDENTIALS', '') or '').strip()
    local_creds_file_exists = os.path.exists('firebase-credentials.json')
    if local_creds_file_exists:
        logger.warning(
            "Local firebase-credentials.json detected. Prefer FIREBASE_CREDENTIALS environment variable for safer deployments."
        )
    if firebase_creds_raw:
        cred = credentials.Certificate(json.loads(firebase_creds_raw))
    elif local_creds_file_exists:
        cred = credentials.Certificate('firebase-credentials.json')
    else:
        raise ValueError("FIREBASE_CREDENTIALS is not set and firebase-credentials.json was not found.")
    if not firebase_admin._apps:
        firebase_admin.initialize_app(cred)
    db = firestore.client()
except Exception as e:
    firebase_init_error = str(e)
    logger.warning(f"âš ï¸ Firebase initialization skipped: {firebase_init_error}")

# --- Stripe Setup ---
stripe.api_key = os.getenv('STRIPE_SECRET_KEY')
STRIPE_PUBLISHABLE_KEY = os.getenv('STRIPE_PUBLISHABLE_KEY', '')
STRIPE_WEBHOOK_SECRET = os.getenv('STRIPE_WEBHOOK_SECRET', '')
ADMIN_EMAILS = {email.strip().lower() for email in os.getenv('ADMIN_EMAILS', '').split(',') if email.strip()}
ADMIN_UIDS = {uid.strip() for uid in os.getenv('ADMIN_UIDS', '').split(',') if uid.strip()}
ADMIN_SESSION_COOKIE_NAME = 'lp_admin_session'
ADMIN_SESSION_DURATION_SECONDS = int(os.getenv('ADMIN_SESSION_DURATION_SECONDS', str(8 * 60 * 60)) or (8 * 60 * 60))
ADMIN_PAGE_UNAUTHORIZED_MODE = str(os.getenv('ADMIN_PAGE_UNAUTHORIZED_MODE', 'redirect')).strip().lower()

# --- In-Memory Storage (jobs only â€” credits are in Firestore now) ---
jobs = {}
JOBS_LOCK = threading.RLock()
RUNTIME_JOBS_COLLECTION = 'runtime_jobs'
RUNTIME_JOB_RECOVERY_BATCH_LIMIT = int(os.getenv('RUNTIME_JOB_RECOVERY_BATCH_LIMIT', '200') or 200)
RUNTIME_JOB_PERSISTED_FIELDS = {
    'status',
    'step',
    'step_description',
    'total_steps',
    'mode',
    'user_id',
    'user_email',
    'credit_deducted',
    'credit_refunded',
    'started_at',
    'finished_at',
    'result',
    'slide_text',
    'transcript',
    'flashcards',
    'test_questions',
    'flashcard_selection',
    'question_selection',
    'study_features',
    'output_language',
    'study_generation_error',
    'study_pack_id',
    'error',
    'billing_receipt',
    'interview_features',
    'interview_features_successful',
    'interview_summary',
    'interview_sections',
    'interview_combined',
    'interview_features_cost',
    'extra_slides_refunded',
    'audio_storage_key',
    'notes_audio_map',
    'transcript_segments',
    'token_usage_by_stage',
    'token_input_total',
    'token_output_total',
    'token_total',
    'export_manifest',
}
RUNTIME_JOB_MAX_STRING_LENGTH = 200000
AUDIO_STREAM_TOKEN_TTL_SECONDS = 3600
AUDIO_STREAM_TOKENS = {}
ALLOW_LEGACY_AUDIO_STREAM_TOKENS = str(os.getenv('ALLOW_LEGACY_AUDIO_STREAM_TOKENS', '0')).strip().lower() in {'1', 'true', 'yes', 'on'}
AUDIO_IMPORT_TOKEN_TTL_SECONDS = 30 * 60
AUDIO_IMPORT_TOKENS = {}
AUDIO_IMPORT_LOCK = threading.Lock()
FEATURE_AUDIO_SECTION_SYNC = os.getenv('FEATURE_AUDIO_SECTION_SYNC', '0').strip().lower() in {'1', 'true', 'yes', 'on'}
MAX_PROGRESS_PACKS_PER_SYNC = 300
MAX_PROGRESS_CARDS_PER_PACK = 2500
PROGRESS_DATE_RE = re.compile(r'^\d{4}-\d{2}-\d{2}$')
ANALYTICS_NAME_RE = re.compile(r'^[a-z0-9_]{2,64}$')
ANALYTICS_SESSION_ID_RE = re.compile(r'^[A-Za-z0-9_-]{6,80}$')
ANALYTICS_ALLOWED_EVENTS = {
    'auth_modal_opened',
    'auth_success',
    'auth_failed',
    'checkout_started',
    'payment_confirmed',
    'payment_cancelled',
    'process_clicked',
    'processing_started',
    'processing_completed',
    'processing_failed',
    'processing_timeout',
    'processing_retry_requested',
    'study_mode_opened',
    'payment_confirmed_backend',
    'processing_started_backend',
    'processing_completed_backend',
    'processing_failed_backend',
    'processing_finished_backend',
}
ANALYTICS_FUNNEL_STAGES = [
    {'event': 'auth_modal_opened', 'label': 'Opened sign-in'},
    {'event': 'auth_success', 'label': 'Signed in'},
    {'event': 'checkout_started', 'label': 'Started checkout'},
    {'event': 'payment_confirmed', 'label': 'Payment confirmed'},
    {'event': 'process_clicked', 'label': 'Clicked process'},
    {'event': 'processing_started', 'label': 'Upload accepted'},
    {'event': 'processing_completed', 'label': 'Processing complete'},
    {'event': 'study_mode_opened', 'label': 'Opened study mode'},
]
ANALYTICS_FUNNEL_EVENT_NAMES = {stage['event'] for stage in ANALYTICS_FUNNEL_STAGES}

def safe_int_env(name, default=0, minimum=1, maximum=100000):
    raw = os.getenv(name, str(default)).strip()
    try:
        value = int(raw)
    except Exception:
        value = int(default)
    return min(max(value, minimum), maximum)

UPLOAD_RATE_LIMIT_WINDOW_SECONDS = safe_int_env('UPLOAD_RATE_LIMIT_WINDOW_SECONDS', 600, minimum=10, maximum=86400)
UPLOAD_RATE_LIMIT_MAX_REQUESTS = safe_int_env('UPLOAD_RATE_LIMIT_MAX_REQUESTS', 10, minimum=1, maximum=1000)
MAX_ACTIVE_JOBS_PER_USER = safe_int_env('MAX_ACTIVE_JOBS_PER_USER', 2, minimum=1, maximum=20)
CHECKOUT_RATE_LIMIT_WINDOW_SECONDS = safe_int_env('CHECKOUT_RATE_LIMIT_WINDOW_SECONDS', 600, minimum=10, maximum=86400)
CHECKOUT_RATE_LIMIT_MAX_REQUESTS = safe_int_env('CHECKOUT_RATE_LIMIT_MAX_REQUESTS', 6, minimum=1, maximum=100)
ANALYTICS_RATE_LIMIT_WINDOW_SECONDS = safe_int_env('ANALYTICS_RATE_LIMIT_WINDOW_SECONDS', 60, minimum=10, maximum=3600)
ANALYTICS_RATE_LIMIT_MAX_REQUESTS = safe_int_env('ANALYTICS_RATE_LIMIT_MAX_REQUESTS', 240, minimum=10, maximum=5000)
VIDEO_IMPORT_RATE_LIMIT_WINDOW_SECONDS = safe_int_env('VIDEO_IMPORT_RATE_LIMIT_WINDOW_SECONDS', 600, minimum=30, maximum=86400)
VIDEO_IMPORT_RATE_LIMIT_MAX_REQUESTS = safe_int_env('VIDEO_IMPORT_RATE_LIMIT_MAX_REQUESTS', 8, minimum=1, maximum=200)
UPLOAD_MIN_FREE_DISK_BYTES = safe_int_env('UPLOAD_MIN_FREE_DISK_BYTES', 1024 * 1024 * 1024, minimum=50 * 1024 * 1024, maximum=500 * 1024 * 1024 * 1024)
UPLOAD_DAILY_BYTE_CAP = safe_int_env('UPLOAD_DAILY_BYTE_CAP', 2 * 1024 * 1024 * 1024, minimum=100 * 1024 * 1024, maximum=5 * 1024 * 1024 * 1024 * 1024)
UPLOAD_DAILY_COUNTER_COLLECTION = 'upload_usage_daily'
ACCOUNT_EXPORT_MAX_DOCS_PER_COLLECTION = safe_int_env('ACCOUNT_EXPORT_MAX_DOCS_PER_COLLECTION', 10000, minimum=100, maximum=50000)
ACCOUNT_DELETE_MAX_DOCS_PER_COLLECTION = safe_int_env('ACCOUNT_DELETE_MAX_DOCS_PER_COLLECTION', 10000, minimum=100, maximum=50000)
RATE_LIMIT_EVENTS = {}
RATE_LIMIT_LOCK = threading.Lock()
RATE_LIMIT_COUNTER_COLLECTION = 'rate_limit_counters'
RATE_LIMIT_FIRESTORE_ENABLED = str(os.getenv('RATE_LIMIT_FIRESTORE_ENABLED', '1')).strip().lower() in {'1', 'true', 'yes', 'on'}

SENTRY_BACKEND_DSN = os.getenv('SENTRY_DSN_BACKEND', '').strip()
SENTRY_FRONTEND_DSN = os.getenv('SENTRY_DSN_FRONTEND', '').strip()
SENTRY_ENVIRONMENT = (os.getenv('SENTRY_ENVIRONMENT', os.getenv('FLASK_ENV', 'production')) or 'production').strip()
SENTRY_RELEASE = (os.getenv('SENTRY_RELEASE', 'lecture-processor') or 'lecture-processor').strip()
LEGAL_CONTACT_EMAIL = (os.getenv('LEGAL_CONTACT_EMAIL', os.getenv('SUPPORT_EMAIL', '')) or '').strip()
DEV_ENV_NAMES = {'development', 'dev', 'local', 'test'}
APP_BOOT_TS = time.time()


def parse_cors_allowed_origins():
    raw = (os.getenv('CORS_ALLOWED_ORIGINS', '') or '').strip()
    if raw:
        origins = [part.strip().lower() for part in raw.split(',') if part.strip()]
        return set(origins)
    return {
        'http://127.0.0.1:5000',
        'http://localhost:5000',
        'http://127.0.0.1:10000',
        'http://localhost:10000',
        'https://lecture-processor.onrender.com',
        'https://lecture-processor-1.onrender.com',
    }


CORS_ALLOWED_ORIGINS = parse_cors_allowed_origins()

def safe_float_env(name, default=0.0):
    raw = os.getenv(name, str(default)).strip()
    try:
        value = float(raw)
    except Exception:
        return default
    return min(max(value, 0.0), 1.0)

SENTRY_TRACES_SAMPLE_RATE = safe_float_env('SENTRY_TRACES_SAMPLE_RATE', 0.0)

if SENTRY_BACKEND_DSN and sentry_sdk and FlaskIntegration:
    sentry_sdk.init(
        dsn=SENTRY_BACKEND_DSN,
        integrations=[FlaskIntegration()],
        traces_sample_rate=SENTRY_TRACES_SAMPLE_RATE,
        send_default_pii=False,
        environment=SENTRY_ENVIRONMENT,
        release=SENTRY_RELEASE,
    )

def is_dev_environment():
    env_value = str(SENTRY_ENVIRONMENT or '').strip().lower()
    flask_debug = str(os.getenv('FLASK_DEBUG', '0')).strip().lower() in {'1', 'true', 'yes', 'on'}
    return env_value in DEV_ENV_NAMES or flask_debug


def resolve_js_asset(filename):
    """Prefer minified JS assets in non-dev environments when available."""
    safe_name = str(filename or '').strip()
    if not safe_name.endswith('.js'):
        return safe_name
    if is_dev_environment():
        return safe_name
    min_name = safe_name[:-3] + '.min.js'
    min_path = os.path.join(PROJECT_ROOT_DIR, 'static', min_name)
    if os.path.exists(min_path):
        return min_name
    return safe_name


def infer_stripe_key_mode(key_value):
    key = str(key_value or '').strip()
    if not key:
        return 'missing'
    if key.startswith('sk_live_') or key.startswith('pk_live_'):
        return 'live'
    if key.startswith('sk_test_') or key.startswith('pk_test_'):
        return 'test'
    return 'unknown'


def build_admin_deployment_info(request_host=''):
    request_host = str(request_host or '').strip()
    request_hostname = request_host.split(':', 1)[0].strip().lower()
    render_hostname = str(os.getenv('RENDER_EXTERNAL_HOSTNAME', '') or '').strip().lower()
    render_external_url = str(os.getenv('RENDER_EXTERNAL_URL', '') or '').strip()
    render_service_id = str(os.getenv('RENDER_SERVICE_ID', '') or '').strip()
    render_deploy_id = str(os.getenv('RENDER_DEPLOY_ID', '') or '').strip()
    render_instance_id = str(os.getenv('RENDER_INSTANCE_ID', '') or '').strip()
    render_service_name = str(os.getenv('RENDER_SERVICE_NAME', '') or '').strip()
    render_git_commit = str(os.getenv('RENDER_GIT_COMMIT', '') or '').strip()
    render_git_branch = str(os.getenv('RENDER_GIT_BRANCH', '') or '').strip()
    render_detected = bool(str(os.getenv('RENDER', '') or '').strip() or render_service_id or render_deploy_id)
    host_matches_render = None
    if render_hostname and request_hostname:
        host_matches_render = request_hostname == render_hostname
    return {
        'runtime': 'render' if render_detected else 'local',
        'request_host': request_host,
        'request_hostname': request_hostname,
        'render_external_hostname': render_hostname,
        'render_external_url': render_external_url,
        'host_matches_render': host_matches_render,
        'service_id': render_service_id,
        'service_name': render_service_name,
        'deploy_id': render_deploy_id,
        'instance_id': render_instance_id,
        'git_branch': render_git_branch,
        'git_commit': render_git_commit,
        'git_commit_short': (render_git_commit[:12] if render_git_commit else ''),
        'app_boot_ts': APP_BOOT_TS,
        'app_uptime_seconds': max(0, round(time.time() - APP_BOOT_TS, 1)),
    }


def build_admin_runtime_checks():
    secret_key_mode = infer_stripe_key_mode(stripe.api_key)
    publishable_key_mode = infer_stripe_key_mode(STRIPE_PUBLISHABLE_KEY)
    webhook_configured = bool(str(STRIPE_WEBHOOK_SECRET or '').strip())
    stripe_keys_match = (
        secret_key_mode in {'live', 'test'}
        and publishable_key_mode in {'live', 'test'}
        and secret_key_mode == publishable_key_mode
    )
    soffice_available = bool(get_soffice_binary())
    ffmpeg_available = bool(get_ffmpeg_binary())
    ytdlp_available = bool(shutil.which('yt-dlp'))
    return {
        'firebase_ready': bool(db),
        'gemini_ready': bool(client),
        'stripe_secret_mode': secret_key_mode,
        'stripe_publishable_mode': publishable_key_mode,
        'stripe_keys_match': stripe_keys_match,
        'stripe_webhook_configured': webhook_configured,
        'pptx_conversion_available': soffice_available,
        'video_import_available': (ffmpeg_available and ytdlp_available),
        'ffmpeg_available': ffmpeg_available,
        'yt_dlp_available': ytdlp_available,
    }

def apply_cors_headers(response):
    origin = str(request.headers.get('Origin', '') or '').strip()
    if not origin:
        return response
    if not request.path.startswith('/api/'):
        return response
    if origin.lower() not in CORS_ALLOWED_ORIGINS:
        return response
    response.headers['Access-Control-Allow-Origin'] = origin
    response.headers['Vary'] = 'Origin'
    response.headers['Access-Control-Allow-Headers'] = 'Authorization, Content-Type'
    response.headers['Access-Control-Allow-Methods'] = 'GET, POST, PUT, PATCH, DELETE, OPTIONS'
    return response


@app.before_request
def handle_api_options_preflight():
    if request.method == 'OPTIONS' and request.path.startswith('/api/'):
        return apply_cors_headers(app.make_default_options_response())

@app.before_request
def attach_sentry_route_context():
    request_id = str(request.headers.get('X-Request-ID', '') or '').strip()[:120] or uuid.uuid4().hex
    g.request_id = request_id
    if not sentry_sdk:
        return
    try:
        with sentry_sdk.configure_scope() as scope:
            scope.set_tag('request.id', request_id)
            scope.set_tag('route.path', request.path)
            scope.set_tag('route.method', request.method)
            scope.set_tag('route.endpoint', request.endpoint or '')
            scope.set_tag('route.auth_header_present', 'true' if request.headers.get('Authorization') else 'false')
            scope.set_tag('route.environment', SENTRY_ENVIRONMENT or 'production')
            if request.content_type:
                scope.set_tag('route.content_type', str(request.content_type).split(';')[0][:80])
    except Exception:
        pass

@app.after_request
def attach_sentry_response_context(response):
    request_id = str(getattr(g, 'request_id', '') or '').strip()
    if request_id:
        response.headers['X-Request-ID'] = request_id
    if sentry_sdk:
        try:
            with sentry_sdk.configure_scope() as scope:
                scope.set_tag('route.status_code', str(response.status_code))
        except Exception:
            pass
    return apply_cors_headers(response)

@app.errorhandler(RequestEntityTooLarge)
def handle_request_entity_too_large(_error):
    return jsonify({'error': 'Upload too large. Maximum total upload size is 560MB (up to 50MB PDF and 500MB audio).'}), 413

# --- Credit Bundles (what users can buy) ---
CREDIT_BUNDLES = {
    'lecture_5': {
        'name': 'Lecture Notes â€” 5 Pack',
        'description': '5 standard lecture credits',
        'credits': {'lecture_credits_standard': 5},
        'price_cents': 999,
        'currency': 'eur',
    },
    'lecture_10': {
        'name': 'Lecture Notes â€” 10 Pack',
        'description': '10 standard lecture credits (best value)',
        'credits': {'lecture_credits_standard': 10},
        'price_cents': 1699,
        'currency': 'eur',
    },
    'slides_10': {
        'name': 'Slides Extraction â€” 10 Pack',
        'description': '10 slides extraction credits',
        'credits': {'slides_credits': 10},
        'price_cents': 499,
        'currency': 'eur',
    },
    'slides_25': {
        'name': 'Slides Extraction â€” 25 Pack',
        'description': '25 slides extraction credits (best value)',
        'credits': {'slides_credits': 25},
        'price_cents': 999,
        'currency': 'eur',
    },
    'interview_3': {
        'name': 'Interview Transcription â€” 3 Pack',
        'description': '3 interview transcription credits',
        'credits': {'interview_credits_short': 3},
        'price_cents': 799,
        'currency': 'eur',
    },
    'interview_8': {
        'name': 'Interview Transcription â€” 8 Pack',
        'description': '8 interview transcription credits (best value)',
        'credits': {'interview_credits_short': 8},
        'price_cents': 1799,
        'currency': 'eur',
    },
}

# --- Email Allowlist ---
EMAIL_ALLOWLIST_CONFIG_PATH = os.path.join(
    PROJECT_ROOT_DIR,
    'config',
    'allowed_email_domains.json',
)


def load_email_allowlist_config(path):
    try:
        with open(path, 'r', encoding='utf-8') as handle:
            data = json.load(handle)
    except Exception as e:
        raise RuntimeError(f"Could not read allowlist config at {path}: {e}")
    if not isinstance(data, dict):
        raise RuntimeError(f"Allowlist config at {path} must be a JSON object.")
    raw_domains = data.get('domains', [])
    raw_suffixes = data.get('suffixes', [])
    if not isinstance(raw_domains, list) or not isinstance(raw_suffixes, list):
        raise RuntimeError(f"Allowlist config at {path} must contain list values for 'domains' and 'suffixes'.")
    domains = {str(item).strip().lower() for item in raw_domains if str(item).strip()}
    suffixes = [str(item).strip().lower() for item in raw_suffixes if str(item).strip()]
    if not domains:
        raise RuntimeError(f"Allowlist config at {path} has an empty domains list.")
    if not suffixes:
        raise RuntimeError(f"Allowlist config at {path} has an empty suffixes list.")
    return domains, suffixes


ALLOWED_EMAIL_DOMAINS, ALLOWED_EMAIL_PATTERNS = load_email_allowlist_config(EMAIL_ALLOWLIST_CONFIG_PATH)

def is_email_allowed(email):
    if not email: return False
    email = email.lower()
    domain = email.split('@')[-1] if '@' in email else ''
    if domain in ALLOWED_EMAIL_DOMAINS: return True
    for pattern in ALLOWED_EMAIL_PATTERNS:
        if domain.endswith(pattern): return True
    return False

# --- AI Model Config ---
MODEL_SLIDES = 'gemini-2.5-flash-lite'
MODEL_AUDIO = 'gemini-3-flash-preview'
MODEL_INTEGRATION = 'gemini-2.5-pro'
MODEL_INTERVIEW = 'gemini-2.5-pro'
MODEL_STUDY = 'gemini-2.5-flash-lite'

FREE_LECTURE_CREDITS = 1
FREE_SLIDES_CREDITS = 2
FREE_INTERVIEW_CREDITS = 0

OUTPUT_LANGUAGE_MAP = {
    'dutch': 'Dutch',
    'english': 'English',
    'spanish': 'Spanish',
    'french': 'French',
    'german': 'German',
    'chinese': 'Chinese',
}
DEFAULT_OUTPUT_LANGUAGE_KEY = 'english'
OUTPUT_LANGUAGE_KEYS = set(OUTPUT_LANGUAGE_MAP.keys()) | {'other'}
MAX_OUTPUT_LANGUAGE_CUSTOM_LENGTH = 40
VIDEO_IMPORT_MAX_URL_LENGTH = 4096
VIDEO_IMPORT_ALLOWED_HOST_SUFFIXES = tuple(
    part.strip().lower()
    for part in (os.getenv('VIDEO_IMPORT_ALLOWED_HOST_SUFFIXES', 'kaltura.com,ovp.kaltura.com,brightspace.com,d2l.com') or '').split(',')
    if part.strip()
)

# --- Prompts ---
PROMPT_REGISTRY_VERSION = prompt_registry.PROMPT_REGISTRY_VERSION
PROMPT_SLIDE_EXTRACTION = prompt_registry.PROMPT_SLIDE_EXTRACTION
PROMPT_AUDIO_TRANSCRIPTION = prompt_registry.PROMPT_AUDIO_TRANSCRIPTION
PROMPT_AUDIO_TRANSCRIPTION_TIMESTAMPED = prompt_registry.PROMPT_AUDIO_TRANSCRIPTION_TIMESTAMPED
PROMPT_INTERVIEW_TRANSCRIPTION = prompt_registry.PROMPT_INTERVIEW_TRANSCRIPTION
PROMPT_INTERVIEW_SUMMARY = prompt_registry.PROMPT_INTERVIEW_SUMMARY
PROMPT_INTERVIEW_SECTIONED = prompt_registry.PROMPT_INTERVIEW_SECTIONED
PROMPT_MERGE_TEMPLATE = prompt_registry.PROMPT_MERGE_TEMPLATE
PROMPT_MERGE_WITH_AUDIO_MARKERS = prompt_registry.PROMPT_MERGE_WITH_AUDIO_MARKERS
PROMPT_STUDY_TEMPLATE = prompt_registry.PROMPT_STUDY_TEMPLATE


def get_prompt_inventory():
    return prompt_registry.get_prompt_inventory()


def get_prompt_inventory_markdown():
    return prompt_registry.get_prompt_inventory_markdown()

# =============================================
# FIRESTORE USER FUNCTIONS
# =============================================

JOB_TTL_SECONDS = safe_int_env('JOB_TTL_SECONDS', 30 * 60, minimum=5 * 60, maximum=24 * 60 * 60)
JOB_CLEANUP_INTERVAL_SECONDS = 5 * 60  # run cleanup every 5 minutes
PROCESSING_PUBLIC_ERROR_MESSAGE = 'Processing failed. Your credit has been refunded.'

def cleanup_old_jobs():
    """Evict completed/errored jobs older than JOB_TTL_SECONDS to prevent OOM."""
    now_ts = time.time()
    expired_ids = []
    with JOBS_LOCK:
        for job_id, job in list(jobs.items()):
            status = job.get('status', '')
            if status not in ('complete', 'error'):
                continue
            finished_at = job.get('finished_at', job.get('started_at', now_ts))
            if now_ts - finished_at > JOB_TTL_SECONDS:
                expired_ids.append(job_id)
        for job_id in expired_ids:
            jobs.pop(job_id, None)
    for job_id in expired_ids:
        delete_runtime_job_snapshot(job_id)

def cleanup_expired_audio_stream_tokens():
    """Evict expired audio stream tokens to prevent unbounded memory growth."""
    now_ts = time.time()
    expired = [t for t, d in list(AUDIO_STREAM_TOKENS.items()) if now_ts > d.get('expires_at', 0)]
    for token in expired:
        AUDIO_STREAM_TOKENS.pop(token, None)

def _run_periodic_cleanup():
    """Background thread: periodically evict stale jobs and audio stream tokens."""
    while True:
        time.sleep(JOB_CLEANUP_INTERVAL_SECONDS)
        try:
            cleanup_old_jobs()
            cleanup_expired_audio_stream_tokens()
        except Exception:
            logger.warning("Periodic cleanup failed", exc_info=True)

_cleanup_thread = threading.Thread(target=_run_periodic_cleanup, daemon=True)
_cleanup_thread.start()

def build_default_user_data(uid, email):
    """Return the canonical default user document structure."""
    return {
        'uid': uid,
        'email': email,
        'lecture_credits_standard': FREE_LECTURE_CREDITS,
        'lecture_credits_extended': 0,
        'slides_credits': FREE_SLIDES_CREDITS,
        'interview_credits_short': FREE_INTERVIEW_CREDITS,
        'interview_credits_medium': 0,
        'interview_credits_long': 0,
        'total_processed': 0,
        'created_at': time.time(),
        'preferred_output_language': DEFAULT_OUTPUT_LANGUAGE_KEY,
        'preferred_output_language_custom': '',
        'onboarding_completed': False,
    }

def get_or_create_user(uid, email):
    """Get a user from Firestore, or create them with free credits if they don't exist."""
    user_ref = users_repo.doc_ref(db, uid)
    user_doc = user_ref.get()

    if user_doc.exists:
        user_data = user_doc.to_dict()
        updates = {}
        # Update email if it changed
        if user_data.get('email') != email and email:
            updates['email'] = email
        pref_key = sanitize_output_language_pref_key(user_data.get('preferred_output_language', DEFAULT_OUTPUT_LANGUAGE_KEY))
        pref_custom = sanitize_output_language_pref_custom(user_data.get('preferred_output_language_custom', ''))
        if pref_key != str(user_data.get('preferred_output_language', '') or '').strip().lower():
            updates['preferred_output_language'] = pref_key
        if pref_key != 'other':
            pref_custom = ''
        if pref_custom != str(user_data.get('preferred_output_language_custom', '') or '').strip():
            updates['preferred_output_language_custom'] = pref_custom
        if not isinstance(user_data.get('onboarding_completed'), bool):
            updates['onboarding_completed'] = False
        if updates:
            user_ref.update(updates)
            user_data.update(updates)
        return user_data
    else:
        # New user â€” create with free credits
        user_data = build_default_user_data(uid, email)
        user_ref.set(user_data)
        logger.info(f"New user created: {uid} ({email})")
        return user_data

def grant_credits_to_user(uid, bundle_id):
    """Grant credits from a purchased bundle to a user in Firestore."""
    bundle = CREDIT_BUNDLES.get(bundle_id)
    if not bundle:
        logger.warning(f"Warning: Unknown bundle_id '{bundle_id}' in grant_credits_to_user")
        return False

    user_ref = users_repo.doc_ref(db, uid)
    user_doc = user_ref.get()

    if not user_doc.exists:
        # User not in Firestore yet â€” create with defaults first
        user_data = build_default_user_data(uid, '')
        user_ref.set(user_data)

    # Add the purchased credits
    for credit_key, credit_amount in bundle['credits'].items():
        user_ref.update({credit_key: firestore.Increment(credit_amount)})
        logger.info(f"Granted {credit_amount} '{credit_key}' credits to user {uid}.")

    return True

def deduct_credit(uid, credit_type_primary, credit_type_fallback=None):
    """Deduct one credit atomically using a Firestore transaction. Returns the credit type deducted, or None."""
    @firestore.transactional
    def _deduct_in_transaction(transaction, user_ref):
        snapshot = user_ref.get(transaction=transaction)
        if not snapshot.exists:
            return None
        data = snapshot.to_dict()
        if data.get(credit_type_primary, 0) > 0:
            transaction.update(user_ref, {
                credit_type_primary: firestore.Increment(-1),
                'total_processed': firestore.Increment(1),
            })
            return credit_type_primary
        elif credit_type_fallback and data.get(credit_type_fallback, 0) > 0:
            transaction.update(user_ref, {
                credit_type_fallback: firestore.Increment(-1),
                'total_processed': firestore.Increment(1),
            })
            return credit_type_fallback
        return None

    user_ref = users_repo.doc_ref(db, uid)
    transaction = db.transaction()
    return _deduct_in_transaction(transaction, user_ref)

def deduct_interview_credit(uid):
    """Deduct one interview credit atomically, checking short -> medium -> long. Returns the credit type deducted, or None."""
    @firestore.transactional
    def _deduct_in_transaction(transaction, user_ref):
        snapshot = user_ref.get(transaction=transaction)
        if not snapshot.exists:
            return None
        data = snapshot.to_dict()
        for credit_type in ('interview_credits_short', 'interview_credits_medium', 'interview_credits_long'):
            if data.get(credit_type, 0) > 0:
                transaction.update(user_ref, {
                    credit_type: firestore.Increment(-1),
                    'total_processed': firestore.Increment(1),
                })
                return credit_type
        return None

    user_ref = users_repo.doc_ref(db, uid)
    transaction = db.transaction()
    return _deduct_in_transaction(transaction, user_ref)

def refund_credit(uid, credit_type):
    """Refund one credit back to the user after a failed processing job."""
    if not uid or not credit_type:
        return
    try:
        user_ref = users_repo.doc_ref(db, uid)
        user_ref.update({
            credit_type: firestore.Increment(1),
            'total_processed': firestore.Increment(-1),
        })
        logger.info(f"âœ… Refunded 1 '{credit_type}' credit to user {uid} due to processing failure.")
    except Exception as e:
        logger.error(f"âŒ Failed to refund credit '{credit_type}' to user {uid}: {e}")

def save_purchase_record(uid, bundle_id, stripe_session_id):
    """Save a purchase record to Firestore for purchase history."""
    bundle = CREDIT_BUNDLES.get(bundle_id)
    if not bundle:
        return
    try:
        record = {
            'uid': uid,
            'bundle_id': bundle_id,
            'bundle_name': bundle['name'],
            'price_cents': bundle['price_cents'],
            'currency': bundle['currency'],
            'credits': bundle['credits'],
            'stripe_session_id': stripe_session_id,
            'created_at': time.time(),
        }
        if stripe_session_id:
            purchases_repo.set_doc(db, stripe_session_id, record, merge=True)
        else:
            purchases_repo.add_doc(db, record)
        logger.info(f"ðŸ“ Saved purchase record for user {uid}: {bundle['name']}")
    except Exception as e:
        logger.error(f"âŒ Failed to save purchase record for user {uid}: {e}")

def purchase_record_exists_for_session(stripe_session_id):
    if not stripe_session_id:
        return False
    try:
        doc = purchases_repo.get_doc(db, stripe_session_id)
        if doc.exists:
            return True
        for _ in purchases_repo.query_by_session_id(db, stripe_session_id, limit=1):
            return True
        return False
    except Exception as e:
        logger.warning(f"âš ï¸ Could not check purchase record for session {stripe_session_id}: {e}")
        return False

def process_checkout_session_credits(stripe_session):
    metadata = stripe_session.get('metadata', {}) or {}
    uid = metadata.get('uid', '')
    bundle_id = metadata.get('bundle_id', '')
    stripe_session_id = stripe_session.get('id', '')
    payment_status = (stripe_session.get('payment_status') or '').lower()
    session_status = (stripe_session.get('status') or '').lower()

    if not uid or not bundle_id:
        return False, 'Missing checkout metadata.'
    if bundle_id not in CREDIT_BUNDLES:
        return False, 'Unknown credit bundle.'
    if payment_status != 'paid' and session_status != 'complete':
        return False, 'Checkout session is not paid yet.'
    if purchase_record_exists_for_session(stripe_session_id):
        return True, 'already_processed'

    success = grant_credits_to_user(uid, bundle_id)
    if not success:
        return False, 'Could not grant credits.'
    save_purchase_record(uid, bundle_id, stripe_session_id)
    bundle = CREDIT_BUNDLES.get(bundle_id, {})
    log_analytics_event(
        'payment_confirmed_backend',
        source='backend',
        uid=uid,
        session_id=stripe_session_id,
        properties={
            'bundle_id': bundle_id,
            'price_cents': int(bundle.get('price_cents', 0) or 0),
        }
    )
    return True, 'granted'

def sanitize_analytics_event_name(raw_name):
    return analytics_service.sanitize_event_name(
        raw_name,
        name_re=ANALYTICS_NAME_RE,
        allowed_events=ANALYTICS_ALLOWED_EVENTS,
    )

def sanitize_analytics_session_id(raw_session_id):
    return analytics_service.sanitize_session_id(
        raw_session_id,
        session_id_re=ANALYTICS_SESSION_ID_RE,
    )

def sanitize_analytics_properties(raw_props):
    return analytics_service.sanitize_properties(raw_props, name_re=ANALYTICS_NAME_RE)

def log_analytics_event(event_name, source='frontend', uid='', email='', session_id='', properties=None, created_at=None):
    return analytics_service.log_analytics_event(
        event_name,
        source=source,
        uid=uid,
        email=email,
        session_id=session_id,
        properties=properties,
        created_at=created_at,
        db=db,
        name_re=ANALYTICS_NAME_RE,
        session_id_re=ANALYTICS_SESSION_ID_RE,
        allowed_events=ANALYTICS_ALLOWED_EVENTS,
        logger=logger,
        time_module=time,
    )

def log_rate_limit_hit(limit_name, retry_after=0):
    return analytics_service.log_rate_limit_hit(
        limit_name,
        retry_after=retry_after,
        db=db,
        logger=logger,
        time_module=time,
    )

def save_job_log(job_id, job_data, finished_at):
    """Save a processing job log to Firestore for analytics."""
    try:
        started_at = job_data.get('started_at', 0)
        duration = round(finished_at - started_at, 1) if started_at else 0
        job_logs_repo.set_job_log(db, job_id, {
            'job_id': job_id,
            'uid': job_data.get('user_id', ''),
            'email': job_data.get('user_email', ''),
            'mode': job_data.get('mode', ''),
            'status': job_data.get('status', ''),
            'credit_deducted': job_data.get('credit_deducted', ''),
            'credit_refunded': job_data.get('credit_refunded', False),
            'error_message': job_data.get('error', ''),
            'started_at': started_at,
            'finished_at': finished_at,
            'duration_seconds': duration,
        })
        status = str(job_data.get('status', '') or '').lower()
        backend_event = 'processing_finished_backend'
        if status == 'complete':
            backend_event = 'processing_completed_backend'
        elif status == 'error':
            backend_event = 'processing_failed_backend'
        log_analytics_event(
            backend_event,
            source='backend',
            uid=job_data.get('user_id', ''),
            email=job_data.get('user_email', ''),
            session_id=job_id,
            properties={
                'job_id': job_id,
                'mode': job_data.get('mode', ''),
                'duration_seconds': duration,
                'credit_refunded': bool(job_data.get('credit_refunded', False)),
            },
            created_at=finished_at,
        )
        logger.info(f"ðŸ“Š Logged job {job_id}: mode={job_data.get('mode')}, status={job_data.get('status')}, duration={duration}s")
    except Exception as e:
        logger.error(f"âŒ Failed to log job {job_id}: {e}")


def recover_stale_runtime_jobs():
    """Recover jobs left in starting/processing state after a restart."""
    if db is None:
        return 0
    now_ts = time.time()
    recovered = 0
    try:
        stale_docs = runtime_jobs_repo.query_statuses(
            db,
            RUNTIME_JOBS_COLLECTION,
            {'starting', 'processing'},
            limit=RUNTIME_JOB_RECOVERY_BATCH_LIMIT,
        )
    except Exception:
        logger.warning("Runtime-job recovery query failed", exc_info=True)
        return 0

    for doc in stale_docs:
        job_id = doc.id
        job_data = doc.to_dict() or {}
        if not isinstance(job_data, dict):
            continue
        status = str(job_data.get('status', '') or '').lower()
        if status not in {'starting', 'processing'}:
            continue

        uid = str(job_data.get('user_id', '') or '').strip()
        credit_type = str(job_data.get('credit_deducted', '') or '').strip()
        already_refunded = bool(job_data.get('credit_refunded', False))

        if uid and credit_type and not already_refunded:
            refund_credit(uid, credit_type)
            add_job_credit_refund(job_data, credit_type, 1)
            job_data['credit_refunded'] = True

        extra_spent = int(job_data.get('interview_features_cost', 0) or 0)
        extra_refunded = int(job_data.get('extra_slides_refunded', 0) or 0)
        extra_to_refund = max(0, extra_spent - extra_refunded)
        if uid and extra_to_refund > 0:
            refund_slides_credits(uid, extra_to_refund)
            job_data['extra_slides_refunded'] = extra_refunded + extra_to_refund
            add_job_credit_refund(job_data, 'slides_credits', extra_to_refund)

        ensure_job_billing_receipt(job_data, {credit_type: 1} if credit_type else None)
        job_data['status'] = 'error'
        job_data['step_description'] = 'Interrupted by server restart'
        job_data['error'] = 'Processing was interrupted by a server restart. Your credit has been refunded.'
        job_data['finished_at'] = now_ts
        job_data['job_id'] = job_id

        set_job(job_id, job_data)
        save_job_log(job_id, job_data, now_ts)
        recovered += 1

    if recovered:
        logger.warning("Recovered %s stale runtime jobs after startup.", recovered)
    return recovered

# =============================================
# HELPER FUNCTIONS
# =============================================

def verify_firebase_token(request):
    return auth_service.verify_firebase_token(request, auth_module=auth, logger=logger)

def is_admin_user(decoded_token):
    if not decoded_token:
        return False
    uid = decoded_token.get('uid', '')
    email = decoded_token.get('email', '').lower()
    return uid in ADMIN_UIDS or email in ADMIN_EMAILS


def _extract_bearer_token(req):
    auth_header = req.headers.get('Authorization', '')
    if isinstance(auth_header, str) and auth_header.startswith('Bearer '):
        token = auth_header.split('Bearer ', 1)[1].strip()
        if token:
            return token
    payload = req.get_json(silent=True) or {}
    body_token = str(payload.get('id_token', '') or payload.get('idToken', '') or '').strip()
    if body_token:
        return body_token
    return ''


def verify_admin_session_cookie(req):
    session_cookie = req.cookies.get(ADMIN_SESSION_COOKIE_NAME, '')
    if not session_cookie:
        return None
    try:
        decoded_token = auth.verify_session_cookie(session_cookie, check_revoked=True)
    except Exception:
        return None
    if not is_admin_user(decoded_token):
        return None
    return decoded_token

def get_admin_window(window_key):
    windows = {
        '24h': 24 * 60 * 60,
        '7d': 7 * 24 * 60 * 60,
        '30d': 30 * 24 * 60 * 60,
    }
    safe_key = window_key if window_key in windows else '7d'
    return safe_key, windows[safe_key]

def get_timestamp(value):
    return value if isinstance(value, (int, float)) else 0

def build_time_buckets(window_key, now_ts):
    labels = []
    keys = []
    if window_key == '24h':
        now_dt = datetime.fromtimestamp(now_ts, tz=timezone.utc).replace(minute=0, second=0, microsecond=0)
        start_dt = now_dt - timedelta(hours=23)
        for i in range(24):
            current = start_dt + timedelta(hours=i)
            labels.append(current.strftime('%H:%M'))
            keys.append(current.strftime('%Y-%m-%d %H:00'))
        granularity = 'hour'
    else:
        days = 7 if window_key == '7d' else 30
        now_dt = datetime.fromtimestamp(now_ts, tz=timezone.utc).replace(hour=0, minute=0, second=0, microsecond=0)
        start_dt = now_dt - timedelta(days=days - 1)
        for i in range(days):
            current = start_dt + timedelta(days=i)
            labels.append(current.strftime('%d %b'))
            keys.append(current.strftime('%Y-%m-%d'))
        granularity = 'day'
    return labels, keys, granularity

def get_bucket_key(timestamp, window_key):
    dt = datetime.fromtimestamp(timestamp, tz=timezone.utc)
    if window_key == '24h':
        return dt.replace(minute=0, second=0, microsecond=0).strftime('%Y-%m-%d %H:00')
    return dt.replace(hour=0, minute=0, second=0, microsecond=0).strftime('%Y-%m-%d')


def query_docs_in_window(collection_name, timestamp_field, window_start, window_end=None, order_desc=False, limit=None):
    return admin_repo.query_docs_in_window(
        db,
        collection_name=collection_name,
        timestamp_field=timestamp_field,
        window_start=window_start,
        window_end=window_end,
        order_desc=order_desc,
        limit=limit,
        firestore_module=firestore,
    )


def safe_query_docs_in_window(collection_name, timestamp_field, window_start, window_end=None, order_desc=False, limit=None):
    if db is None:
        return []
    try:
        return query_docs_in_window(
            collection_name=collection_name,
            timestamp_field=timestamp_field,
            window_start=window_start,
            window_end=window_end,
            order_desc=order_desc,
            limit=limit,
        )
    except Exception:
        # Fallback for missing indexes in early environments.
        fallback_cap = safe_int_env('ADMIN_FALLBACK_QUERY_CAP', 5000, minimum=100, maximum=50000)
        logger.warning(
            "Using fallback in-memory query for %s due to index/query issue; cap=%s",
            collection_name,
            fallback_cap,
            exc_info=True,
        )
        docs = []
        for doc in admin_repo.stream_collection(db, collection_name):
            data = doc.to_dict() or {}
            ts = get_timestamp(data.get(timestamp_field))
            if ts < window_start:
                continue
            if window_end is not None and ts > window_end:
                continue
            docs.append(doc)
            if len(docs) >= fallback_cap:
                break
        docs.sort(
            key=lambda d: get_timestamp((d.to_dict() or {}).get(timestamp_field)),
            reverse=order_desc,
        )
        if isinstance(limit, int) and limit > 0:
            docs = docs[:limit]
        return docs


def safe_count_collection(collection_name):
    if db is None:
        return 0
    try:
        return admin_repo.count_collection(db, collection_name)
    except Exception:
        pass
    return len(list(admin_repo.stream_collection(db, collection_name)))


def safe_count_window(collection_name, timestamp_field, window_start):
    if db is None:
        return 0
    try:
        return admin_repo.count_window(db, collection_name, timestamp_field, window_start)
    except Exception:
        pass
    docs = safe_query_docs_in_window(collection_name, timestamp_field, window_start)
    return len(docs)

def build_admin_funnel_steps(analytics_docs, window_start):
    funnel_actor_sets = {stage['event']: set() for stage in ANALYTICS_FUNNEL_STAGES}
    analytics_event_count = 0

    for doc in analytics_docs:
        event = doc.to_dict() or {}
        created_at = get_timestamp(event.get('created_at'))
        if created_at < window_start:
            continue
        event_name = sanitize_analytics_event_name(event.get('event', ''))
        if not event_name:
            continue
        analytics_event_count += 1
        if event_name not in funnel_actor_sets:
            continue
        uid = str(event.get('uid', '') or '').strip()
        session_id = sanitize_analytics_session_id(event.get('session_id', ''))
        actor_id = uid or session_id or f"doc:{doc.id}"
        funnel_actor_sets[event_name].add(actor_id)

    funnel_steps = []
    previous_count = 0
    for idx, stage in enumerate(ANALYTICS_FUNNEL_STAGES):
        count = len(funnel_actor_sets.get(stage['event'], set()))
        if idx == 0:
            conversion = 100.0 if count > 0 else 0.0
        elif previous_count > 0:
            conversion = round(min((count / previous_count) * 100.0, 100.0), 1)
        else:
            conversion = 0.0
        funnel_steps.append({
            'event': stage['event'],
            'label': stage['label'],
            'count': count,
            'conversion_from_prev': conversion,
        })
        previous_count = count

    return funnel_steps, analytics_event_count

def build_admin_funnel_daily_rows(analytics_docs, window_start, window_key, now_ts):
    _labels, bucket_keys, granularity = build_time_buckets(window_key, now_ts)
    counts_by_bucket = {}

    for doc in analytics_docs:
        event = doc.to_dict() or {}
        created_at = get_timestamp(event.get('created_at'))
        if created_at < window_start or created_at > now_ts:
            continue
        event_name = sanitize_analytics_event_name(event.get('event', ''))
        if event_name not in ANALYTICS_FUNNEL_EVENT_NAMES:
            continue

        bucket_key = get_bucket_key(created_at, window_key)
        if bucket_key not in counts_by_bucket:
            counts_by_bucket[bucket_key] = {}
        if event_name not in counts_by_bucket[bucket_key]:
            counts_by_bucket[bucket_key][event_name] = {'event_count': 0, 'actors': set()}

        uid = str(event.get('uid', '') or '').strip()
        session_id = sanitize_analytics_session_id(event.get('session_id', ''))
        actor_id = uid or session_id or f"doc:{doc.id}"

        counts_by_bucket[bucket_key][event_name]['event_count'] += 1
        counts_by_bucket[bucket_key][event_name]['actors'].add(actor_id)

    rows = []
    for bucket_key in bucket_keys:
        stage_counts = counts_by_bucket.get(bucket_key, {})
        prev_unique = 0
        for idx, stage in enumerate(ANALYTICS_FUNNEL_STAGES):
            stage_data = stage_counts.get(stage['event'], {'event_count': 0, 'actors': set()})
            unique_actor_count = len(stage_data.get('actors', set()))
            event_count = int(stage_data.get('event_count', 0) or 0)
            if idx == 0:
                conversion = 100.0 if unique_actor_count > 0 else 0.0
            elif prev_unique > 0:
                conversion = round(min((unique_actor_count / prev_unique) * 100.0, 100.0), 1)
            else:
                conversion = 0.0

            rows.append({
                'bucket_key': bucket_key,
                'granularity': granularity,
                'event': stage['event'],
                'label': stage['label'],
                'unique_actor_count': unique_actor_count,
                'event_count': event_count,
                'conversion_from_prev': conversion,
            })
            prev_unique = unique_actor_count

    return rows, granularity

def _runtime_job_storage_enabled():
    return db is not None


def _runtime_job_sanitize_value(value):
    if isinstance(value, str):
        if len(value) > RUNTIME_JOB_MAX_STRING_LENGTH:
            return value[:RUNTIME_JOB_MAX_STRING_LENGTH]
        return value
    if isinstance(value, list):
        return [_runtime_job_sanitize_value(v) for v in value]
    if isinstance(value, dict):
        return {str(k): _runtime_job_sanitize_value(v) for k, v in value.items()}
    return value


def _build_runtime_job_payload(job_id, job_data):
    payload = {'job_id': job_id, 'updated_at': time.time()}
    if not isinstance(job_data, dict):
        payload['status'] = 'unknown'
        return payload
    for field in RUNTIME_JOB_PERSISTED_FIELDS:
        if field in job_data:
            payload[field] = _runtime_job_sanitize_value(job_data.get(field))
    return payload


def persist_runtime_job_snapshot(job_id, job_data):
    if not _runtime_job_storage_enabled() or not job_id:
        return
    try:
        runtime_jobs_repo.set_doc(
            db,
            RUNTIME_JOBS_COLLECTION,
            job_id,
            _build_runtime_job_payload(job_id, job_data),
            merge=True,
        )
    except Exception:
        logger.warning("Failed to persist runtime job snapshot for %s", job_id, exc_info=True)


def load_runtime_job_snapshot(job_id):
    if not _runtime_job_storage_enabled() or not job_id:
        return None
    try:
        doc = runtime_jobs_repo.get_doc(db, RUNTIME_JOBS_COLLECTION, job_id)
        if not doc.exists:
            return None
        data = doc.to_dict() or {}
        if not isinstance(data, dict):
            return None
        data.setdefault('job_id', job_id)
        return data
    except Exception:
        logger.warning("Failed to load runtime job snapshot for %s", job_id, exc_info=True)
        return None


def delete_runtime_job_snapshot(job_id):
    if not _runtime_job_storage_enabled() or not job_id:
        return
    try:
        runtime_jobs_repo.delete_doc(db, RUNTIME_JOBS_COLLECTION, job_id)
    except Exception:
        logger.warning("Failed to delete runtime job snapshot for %s", job_id, exc_info=True)


def update_job_fields(job_id, **fields):
    if not fields:
        return get_job_snapshot(job_id)

    def _mutator(job):
        job.update(fields)

    snapshot = mutate_job(job_id, _mutator)
    return snapshot


def get_job_snapshot(job_id):
    snapshot = job_state_service.get_job_snapshot(job_id, jobs_store=jobs, lock=JOBS_LOCK)
    if snapshot is not None:
        return snapshot
    runtime_snapshot = load_runtime_job_snapshot(job_id)
    if runtime_snapshot is not None:
        # Warm the in-memory cache for subsequent reads.
        job_state_service.set_job(job_id, dict(runtime_snapshot), jobs_store=jobs, lock=JOBS_LOCK)
        return runtime_snapshot
    return None


def mutate_job(job_id, mutator_fn):
    snapshot = job_state_service.mutate_job(job_id, mutator_fn, jobs_store=jobs, lock=JOBS_LOCK)
    if snapshot is not None:
        persist_runtime_job_snapshot(job_id, snapshot)
    return snapshot


def set_job(job_id, value):
    snapshot = job_state_service.set_job(job_id, value, jobs_store=jobs, lock=JOBS_LOCK)
    if isinstance(snapshot, dict):
        persist_runtime_job_snapshot(job_id, snapshot)
    return snapshot


def delete_job(job_id):
    deleted = job_state_service.delete_job(job_id, jobs_store=jobs, lock=JOBS_LOCK)
    delete_runtime_job_snapshot(job_id)
    return deleted


def _window_counter_id(key, window_seconds, window_start):
    return rate_limit_service.window_counter_id(key, window_seconds, window_start)


def _check_rate_limit_firestore(key, limit, window_seconds, now_ts):
    return rate_limit_service.check_rate_limit_firestore(
        key,
        limit,
        window_seconds,
        now_ts,
        firestore_enabled=RATE_LIMIT_FIRESTORE_ENABLED,
        db=db,
        firestore_module=firestore,
        counter_collection=RATE_LIMIT_COUNTER_COLLECTION,
    )


def check_rate_limit(key, limit, window_seconds):
    return rate_limit_service.check_rate_limit(
        key,
        limit,
        window_seconds,
        firestore_enabled=RATE_LIMIT_FIRESTORE_ENABLED,
        db=db,
        firestore_module=firestore,
        counter_collection=RATE_LIMIT_COUNTER_COLLECTION,
        in_memory_events=RATE_LIMIT_EVENTS,
        in_memory_lock=RATE_LIMIT_LOCK,
        time_module=time,
    )

def build_rate_limited_response(message, retry_after):
    response = jsonify({
        'error': message,
        'retry_after_seconds': int(max(1, retry_after)),
    })
    response.status_code = 429
    response.headers['Retry-After'] = str(int(max(1, retry_after)))
    return response

def normalize_rate_limit_key_part(value, fallback='anon', max_len=120):
    raw = str(value or '').strip().lower()
    if not raw:
        return fallback
    safe = re.sub(r'[^a-z0-9_.:@-]+', '_', raw)
    return safe[:max_len] if safe else fallback


def has_sufficient_upload_disk_space(required_bytes=0):
    """Return (ok, free_bytes, threshold_bytes) for upload safety checks."""
    try:
        usage = shutil.disk_usage(UPLOAD_FOLDER if os.path.exists(UPLOAD_FOLDER) else '/')
        free_bytes = int(usage.free)
    except Exception:
        return True, 0, UPLOAD_MIN_FREE_DISK_BYTES
    try:
        required = int(required_bytes or 0)
    except Exception:
        required = 0
    needed = max(UPLOAD_MIN_FREE_DISK_BYTES, required + UPLOAD_MIN_FREE_DISK_BYTES)
    return free_bytes >= needed, free_bytes, needed


def reserve_daily_upload_bytes(uid, requested_bytes):
    """Atomically reserve upload bytes for a user in the current UTC day."""
    if db is None:
        return True, 0
    if not uid:
        return False, 0
    try:
        requested = int(requested_bytes or 0)
    except Exception:
        requested = 0
    requested = max(0, requested)
    if requested <= 0:
        return True, 0
    now_ts = time.time()
    day_key = datetime.fromtimestamp(now_ts, tz=timezone.utc).strftime('%Y-%m-%d')
    doc_id = f"{uid}:{day_key}"
    retry_after = max(1, int((datetime.fromtimestamp(now_ts, tz=timezone.utc).replace(hour=23, minute=59, second=59, microsecond=0).timestamp()) - now_ts))
    counter_ref = db.collection(UPLOAD_DAILY_COUNTER_COLLECTION).document(doc_id)
    transaction = db.transaction()

    @firestore.transactional
    def _txn(txn):
        snapshot = counter_ref.get(transaction=txn)
        used = 0
        if snapshot.exists:
            used = int((snapshot.to_dict() or {}).get('bytes_used', 0) or 0)
        if used + requested > UPLOAD_DAILY_BYTE_CAP:
            return False, retry_after
        txn.set(counter_ref, {
            'uid': uid,
            'day': day_key,
            'bytes_used': used + requested,
            'updated_at': now_ts,
            'expires_at': now_ts + (3 * 24 * 60 * 60),
        }, merge=True)
        return True, 0

    try:
        return _txn(transaction)
    except Exception:
        logger.warning("Upload daily byte reservation failed for uid=%s", uid, exc_info=True)
        return True, 0


def count_active_jobs_for_user(uid):
    return job_state_service.count_active_jobs_for_user(uid, jobs_store=jobs, lock=JOBS_LOCK)

def list_docs_by_uid(collection_name, uid, max_docs):
    docs = admin_repo.query_by_uid(db, collection_name, uid, max_docs + 1)
    truncated = len(docs) > max_docs
    limited = docs[:max_docs]
    records = []
    for doc in limited:
        data = doc.to_dict() or {}
        data['_id'] = doc.id
        records.append(data)
    return records, truncated

def delete_docs_by_uid(collection_name, uid, max_docs):
    docs = admin_repo.query_by_uid(db, collection_name, uid, max_docs + 1)
    truncated = len(docs) > max_docs
    limited = docs[:max_docs]
    deleted = 0
    for doc in limited:
        try:
            doc.reference.delete()
            deleted += 1
        except Exception as e:
            logger.warning(f"Warning: could not delete doc in {collection_name}/{doc.id}: {e}")
    return deleted, truncated

def remove_upload_artifacts_for_job_ids(job_ids):
    if not job_ids:
        return 0
    try:
        names = os.listdir(UPLOAD_FOLDER)
    except Exception:
        return 0
    prefixes = tuple(f"{str(job_id).strip()}_" for job_id in job_ids if str(job_id or '').strip())
    if not prefixes:
        return 0
    removed = 0
    for name in names:
        if not name.startswith(prefixes):
            continue
        file_path = os.path.join(UPLOAD_FOLDER, name)
        try:
            if os.path.isfile(file_path):
                os.remove(file_path)
                removed += 1
        except Exception as e:
            logger.warning(f"Warning: could not delete upload artifact {file_path}: {e}")
    return removed

def anonymize_purchase_docs_by_uid(uid, max_docs):
    docs = purchases_repo.query_by_uid(db, uid, max_docs + 1)
    truncated = len(docs) > max_docs
    limited = docs[:max_docs]
    anonymized = 0
    for doc in limited:
        try:
            doc.reference.set({
                'uid': '',
                'user_erased': True,
                'erased_at': time.time(),
            }, merge=True)
            anonymized += 1
        except Exception as e:
            logger.warning(f"Warning: could not anonymize purchase doc {doc.id}: {e}")
    return anonymized, truncated

def collect_user_export_payload(uid, email):
    user_doc = users_repo.get_doc(db, uid)
    user_profile = user_doc.to_dict() if user_doc.exists else {}

    study_progress_doc = get_study_progress_doc(uid).get()
    study_progress = study_progress_doc.to_dict() if study_progress_doc.exists else {}

    purchases, purchases_truncated = list_docs_by_uid('purchases', uid, ACCOUNT_EXPORT_MAX_DOCS_PER_COLLECTION)
    job_logs, job_logs_truncated = list_docs_by_uid('job_logs', uid, ACCOUNT_EXPORT_MAX_DOCS_PER_COLLECTION)
    analytics_events, analytics_truncated = list_docs_by_uid('analytics_events', uid, ACCOUNT_EXPORT_MAX_DOCS_PER_COLLECTION)
    study_folders, folders_truncated = list_docs_by_uid('study_folders', uid, ACCOUNT_EXPORT_MAX_DOCS_PER_COLLECTION)
    study_packs, packs_truncated = list_docs_by_uid('study_packs', uid, ACCOUNT_EXPORT_MAX_DOCS_PER_COLLECTION)
    card_states, card_states_truncated = list_docs_by_uid('study_card_states', uid, ACCOUNT_EXPORT_MAX_DOCS_PER_COLLECTION)

    for pack in study_packs:
        audio_key = get_audio_storage_key_from_pack(pack)
        audio_path = resolve_audio_storage_path_from_key(audio_key) if audio_key else ''
        pack['audio_filename'] = os.path.basename(audio_path) if audio_path else ''
        pack.pop('audio_storage_path', None)
        pack.pop('audio_storage_key', None)

    return {
        'meta': {
            'exported_at': time.time(),
            'version': 1,
            'uid': uid,
            'email': email,
            'source': 'lecture-processor',
            'limits': {
                'max_docs_per_collection': ACCOUNT_EXPORT_MAX_DOCS_PER_COLLECTION,
            },
            'truncated': {
                'purchases': purchases_truncated,
                'job_logs': job_logs_truncated,
                'analytics_events': analytics_truncated,
                'study_folders': folders_truncated,
                'study_packs': packs_truncated,
                'study_card_states': card_states_truncated,
            },
        },
        'account': {
            'profile': user_profile,
            'study_progress': study_progress,
        },
        'collections': {
            'purchases': purchases,
            'job_logs': job_logs,
            'analytics_events': analytics_events,
            'study_folders': study_folders,
            'study_packs': study_packs,
            'study_card_states': card_states,
        },
    }

def parse_requested_amount(raw_value, allowed, default):
    value = str(raw_value or default).strip().lower()
    return value if value in allowed else default

def parse_study_features(raw_value):
    value = str(raw_value or 'none').strip().lower()
    return value if value in {'none', 'flashcards', 'test', 'both'} else 'none'

def normalize_output_language_choice(raw_value, custom_value=''):
    key = str(raw_value or DEFAULT_OUTPUT_LANGUAGE_KEY).strip().lower()
    custom = str(custom_value or '').strip()[:MAX_OUTPUT_LANGUAGE_CUSTOM_LENGTH]
    if key in OUTPUT_LANGUAGE_MAP:
        return key, '', OUTPUT_LANGUAGE_MAP[key]
    if key == 'other':
        if custom:
            return 'other', custom, custom
        return DEFAULT_OUTPUT_LANGUAGE_KEY, '', OUTPUT_LANGUAGE_MAP[DEFAULT_OUTPUT_LANGUAGE_KEY]
    return DEFAULT_OUTPUT_LANGUAGE_KEY, '', OUTPUT_LANGUAGE_MAP[DEFAULT_OUTPUT_LANGUAGE_KEY]

def parse_output_language(raw_value, custom_value=''):
    _key, _custom, resolved = normalize_output_language_choice(raw_value, custom_value)
    return resolved

def sanitize_output_language_pref_key(raw_value):
    key = str(raw_value or DEFAULT_OUTPUT_LANGUAGE_KEY).strip().lower()
    return key if key in OUTPUT_LANGUAGE_KEYS else DEFAULT_OUTPUT_LANGUAGE_KEY

def sanitize_output_language_pref_custom(raw_value):
    return str(raw_value or '').strip()[:MAX_OUTPUT_LANGUAGE_CUSTOM_LENGTH]

def build_user_preferences_payload(user_data):
    key, custom, resolved = normalize_output_language_choice(
        user_data.get('preferred_output_language', DEFAULT_OUTPUT_LANGUAGE_KEY),
        user_data.get('preferred_output_language_custom', ''),
    )
    return {
        'output_language': key,
        'output_language_custom': custom,
        'output_language_label': resolved,
        'onboarding_completed': bool(user_data.get('onboarding_completed', False)),
    }

def parse_interview_features(raw_value):
    value = str(raw_value or 'none').strip().lower()
    if value in {'none', ''}:
        return []
    if value == 'both':
        return ['summary', 'sections']
    parts = [part.strip() for part in value.split(',') if part.strip()]
    features = []
    for part in parts:
        if part in {'summary', 'sections'} and part not in features:
            features.append(part)
    return features

def host_matches_allowed_suffix(hostname):
    if not hostname:
        return False
    host = hostname.strip().lower()
    return any(host == suffix or host.endswith('.' + suffix) for suffix in VIDEO_IMPORT_ALLOWED_HOST_SUFFIXES)

def is_blocked_hostname(hostname):
    host = str(hostname or '').strip().lower()
    if not host:
        return True
    if host in {'localhost', 'localhost.localdomain'}:
        return True
    if host.endswith('.local') or host.endswith('.internal'):
        return True
    try:
        ip = ipaddress.ip_address(host)
        if ip.is_private or ip.is_loopback or ip.is_link_local or ip.is_multicast or ip.is_reserved:
            return True
    except ValueError:
        pass
    return False

def validate_video_import_url(raw_url):
    url = str(raw_url or '').strip()
    if not url:
        return '', 'Please paste a video URL.'
    if len(url) > VIDEO_IMPORT_MAX_URL_LENGTH:
        return '', 'Video URL is too long.'
    try:
        parsed = urlparse(url)
    except Exception:
        return '', 'Video URL is invalid.'
    if parsed.scheme.lower() != 'https':
        return '', 'Only HTTPS video URLs are supported.'
    if parsed.username or parsed.password:
        return '', 'Video URL credentials are not allowed.'
    host = (parsed.hostname or '').strip().lower()
    if not host:
        return '', 'Video URL host is missing.'
    if is_blocked_hostname(host):
        return '', 'This video host is not allowed.'
    if VIDEO_IMPORT_ALLOWED_HOST_SUFFIXES and not host_matches_allowed_suffix(host):
        return '', 'Only Brightspace/Kaltura video hosts are supported for automatic import.'
    # SSRF guard: pre-resolve DNS and check the resolved IP isn't private/internal (Issue 5)
    try:
        import socket
        resolved_ips = socket.getaddrinfo(host, 443, proto=socket.IPPROTO_TCP)
        for family, kind, proto, canonname, sockaddr in resolved_ips:
            ip_str = sockaddr[0]
            ip_obj = ipaddress.ip_address(ip_str)
            if ip_obj.is_private or ip_obj.is_loopback or ip_obj.is_link_local or ip_obj.is_multicast or ip_obj.is_reserved:
                return '', 'This video host resolves to a restricted network address.'
    except socket.gaierror:
        return '', 'Could not resolve the video URL host.'
    except Exception:
        pass
    return url, ''

def cleanup_expired_audio_import_tokens():
    now_ts = time.time()
    expired = []
    with AUDIO_IMPORT_LOCK:
        for token, data in list(AUDIO_IMPORT_TOKENS.items()):
            if now_ts > float(data.get('expires_at', 0) or 0):
                expired.append(data.get('path', ''))
                AUDIO_IMPORT_TOKENS.pop(token, None)
    for path in expired:
        try:
            if path and os.path.exists(path):
                os.remove(path)
        except Exception:
            pass

def register_audio_import_token(uid, file_path, source_url='', original_name=''):
    token = str(uuid.uuid4())
    with AUDIO_IMPORT_LOCK:
        AUDIO_IMPORT_TOKENS[token] = {
            'uid': str(uid or ''),
            'path': str(file_path or ''),
            'source_url': str(source_url or '')[:VIDEO_IMPORT_MAX_URL_LENGTH],
            'original_name': str(original_name or '')[:240],
            'created_at': time.time(),
            'expires_at': time.time() + AUDIO_IMPORT_TOKEN_TTL_SECONDS,
        }
    return token

def get_audio_import_token_path(uid, token, consume=False):
    cleanup_expired_audio_import_tokens()
    safe_uid = str(uid or '')
    safe_token = str(token or '').strip()
    if not safe_token:
        return '', 'Missing imported audio token.'
    with AUDIO_IMPORT_LOCK:
        entry = AUDIO_IMPORT_TOKENS.get(safe_token)
        if not entry:
            return '', 'Imported audio token expired or invalid. Please import again.'
        if entry.get('uid', '') != safe_uid:
            return '', 'Imported audio token does not belong to this account.'
        file_path = str(entry.get('path', '') or '').strip()
        if consume:
            AUDIO_IMPORT_TOKENS.pop(safe_token, None)
    if not file_path or not os.path.exists(file_path):
        return '', 'Imported audio file is no longer available. Please import again.'
    return file_path, ''

def release_audio_import_token(uid, token):
    safe_uid = str(uid or '')
    safe_token = str(token or '').strip()
    if not safe_token:
        return False
    file_path = ''
    with AUDIO_IMPORT_LOCK:
        entry = AUDIO_IMPORT_TOKENS.get(safe_token)
        if not entry or entry.get('uid', '') != safe_uid:
            return False
        file_path = str(entry.get('path', '') or '').strip()
        AUDIO_IMPORT_TOKENS.pop(safe_token, None)
    try:
        if file_path and os.path.exists(file_path):
            os.remove(file_path)
    except Exception:
        pass
    return True

def get_ffmpeg_binary():
    return file_service.get_ffmpeg_binary(which_func=shutil.which, imageio_ffmpeg_module=imageio_ffmpeg)

def get_ffprobe_binary():
    return file_service.get_ffprobe_binary(ffmpeg_binary_getter=get_ffmpeg_binary)

def download_audio_from_video_url(source_url, file_prefix):
    return file_service.download_audio_from_video_url(
        source_url,
        file_prefix,
        upload_folder=UPLOAD_FOLDER,
        max_audio_upload_bytes=MAX_AUDIO_UPLOAD_BYTES,
        ffmpeg_binary_getter=get_ffmpeg_binary,
        file_looks_like_audio_fn=file_looks_like_audio,
        get_saved_file_size_fn=get_saved_file_size,
        which_func=shutil.which,
        subprocess_module=subprocess,
    )

def deduct_slides_credits(uid, amount):
    """Deduct slides credits atomically using a Firestore transaction."""
    if amount <= 0:
        return True

    @firestore.transactional
    def _deduct_in_transaction(transaction, user_ref):
        snapshot = user_ref.get(transaction=transaction)
        if not snapshot.exists:
            return False
        data = snapshot.to_dict()
        current = data.get('slides_credits', 0)
        if current < amount:
            return False
        transaction.update(user_ref, {'slides_credits': firestore.Increment(-amount)})
        return True

    user_ref = users_repo.doc_ref(db, uid)
    transaction = db.transaction()
    return _deduct_in_transaction(transaction, user_ref)

def refund_slides_credits(uid, amount):
    if not uid or amount <= 0:
        return
    try:
        users_repo.update_doc(db, uid, {'slides_credits': firestore.Increment(amount)})
        logger.info(f"âœ… Refunded {amount} slides credits to user {uid}.")
    except Exception as e:
        logger.error(f"âŒ Failed to refund {amount} slides credits to user {uid}: {e}")

def normalize_credit_ledger(credit_map):
    normalized = {}
    if not isinstance(credit_map, dict):
        return normalized
    for credit_type, raw_amount in credit_map.items():
        key = str(credit_type or '').strip()
        if not key:
            continue
        try:
            amount = int(raw_amount)
        except Exception:
            continue
        if amount > 0:
            normalized[key] = amount
    return normalized

def initialize_billing_receipt(charged_map=None):
    return {
        'charged': normalize_credit_ledger(charged_map or {}),
        'refunded': {},
        'updated_at': time.time(),
    }

def ensure_job_billing_receipt(job_data, charged_map=None):
    receipt = job_data.get('billing_receipt')
    if not isinstance(receipt, dict):
        receipt = initialize_billing_receipt(charged_map or {})
        job_data['billing_receipt'] = receipt
        return receipt
    charged = receipt.get('charged', {})
    if not isinstance(charged, dict):
        charged = {}
    for credit_type, amount in normalize_credit_ledger(charged_map or {}).items():
        charged[credit_type] = max(int(charged.get(credit_type, 0) or 0), amount)
    receipt['charged'] = charged
    if not isinstance(receipt.get('refunded'), dict):
        receipt['refunded'] = {}
    receipt['updated_at'] = time.time()
    return receipt

def add_job_credit_refund(job_data, credit_type, amount=1):
    if not credit_type:
        return
    try:
        amount_int = int(amount)
    except Exception:
        return
    if amount_int <= 0:
        return
    receipt = ensure_job_billing_receipt(job_data)
    refunded = receipt.setdefault('refunded', {})
    refunded[credit_type] = int(refunded.get(credit_type, 0) or 0) + amount_int
    receipt['updated_at'] = time.time()

def get_billing_receipt_snapshot(job_data):
    receipt = job_data.get('billing_receipt')
    if not isinstance(receipt, dict):
        return {'charged': {}, 'refunded': {}}
    snapshot = {
        'charged': normalize_credit_ledger(receipt.get('charged', {})),
        'refunded': normalize_credit_ledger(receipt.get('refunded', {})),
    }
    updated_at = receipt.get('updated_at')
    if updated_at:
        snapshot['updated_at'] = updated_at
    return snapshot

# --- Model thinking policy ---
# Each model gets a specific thinking configuration.
MODEL_THINKING_POLICY = {
    'gemini-2.5-flash-lite': {'thinking_budget': 24576},
    'gemini-2.5-pro':        {'thinking_budget': 32768},
    'gemini-3-flash-preview': {'thinking_level': 'high'},
}

def _build_thinking_config(model_name):
    """Build a ThinkingConfig for the given model based on MODEL_THINKING_POLICY."""
    policy = MODEL_THINKING_POLICY.get(model_name)
    if not policy or not hasattr(types, 'ThinkingConfig'):
        return None
    try:
        return types.ThinkingConfig(**policy)
    except Exception:
        return None

def extract_token_usage(response):
    """Extract token counts from a Gemini response's usage_metadata."""
    meta = getattr(response, 'usage_metadata', None)
    if not meta:
        return {'input_tokens': 0, 'output_tokens': 0, 'total_tokens': 0}
    return {
        'input_tokens': getattr(meta, 'prompt_token_count', 0) or 0,
        'output_tokens': getattr(meta, 'candidates_token_count', 0) or 0,
        'total_tokens': getattr(meta, 'total_token_count', 0) or 0,
    }

class TokenAccumulator:
    """Accumulates token usage across multiple AI calls in a processing job."""
    def __init__(self):
        self.stages = {}
        self.input_total = 0
        self.output_total = 0
        self.total = 0

    def record(self, stage_name, response):
        usage = extract_token_usage(response)
        self.stages[stage_name] = usage
        self.input_total += usage['input_tokens']
        self.output_total += usage['output_tokens']
        self.total += usage['total_tokens']

    def as_dict(self):
        return {
            'token_usage_by_stage': self.stages,
            'token_input_total': self.input_total,
            'token_output_total': self.output_total,
            'token_total': self.total,
        }

def generate_with_policy(model, contents, max_output_tokens=65536):
    """Unified generation wrapper that applies model-specific thinking config."""
    base_config = {'max_output_tokens': max_output_tokens}
    thinking = _build_thinking_config(model)
    if thinking:
        base_config['thinking_config'] = thinking
    try:
        config = types.GenerateContentConfig(**base_config)
    except Exception:
        config = types.GenerateContentConfig(max_output_tokens=max_output_tokens)
    return client.models.generate_content(
        model=model,
        contents=contents,
        config=config
    )

def generate_with_optional_thinking(model, prompt_text, max_output_tokens=65536, thinking_budget=None):
    """Convenience wrapper for text-only prompts. Uses model policy for thinking config."""
    contents = [types.Content(role='user', parts=[types.Part.from_text(text=prompt_text)])]
    return generate_with_policy(model, contents, max_output_tokens=max_output_tokens)

def convert_audio_to_mp3_with_ytdlp(local_audio_path):
    return file_service.convert_audio_to_mp3_with_ytdlp(
        local_audio_path,
        ffmpeg_binary_getter=get_ffmpeg_binary,
        logger=logger,
        which_func=shutil.which,
        subprocess_module=subprocess,
    )

def resolve_auto_amount(kind, source_text):
    word_count = len((source_text or '').split())
    if kind == 'flashcards':
        if word_count < 1200:
            return 10
        if word_count < 2600:
            return 20
        return 30
    if word_count < 1200:
        return 5
    if word_count < 2600:
        return 10
    return 15

def resolve_study_amounts(flashcard_selection, question_selection, source_text):
    flashcard_amount = resolve_auto_amount('flashcards', source_text) if flashcard_selection == 'auto' else int(flashcard_selection)
    question_amount = resolve_auto_amount('questions', source_text) if question_selection == 'auto' else int(question_selection)
    return flashcard_amount, question_amount

def extract_json_payload(raw_text):
    if not raw_text:
        return None
    text = raw_text.strip()
    if text.startswith('```'):
        lines = text.splitlines()
        if len(lines) >= 3 and lines[0].startswith('```') and lines[-1].strip() == '```':
            text = '\n'.join(lines[1:-1]).strip()
    start = text.find('{')
    if start == -1:
        return None
    decoder = json.JSONDecoder()
    try:
        parsed, _ = decoder.raw_decode(text[start:])
        return parsed
    except json.JSONDecodeError:
        end = text.rfind('}')
        if end == -1 or end <= start:
            return None
        try:
            return json.loads(text[start:end + 1])
        except json.JSONDecodeError:
            return None

def sanitize_flashcards(items, max_items):
    MAX_TEXT_LEN = 2000
    if not isinstance(items, list):
        return []
    cleaned = []
    seen = set()
    for item in items:
        if not isinstance(item, dict):
            continue
        front = str(item.get('front', '')).strip()[:MAX_TEXT_LEN]
        back = str(item.get('back', '')).strip()[:MAX_TEXT_LEN]
        if not front or not back:
            continue
        key = (front.lower(), back.lower())
        if key in seen:
            continue
        seen.add(key)
        cleaned.append({'front': front, 'back': back})
        if len(cleaned) >= max_items:
            break
    return cleaned

def sanitize_questions(items, max_items):
    MAX_TEXT_LEN = 2000
    if not isinstance(items, list):
        return []
    cleaned = []
    seen = set()
    for item in items:
        if not isinstance(item, dict):
            continue
        question = str(item.get('question', '')).strip()[:MAX_TEXT_LEN]
        options = item.get('options', [])
        answer = str(item.get('answer', '')).strip()[:MAX_TEXT_LEN]
        explanation = str(item.get('explanation', '')).strip()[:MAX_TEXT_LEN]
        if not question or not isinstance(options, list) or len(options) != 4 or not answer:
            continue
        option_strings = [str(option).strip()[:MAX_TEXT_LEN] for option in options]
        if any(not option for option in option_strings):
            continue
        if len(set(option_strings)) != 4:
            continue
        if answer not in option_strings:
            continue
        dedupe_key = question.lower()
        if dedupe_key in seen:
            continue
        seen.add(dedupe_key)
        cleaned.append({
            'question': question,
            'options': option_strings,
            'answer': answer,
            'explanation': explanation,
        })
        if len(cleaned) >= max_items:
            break
    return cleaned

def default_streak_data():
    return {
        'last_study_date': '',
        'current_streak': 0,
        'daily_progress_date': '',
        'daily_progress_count': 0,
    }

def sanitize_progress_date(value):
    text = str(value or '').strip()
    return text if PROGRESS_DATE_RE.match(text) else ''

def sanitize_int(value, default=0, min_value=0, max_value=10_000_000):
    try:
        parsed = int(value)
    except (TypeError, ValueError):
        return default
    if parsed < min_value:
        return min_value
    if parsed > max_value:
        return max_value
    return parsed

def sanitize_streak_data(payload):
    base = default_streak_data()
    if not isinstance(payload, dict):
        return base
    base['last_study_date'] = sanitize_progress_date(payload.get('last_study_date', ''))
    base['current_streak'] = sanitize_int(payload.get('current_streak', 0), default=0, min_value=0, max_value=36500)
    base['daily_progress_date'] = sanitize_progress_date(payload.get('daily_progress_date', ''))
    base['daily_progress_count'] = sanitize_int(payload.get('daily_progress_count', 0), default=0, min_value=0, max_value=100000)
    return base

def sanitize_daily_goal_value(value):
    parsed = sanitize_int(value, default=-1, min_value=-1, max_value=500)
    if parsed < 1:
        return None
    return parsed

def sanitize_pack_id(value):
    pack_id = str(value or '').strip()
    if not pack_id or len(pack_id) > 160:
        return ''
    return pack_id

def sanitize_card_state_entry(payload):
    if not isinstance(payload, dict):
        return None
    seen = sanitize_int(payload.get('seen', 0), default=0, min_value=0, max_value=100000)
    correct = sanitize_int(payload.get('correct', 0), default=0, min_value=0, max_value=100000)
    wrong = sanitize_int(payload.get('wrong', 0), default=0, min_value=0, max_value=100000)
    interval_days = sanitize_int(payload.get('interval_days', 0), default=0, min_value=0, max_value=3650)
    level = str(payload.get('level', '')).strip().lower()
    if level not in {'new', 'familiar', 'mastered'}:
        if interval_days >= 14:
            level = 'mastered'
        elif seen > 0:
            level = 'familiar'
        else:
            level = 'new'
    difficulty = str(payload.get('difficulty', 'medium')).strip().lower()
    if difficulty not in {'easy', 'medium', 'hard'}:
        difficulty = 'medium'
    return {
        'seen': seen,
        'correct': correct,
        'wrong': wrong,
        'level': level,
        'interval_days': interval_days,
        'next_review_date': sanitize_progress_date(payload.get('next_review_date', '')),
        'last_review_date': sanitize_progress_date(payload.get('last_review_date', '')),
        'difficulty': difficulty,
    }

def sanitize_card_state_map(payload):
    if not isinstance(payload, dict):
        return {}
    cleaned = {}
    for raw_card_id, raw_entry in payload.items():
        card_id = str(raw_card_id or '').strip()
        if not card_id or len(card_id) > 64:
            continue
        if not re.match(r'^(fc|q)_\d{1,6}$', card_id):
            continue
        entry = sanitize_card_state_entry(raw_entry)
        if entry is None:
            continue
        cleaned[card_id] = entry
        if len(cleaned) >= MAX_PROGRESS_CARDS_PER_PACK:
            break
    return cleaned

def derive_card_level_from_stats(seen, interval_days):
    if interval_days >= 14:
        return 'mastered'
    if seen > 0:
        return 'familiar'
    return 'new'

def merge_streak_data(server_payload, incoming_payload):
    server = sanitize_streak_data(server_payload)
    incoming = sanitize_streak_data(incoming_payload)

    merged_last_study_date = max(server.get('last_study_date', ''), incoming.get('last_study_date', ''))
    if merged_last_study_date == server.get('last_study_date', '') and merged_last_study_date != incoming.get('last_study_date', ''):
        merged_current_streak = sanitize_int(server.get('current_streak', 0), default=0, min_value=0, max_value=36500)
    elif merged_last_study_date == incoming.get('last_study_date', '') and merged_last_study_date != server.get('last_study_date', ''):
        merged_current_streak = sanitize_int(incoming.get('current_streak', 0), default=0, min_value=0, max_value=36500)
    else:
        merged_current_streak = max(
            sanitize_int(server.get('current_streak', 0), default=0, min_value=0, max_value=36500),
            sanitize_int(incoming.get('current_streak', 0), default=0, min_value=0, max_value=36500),
        )

    merged_daily_progress_date = max(server.get('daily_progress_date', ''), incoming.get('daily_progress_date', ''))
    if merged_daily_progress_date == server.get('daily_progress_date', '') and merged_daily_progress_date != incoming.get('daily_progress_date', ''):
        merged_daily_progress_count = sanitize_int(server.get('daily_progress_count', 0), default=0, min_value=0, max_value=100000)
    elif merged_daily_progress_date == incoming.get('daily_progress_date', '') and merged_daily_progress_date != server.get('daily_progress_date', ''):
        merged_daily_progress_count = sanitize_int(incoming.get('daily_progress_count', 0), default=0, min_value=0, max_value=100000)
    else:
        merged_daily_progress_count = max(
            sanitize_int(server.get('daily_progress_count', 0), default=0, min_value=0, max_value=100000),
            sanitize_int(incoming.get('daily_progress_count', 0), default=0, min_value=0, max_value=100000),
        )
    if not merged_daily_progress_date:
        merged_daily_progress_count = 0

    return sanitize_streak_data({
        'last_study_date': merged_last_study_date,
        'current_streak': merged_current_streak,
        'daily_progress_date': merged_daily_progress_date,
        'daily_progress_count': merged_daily_progress_count,
    })

def merge_timezone_value(server_timezone, incoming_timezone):
    server_value = sanitize_timezone_name(server_timezone)
    incoming_value = sanitize_timezone_name(incoming_timezone)
    return incoming_value or server_value

def sanitize_timezone_name(value):
    timezone_name = str(value or '').strip()[:80]
    if not timezone_name:
        return ''
    if ZoneInfo:
        try:
            ZoneInfo(timezone_name)
            return timezone_name
        except Exception:
            return ''
    return timezone_name

def resolve_progress_timezone(progress_data):
    timezone_name = sanitize_timezone_name((progress_data or {}).get('timezone', ''))
    if timezone_name and ZoneInfo:
        try:
            return ZoneInfo(timezone_name), timezone_name
        except Exception:
            pass
    return timezone.utc, 'UTC'

def resolve_user_timezone(uid):
    safe_uid = str(uid or '').strip()
    if not safe_uid or not db:
        return timezone.utc, 'UTC'
    try:
        progress_doc = get_study_progress_doc(safe_uid).get()
        progress_data = progress_doc.to_dict() if progress_doc.exists else {}
        return resolve_progress_timezone(progress_data)
    except Exception:
        return timezone.utc, 'UTC'

def to_timezone_now(base_now, tzinfo):
    base = base_now
    if base is None:
        return datetime.now(tzinfo)
    if base.tzinfo is None:
        base = base.replace(tzinfo=timezone.utc)
    return base.astimezone(tzinfo)

def card_state_entry_rank(entry):
    if not isinstance(entry, dict):
        return ('', 0, 0, 0, 0, '')
    return (
        sanitize_progress_date(entry.get('last_review_date', '')),
        sanitize_int(entry.get('seen', 0), default=0, min_value=0, max_value=100000),
        sanitize_int(entry.get('correct', 0), default=0, min_value=0, max_value=100000),
        sanitize_int(entry.get('wrong', 0), default=0, min_value=0, max_value=100000),
        sanitize_int(entry.get('interval_days', 0), default=0, min_value=0, max_value=3650),
        sanitize_progress_date(entry.get('next_review_date', '')),
    )

def merge_card_state_entries(server_entry, incoming_entry):
    cleaned_server = sanitize_card_state_entry(server_entry)
    cleaned_incoming = sanitize_card_state_entry(incoming_entry)

    if cleaned_server is None:
        return cleaned_incoming
    if cleaned_incoming is None:
        return cleaned_server

    server_last = sanitize_progress_date(cleaned_server.get('last_review_date', ''))
    incoming_last = sanitize_progress_date(cleaned_incoming.get('last_review_date', ''))
    merged_last = max(server_last, incoming_last)

    if merged_last == server_last and merged_last != incoming_last:
        source_for_schedule = cleaned_server
    elif merged_last == incoming_last and merged_last != server_last:
        source_for_schedule = cleaned_incoming
    else:
        source_for_schedule = cleaned_server if card_state_entry_rank(cleaned_server) >= card_state_entry_rank(cleaned_incoming) else cleaned_incoming

    merged_seen = max(cleaned_server.get('seen', 0), cleaned_incoming.get('seen', 0))
    merged_correct = max(cleaned_server.get('correct', 0), cleaned_incoming.get('correct', 0))
    merged_wrong = max(cleaned_server.get('wrong', 0), cleaned_incoming.get('wrong', 0))
    minimum_seen = merged_correct + merged_wrong
    if merged_seen < minimum_seen:
        merged_seen = minimum_seen

    merged_interval_days = sanitize_int(source_for_schedule.get('interval_days', 0), default=0, min_value=0, max_value=3650)
    merged_next_review_date = sanitize_progress_date(source_for_schedule.get('next_review_date', ''))
    if not merged_next_review_date:
        merged_next_review_date = max(
            sanitize_progress_date(cleaned_server.get('next_review_date', '')),
            sanitize_progress_date(cleaned_incoming.get('next_review_date', '')),
        )

    merged_difficulty = str(source_for_schedule.get('difficulty', 'medium')).strip().lower()
    if merged_difficulty not in {'easy', 'medium', 'hard'}:
        merged_difficulty = 'medium'

    merged_entry = {
        'seen': merged_seen,
        'correct': merged_correct,
        'wrong': merged_wrong,
        'interval_days': merged_interval_days,
        'last_review_date': merged_last,
        'next_review_date': merged_next_review_date,
        'difficulty': merged_difficulty,
        'level': derive_card_level_from_stats(merged_seen, merged_interval_days),
    }
    return sanitize_card_state_entry(merged_entry)

def merge_card_state_maps(server_state, incoming_state):
    cleaned_server = sanitize_card_state_map(server_state)
    cleaned_incoming = sanitize_card_state_map(incoming_state)
    merged = {}
    for card_id in sorted(set(cleaned_server.keys()) | set(cleaned_incoming.keys())):
        merged_entry = merge_card_state_entries(cleaned_server.get(card_id), cleaned_incoming.get(card_id))
        if merged_entry is None:
            continue
        merged[card_id] = merged_entry
        if len(merged) >= MAX_PROGRESS_CARDS_PER_PACK:
            break
    return merged

def count_due_cards_in_state(state, today_local):
    due = 0
    for card_id, entry in (state or {}).items():
        if not str(card_id).startswith('fc_'):
            continue
        seen = sanitize_int((entry or {}).get('seen', 0), default=0, min_value=0, max_value=100000)
        if seen <= 0:
            continue
        next_date = str((entry or {}).get('next_review_date', '') or '').strip()
        if not next_date or next_date <= today_local:
            due += 1
    return due

def compute_study_progress_summary(progress_data, card_state_maps, base_now=None):
    progress = progress_data or {}
    streak_data = sanitize_streak_data(progress.get('streak_data', {}))
    daily_goal = sanitize_daily_goal_value(progress.get('daily_goal'))
    if daily_goal is None:
        daily_goal = 20

    tzinfo, _timezone_name = resolve_progress_timezone(progress)
    now_local = to_timezone_now(base_now, tzinfo)
    today_local = now_local.strftime('%Y-%m-%d')
    yesterday_local = (now_local - timedelta(days=1)).strftime('%Y-%m-%d')

    current_streak = 0
    if streak_data.get('last_study_date') in {today_local, yesterday_local}:
        current_streak = sanitize_int(streak_data.get('current_streak', 0), default=0, min_value=0, max_value=36500)

    today_progress = 0
    if streak_data.get('daily_progress_date') == today_local:
        today_progress = sanitize_int(streak_data.get('daily_progress_count', 0), default=0, min_value=0, max_value=100000)

    due_today = 0
    for raw_state in (card_state_maps or []):
        due_today += count_due_cards_in_state(sanitize_card_state_map(raw_state), today_local)

    return {
        'daily_goal': daily_goal,
        'current_streak': current_streak,
        'today_progress': today_progress,
        'due_today': due_today,
    }

def get_study_progress_doc(uid):
    return study_repo.study_progress_doc_ref(db, uid)

def get_study_card_state_doc(uid, pack_id):
    safe_pack_id = str(pack_id or '').replace('/', '_')
    return study_repo.study_card_state_doc_ref(db, uid, safe_pack_id)

def generate_study_materials(source_text, flashcard_selection, question_selection, study_features='both', output_language='English'):
    if study_features == 'none':
        return [], [], None
    flashcard_amount, question_amount = resolve_study_amounts(flashcard_selection, question_selection, source_text)
    if study_features == 'flashcards':
        question_amount = 0
    elif study_features == 'test':
        flashcard_amount = 0
    MAX_SOURCE_TEXT_LEN = 120000
    was_truncated = len(source_text) > MAX_SOURCE_TEXT_LEN
    prompt = PROMPT_STUDY_TEMPLATE.format(
        flashcard_amount=flashcard_amount,
        question_amount=question_amount,
        output_language=output_language,
        source_text=source_text[:MAX_SOURCE_TEXT_LEN],
    )
    try:
        response = generate_with_policy(MODEL_STUDY, [types.Content(role='user', parts=[types.Part.from_text(text=prompt)])], max_output_tokens=32768)
        parsed = extract_json_payload(response.text)
        if not isinstance(parsed, dict):
            return [], [], 'Study materials JSON parsing failed.'
        flashcards = sanitize_flashcards(parsed.get('flashcards', []), flashcard_amount)
        test_questions = sanitize_questions(parsed.get('test_questions', []), question_amount)
        if not flashcards and not test_questions and study_features != 'none':
            return [], [], 'Study materials were empty after validation.'
        error_msg = None
        if was_truncated:
            error_msg = 'Note: source text was very long and was truncated before study material generation.'
        return flashcards, test_questions, error_msg
    except Exception as e:
        return [], [], f'Study materials generation failed: {e}'

def generate_interview_enhancements(transcript_text, selected_features, output_language='English'):
    summary_text = None
    sectioned_text = None
    errors = []
    for feature in selected_features:
        try:
            if feature == 'summary':
                prompt = PROMPT_INTERVIEW_SUMMARY.format(transcript=transcript_text[:120000], output_language=output_language)
                response = generate_with_optional_thinking(MODEL_STUDY, prompt, max_output_tokens=8192, thinking_budget=384)
                summary_text = (response.text or '').strip()
                if not summary_text:
                    errors.append('Summary generation returned empty output.')
            elif feature == 'sections':
                prompt = PROMPT_INTERVIEW_SECTIONED.format(transcript=transcript_text[:120000], output_language=output_language)
                response = generate_with_optional_thinking(MODEL_STUDY, prompt, max_output_tokens=32768, thinking_budget=384)
                sectioned_text = (response.text or '').strip()
                if not sectioned_text:
                    errors.append('Sectioned transcript generation returned empty output.')
        except Exception as e:
            errors.append(f"{feature} generation failed: {e}")

    successful = []
    if summary_text:
        successful.append('summary')
    if sectioned_text:
        successful.append('sections')

    combined_text = None
    if summary_text and sectioned_text:
        combined_text = f"# Interview Summary\n\n{summary_text}\n\n# Structured Interview Transcript\n\n{sectioned_text}"

    failed_count = max(0, len(selected_features) - len(successful))
    return {
        'summary': summary_text,
        'sections': sectioned_text,
        'combined': combined_text,
        'successful_features': successful,
        'failed_count': failed_count,
        'error': '; '.join(errors) if errors else None,
    }

def allowed_file(filename, allowed_extensions):
    return file_service.allowed_file(filename, allowed_extensions)

def file_has_pdf_signature(path):
    return file_service.file_has_pdf_signature(path)

def file_has_pptx_signature(path):
    return file_service.file_has_pptx_signature(path)

def get_soffice_binary():
    return file_service.get_soffice_binary(env_getter=os.getenv, which_func=shutil.which)

def convert_pptx_to_pdf(source_path, target_pdf_path):
    return file_service.convert_pptx_to_pdf(
        source_path,
        target_pdf_path,
        soffice_binary_getter=get_soffice_binary,
        subprocess_module=subprocess,
    )

def resolve_uploaded_slides_to_pdf(uploaded_file, job_id):
    return file_service.resolve_uploaded_slides_to_pdf(
        uploaded_file,
        job_id,
        upload_folder=UPLOAD_FOLDER,
        allowed_slide_extensions=ALLOWED_SLIDE_EXTENSIONS,
        allowed_slide_mime_types=ALLOWED_SLIDE_MIME_TYPES,
        max_pdf_upload_bytes=MAX_PDF_UPLOAD_BYTES,
        cleanup_files_fn=cleanup_files,
        secure_filename_fn=secure_filename,
        allowed_file_fn=allowed_file,
        file_has_pdf_signature_fn=file_has_pdf_signature,
        file_has_pptx_signature_fn=file_has_pptx_signature,
        convert_pptx_to_pdf_fn=convert_pptx_to_pdf,
        get_saved_file_size_fn=get_saved_file_size,
    )

def file_has_audio_signature(path):
    return file_service.file_has_audio_signature(path)

def file_looks_like_audio(path):
    return file_service.file_looks_like_audio(
        path,
        ffprobe_binary_getter=get_ffprobe_binary,
        subprocess_module=subprocess,
    )

def get_saved_file_size(path):
    return file_service.get_saved_file_size(path)

def get_mime_type(filename):
    return file_service.get_mime_type(filename)

def wait_for_file_processing(uploaded_file):
    max_wait_time = 300
    wait_interval = 5
    total_waited = 0
    while total_waited < max_wait_time:
        file_info = client.files.get(name=uploaded_file.name)
        if file_info.state.name == 'ACTIVE': return True
        elif file_info.state.name == 'FAILED': raise Exception(f"File processing failed: {uploaded_file.name}")
        time.sleep(wait_interval)
        total_waited += wait_interval
    raise Exception(f"File processing timed out after {max_wait_time} seconds")

def cleanup_files(local_paths, gemini_files):
    for path in local_paths:
        try:
            if os.path.exists(path): os.remove(path)
        except Exception as e: print(f"Warning: Could not delete local file {path}: {e}")
    for gemini_file in gemini_files:
        try:
            client.files.delete(name=gemini_file.name)
        except Exception as e: print(f"Warning: Could not delete Gemini file {gemini_file.name}: {e}")

def parse_audio_markers_from_notes(notes_markdown):
    if not notes_markdown:
        return []
    pattern = re.compile(r'#{1,3}\s+(.+?)\s*\n\s*<!--\s*audio:(\d+)-(\d+)\s*-->', re.MULTILINE)
    notes_audio_map = []
    section_index = 0
    for match in pattern.finditer(notes_markdown):
        try:
            start_ms = int(match.group(2))
            end_ms = int(match.group(3))
        except Exception:
            continue
        notes_audio_map.append({
            'section_index': section_index,
            'section_title': match.group(1).strip(),
            'start_ms': max(0, start_ms),
            'end_ms': max(start_ms, end_ms),
        })
        section_index += 1
    return notes_audio_map

def format_transcript_with_timestamps(segments):
    if not isinstance(segments, list):
        return ''
    lines = []
    for seg in segments:
        if not isinstance(seg, dict):
            continue
        text = str(seg.get('text', '') or '').strip()
        if not text:
            continue
        start_ms = int(seg.get('start_ms', 0) or 0)
        end_ms = int(seg.get('end_ms', start_ms) or start_ms)
        lines.append(f"[{start_ms}-{end_ms}] {text}")
    return '\n'.join(lines)

def transcribe_audio_plain(audio_file, audio_mime_type, output_language='English'):
    output_language = OUTPUT_LANGUAGE_MAP.get(str(output_language).lower(), str(output_language))
    prompt = PROMPT_AUDIO_TRANSCRIPTION.format(output_language=output_language)
    response = generate_with_policy(MODEL_AUDIO, [types.Content(role='user', parts=[
        types.Part.from_uri(file_uri=audio_file.uri, mime_type=audio_mime_type),
        types.Part.from_text(text=prompt)
    ])])
    return (getattr(response, 'text', '') or '').strip()

def transcribe_audio_with_timestamps(audio_file, audio_mime_type, output_language='English'):
    output_language = OUTPUT_LANGUAGE_MAP.get(str(output_language).lower(), str(output_language))
    prompt = PROMPT_AUDIO_TRANSCRIPTION_TIMESTAMPED.format(output_language=output_language)
    try:
        response = generate_with_policy(MODEL_AUDIO, [types.Content(role='user', parts=[
            types.Part.from_uri(file_uri=audio_file.uri, mime_type=audio_mime_type),
            types.Part.from_text(text=prompt)
        ])])
        parsed = extract_json_payload(getattr(response, 'text', '') or '')
        if not isinstance(parsed, dict):
            raise ValueError('Timestamped transcription JSON not found')
        raw_segments = parsed.get('transcript_segments', [])
        full_transcript = str(parsed.get('full_transcript', '') or '').strip()
        clean_segments = []
        if isinstance(raw_segments, list):
            for seg in raw_segments:
                if not isinstance(seg, dict):
                    continue
                text = str(seg.get('text', '') or '').strip()
                if not text:
                    continue
                try:
                    start_ms = int(seg.get('start_ms', 0) or 0)
                    end_ms = int(seg.get('end_ms', start_ms) or start_ms)
                except Exception:
                    continue
                clean_segments.append({
                    'start_ms': max(0, start_ms),
                    'end_ms': max(start_ms, end_ms),
                    'text': text,
                })
        if not full_transcript and clean_segments:
            full_transcript = '\n'.join([s['text'] for s in clean_segments]).strip()
        if not full_transcript:
            raise ValueError('Empty transcript')
        return full_transcript, clean_segments
    except Exception as e:
        logger.warning(f"âš ï¸ Timestamp transcription failed, falling back to plain transcript: {e}")
        fallback_prompt = PROMPT_AUDIO_TRANSCRIPTION.format(output_language=output_language)
        fallback_response = generate_with_policy(MODEL_AUDIO, [types.Content(role='user', parts=[
            types.Part.from_uri(file_uri=audio_file.uri, mime_type=audio_mime_type),
            types.Part.from_text(text=fallback_prompt)
        ])])
        return (getattr(fallback_response, 'text', '') or '').strip(), []

def ensure_study_audio_root():
    os.makedirs(STUDY_AUDIO_ROOT, exist_ok=True)


def normalize_audio_storage_key(raw_key):
    key = str(raw_key or '').strip().replace('\\', '/')
    if not key:
        return ''
    key = os.path.normpath(key).replace('\\', '/')
    while key.startswith('./'):
        key = key[2:]
    if key.startswith('/'):
        return ''
    if key == '.' or key.startswith('../') or '/..' in key:
        return ''
    if not key.startswith(f'{STUDY_AUDIO_RELATIVE_DIR}/'):
        return ''
    return key


def resolve_audio_storage_path_from_key(raw_key):
    key = normalize_audio_storage_key(raw_key)
    if not key:
        return ''
    ensure_study_audio_root()
    relative_path = key[len(f'{STUDY_AUDIO_RELATIVE_DIR}/'):]
    absolute_path = os.path.abspath(os.path.join(STUDY_AUDIO_ROOT, relative_path))
    if not absolute_path.startswith(STUDY_AUDIO_ROOT + os.sep):
        return ''
    return absolute_path


def infer_audio_storage_key_from_legacy_path(raw_path):
    path = str(raw_path or '').strip()
    if not path:
        return ''
    absolute_path = os.path.abspath(path)
    ensure_study_audio_root()
    if not absolute_path.startswith(STUDY_AUDIO_ROOT + os.sep):
        return ''
    relative = os.path.relpath(absolute_path, STUDY_AUDIO_ROOT).replace('\\', '/')
    if relative == '.' or relative.startswith('../'):
        return ''
    return normalize_audio_storage_key(f'{STUDY_AUDIO_RELATIVE_DIR}/{relative}')


def get_audio_storage_key_from_pack(pack):
    if not isinstance(pack, dict):
        return ''
    key = normalize_audio_storage_key(pack.get('audio_storage_key', ''))
    if key:
        return key
    return infer_audio_storage_key_from_legacy_path(pack.get('audio_storage_path', ''))


def get_audio_storage_path_from_pack(pack):
    key = get_audio_storage_key_from_pack(pack)
    if key:
        return resolve_audio_storage_path_from_key(key)
    return ''


def ensure_pack_audio_storage_key(pack_ref, pack):
    key = get_audio_storage_key_from_pack(pack)
    if key and not normalize_audio_storage_key(pack.get('audio_storage_key', '')):
        try:
            pack_ref.set({
                'audio_storage_key': key,
                'has_audio_playback': True,
                'updated_at': time.time(),
            }, merge=True)
        except Exception:
            pass
    return key


def remove_pack_audio_file(pack):
    target_path = get_audio_storage_path_from_pack(pack)
    if not target_path:
        return False
    try:
        if os.path.exists(target_path):
            os.remove(target_path)
            return True
    except Exception:
        return False
    return False


def persist_audio_for_study_pack(job_id, audio_source_path):
    if not audio_source_path or not os.path.exists(audio_source_path):
        return ''
    ext = os.path.splitext(audio_source_path)[1].lower() or '.mp3'
    ensure_study_audio_root()
    target_key = normalize_audio_storage_key(f"{STUDY_AUDIO_RELATIVE_DIR}/{job_id}{ext}")
    target_path = resolve_audio_storage_path_from_key(target_key)
    if not target_path:
        return ''
    try:
        shutil.copy2(audio_source_path, target_path)
        return target_key
    except Exception as e:
        logger.warning(f"âš ï¸ Could not persist audio for study pack {job_id}: {e}")
        return ''

def markdown_to_docx(markdown_text, title="Document"):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    lines = markdown_text.split('\n')
    i = 0
    is_transcript = any(len(line.strip()) > 3 and line.strip()[0].isdigit() and ':' in line.strip()[:6] and ' - ' in line for line in lines[:20])

    def add_inline_markdown_runs(paragraph, text):
        raw = str(text or '')
        # Supports inline emphasis markers used in generated output.
        parts = re.split(r'(\*\*.+?\*\*|__.+?__|\*.+?\*|_.+?_)', raw)
        for part in parts:
            if not part:
                continue
            if (part.startswith('**') and part.endswith('**') and len(part) >= 4) or (part.startswith('__') and part.endswith('__') and len(part) >= 4):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
                continue
            if (part.startswith('*') and part.endswith('*') and len(part) >= 3) or (part.startswith('_') and part.endswith('_') and len(part) >= 3):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
                continue
            paragraph.add_run(part.replace('**', '').replace('__', ''))
    
    while i < len(lines):
        line = lines[i].strip()
        numbered_match = re.match(r'^\d+\.\s+(.*)$', line)
        if not line:
            i += 1
            continue
        if line.startswith('### '): doc.add_heading(line[4:], level=3)
        elif line.startswith('## '): doc.add_heading(line[3:], level=2)
        elif line.startswith('# '): doc.add_heading(line[2:], level=1)
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            add_inline_markdown_runs(p, line[2:])
        elif numbered_match:
            p = doc.add_paragraph(style='List Number')
            add_inline_markdown_runs(p, numbered_match.group(1))
        elif is_transcript and len(line) > 3 and line[0].isdigit() and ':' in line[:6]:
            p = doc.add_paragraph()
            add_inline_markdown_runs(p, line)
        else:
            paragraph_lines = [line]
            while i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                if (
                    next_line
                    and not next_line.startswith('#')
                    and not next_line.startswith('- ')
                    and not next_line.startswith('* ')
                    and not re.match(r'^\d+\.\s+', next_line)
                ):
                    paragraph_lines.append(next_line)
                    i += 1
                else:
                    break
            paragraph_text = ' '.join(paragraph_lines)
            p = doc.add_paragraph()
            add_inline_markdown_runs(p, paragraph_text)
        i += 1
    return doc

def normalize_exam_date(raw_value):
    exam_date = str(raw_value or '').strip()
    if not exam_date:
        return ''
    try:
        return datetime.strptime(exam_date, '%Y-%m-%d').strftime('%Y-%m-%d')
    except ValueError:
        raise ValueError('Exam date must use YYYY-MM-DD format')

def markdown_inline_to_pdf_html(text):
    safe_text = html.escape(str(text or ''))
    safe_text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', safe_text)
    safe_text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', safe_text)
    return safe_text

def append_notes_markdown_to_story(story, notes_markdown, styles):
    lines = str(notes_markdown or '').splitlines()
    bullet_items = []

    def flush_bullets():
        nonlocal bullet_items
        if not bullet_items:
            return
        list_flow = ListFlowable(
            [ListItem(Paragraph(item, styles['pdfBody']), leftIndent=6) for item in bullet_items],
            bulletType='bullet',
            leftIndent=14,
            bulletFontSize=8,
            bulletOffsetY=1
        )
        story.append(list_flow)
        story.append(Spacer(1, 4))
        bullet_items = []

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            flush_bullets()
            story.append(Spacer(1, 4))
            continue

        heading_level = 0
        if line.startswith('### '):
            heading_level = 3
        elif line.startswith('## '):
            heading_level = 2
        elif line.startswith('# '):
            heading_level = 1

        if heading_level:
            flush_bullets()
            heading_text = markdown_inline_to_pdf_html(line[heading_level + 1:])
            heading_style = styles['pdfH1'] if heading_level == 1 else styles['pdfH2'] if heading_level == 2 else styles['pdfH3']
            story.append(Paragraph(heading_text, heading_style))
            story.append(Spacer(1, 3))
            continue

        if line.startswith('- ') or line.startswith('* '):
            bullet_items.append(markdown_inline_to_pdf_html(line[2:].strip()))
            continue

        numbered_match = re.match(r'^(\d+)\.\s+(.*)$', line)
        if numbered_match:
            flush_bullets()
            text_html = markdown_inline_to_pdf_html(numbered_match.group(2))
            story.append(Paragraph(f"{numbered_match.group(1)}. {text_html}", styles['pdfBody']))
            story.append(Spacer(1, 2))
            continue

        flush_bullets()
        story.append(Paragraph(markdown_inline_to_pdf_html(line), styles['pdfBody']))
        story.append(Spacer(1, 2))

    flush_bullets()

def build_study_pack_pdf(pack, include_answers=True):
    if not REPORTLAB_AVAILABLE:
        raise RuntimeError(
            "PDF export requires the optional 'reportlab' dependency. "
            "Install it with: pip install reportlab==4.2.5"
        )

    pdf_buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        pdf_buffer,
        pagesize=A4,
        leftMargin=16 * mm,
        rightMargin=16 * mm,
        topMargin=14 * mm,
        bottomMargin=14 * mm,
        title=str(pack.get('title', 'Study Pack')).strip() or 'Study Pack'
    )

    base_styles = getSampleStyleSheet()
    styles = {
        'pdfTitle': ParagraphStyle(
            'PdfTitle',
            parent=base_styles['Heading1'],
            fontName='Helvetica-Bold',
            fontSize=17,
            leading=21,
            spaceAfter=6,
            textColor=colors.HexColor('#111827')
        ),
        'pdfMeta': ParagraphStyle(
            'PdfMeta',
            parent=base_styles['BodyText'],
            fontName='Helvetica',
            fontSize=9.5,
            leading=12.5,
            textColor=colors.HexColor('#4B5563')
        ),
        'pdfSection': ParagraphStyle(
            'PdfSection',
            parent=base_styles['Heading2'],
            fontName='Helvetica-Bold',
            fontSize=12.5,
            leading=16,
            spaceBefore=6,
            spaceAfter=6,
            textColor=colors.HexColor('#111827')
        ),
        'pdfH1': ParagraphStyle(
            'PdfH1',
            parent=base_styles['Heading2'],
            fontName='Helvetica-Bold',
            fontSize=12,
            leading=15,
            textColor=colors.HexColor('#1F2937')
        ),
        'pdfH2': ParagraphStyle(
            'PdfH2',
            parent=base_styles['Heading3'],
            fontName='Helvetica-Bold',
            fontSize=11,
            leading=14,
            textColor=colors.HexColor('#1F2937')
        ),
        'pdfH3': ParagraphStyle(
            'PdfH3',
            parent=base_styles['Heading4'],
            fontName='Helvetica-Bold',
            fontSize=10,
            leading=13,
            textColor=colors.HexColor('#374151')
        ),
        'pdfBody': ParagraphStyle(
            'PdfBody',
            parent=base_styles['BodyText'],
            fontName='Helvetica',
            fontSize=9.5,
            leading=13,
            textColor=colors.HexColor('#111827')
        ),
        'pdfQuestion': ParagraphStyle(
            'PdfQuestion',
            parent=base_styles['BodyText'],
            fontName='Helvetica-Bold',
            fontSize=10,
            leading=13.5,
            textColor=colors.HexColor('#111827')
        ),
        'pdfOption': ParagraphStyle(
            'PdfOption',
            parent=base_styles['BodyText'],
            fontName='Helvetica',
            fontSize=9.5,
            leading=12.5,
            leftIndent=10,
            textColor=colors.HexColor('#1F2937')
        ),
        'pdfOptionCorrect': ParagraphStyle(
            'PdfOptionCorrect',
            parent=base_styles['BodyText'],
            fontName='Helvetica-Bold',
            fontSize=9.5,
            leading=12.5,
            leftIndent=10,
            textColor=colors.HexColor('#065F46')
        ),
    }

    pack_title = str(pack.get('title', 'Study Pack')).strip() or 'Study Pack'
    story = [Paragraph(markdown_inline_to_pdf_html(pack_title), styles['pdfTitle'])]

    mode = str(pack.get('mode', '') or '').strip() or 'Unknown'
    output_language = str(pack.get('output_language', '') or '').strip() or 'Unknown'
    course = str(pack.get('course', '') or '').strip() or '-'
    subject = str(pack.get('subject', '') or '').strip() or '-'
    semester = str(pack.get('semester', '') or '').strip() or '-'
    block = str(pack.get('block', '') or '').strip() or '-'
    created_at = pack.get('created_at', 0)
    created_text = '-'
    try:
        if created_at:
            created_text = datetime.fromtimestamp(float(created_at)).strftime('%Y-%m-%d %H:%M')
    except Exception:
        created_text = '-'

    metadata_rows = [
        [Paragraph('<b>Mode</b>', styles['pdfMeta']), Paragraph(markdown_inline_to_pdf_html(mode), styles['pdfMeta'])],
        [Paragraph('<b>Language</b>', styles['pdfMeta']), Paragraph(markdown_inline_to_pdf_html(output_language), styles['pdfMeta'])],
        [Paragraph('<b>Course</b>', styles['pdfMeta']), Paragraph(markdown_inline_to_pdf_html(course), styles['pdfMeta'])],
        [Paragraph('<b>Subject</b>', styles['pdfMeta']), Paragraph(markdown_inline_to_pdf_html(subject), styles['pdfMeta'])],
        [Paragraph('<b>Semester / Block</b>', styles['pdfMeta']), Paragraph(markdown_inline_to_pdf_html(f"{semester} / {block}"), styles['pdfMeta'])],
        [Paragraph('<b>Created</b>', styles['pdfMeta']), Paragraph(markdown_inline_to_pdf_html(created_text), styles['pdfMeta'])],
    ]
    metadata_table = Table(metadata_rows, colWidths=[36 * mm, 145 * mm], hAlign='LEFT')
    metadata_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('TOPPADDING', (0, 0), (-1, -1), 2),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
        ('GRID', (0, 0), (-1, -1), 0.35, colors.HexColor('#E5E7EB')),
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#F9FAFB')),
    ]))
    story.append(metadata_table)
    story.append(Spacer(1, 10))

    story.append(Paragraph('Integrated Notes', styles['pdfSection']))
    notes_markdown = str(pack.get('notes_markdown', '') or '').strip()
    if notes_markdown:
        append_notes_markdown_to_story(story, notes_markdown, styles)
    else:
        story.append(Paragraph('No integrated notes available.', styles['pdfBody']))
    story.append(Spacer(1, 10))

    story.append(Paragraph('Flashcards', styles['pdfSection']))
    flashcards = pack.get('flashcards', []) if isinstance(pack.get('flashcards', []), list) else []
    if flashcards:
        card_rows = [[Paragraph('<b>Front</b>', styles['pdfMeta']), Paragraph('<b>Back</b>', styles['pdfMeta'])]]
        for card in flashcards:
            card_rows.append([
                Paragraph(markdown_inline_to_pdf_html(str(card.get('front', '') or '')), styles['pdfBody']),
                Paragraph(markdown_inline_to_pdf_html(str(card.get('back', '') or '')), styles['pdfBody']),
            ])
        flashcard_table = Table(card_rows, colWidths=[84 * mm, 97 * mm], repeatRows=1, hAlign='LEFT')
        flashcard_table.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 0.35, colors.HexColor('#D1D5DB')),
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#F3F4F6')),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 5),
            ('RIGHTPADDING', (0, 0), (-1, -1), 5),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ]))
        story.append(flashcard_table)
    else:
        story.append(Paragraph('No flashcards available.', styles['pdfBody']))

    story.append(PageBreak())
    practice_title = 'Practice Questions'
    if not include_answers:
        practice_title += ' (Without Answers)'
    story.append(Paragraph(practice_title, styles['pdfSection']))
    questions = pack.get('test_questions', []) if isinstance(pack.get('test_questions', []), list) else []
    if questions:
        for idx, question in enumerate(questions, 1):
            question_text = str(question.get('question', '') or '').strip() or f'Question {idx}'
            story.append(Paragraph(f"{idx}. {markdown_inline_to_pdf_html(question_text)}", styles['pdfQuestion']))

            options = question.get('options', [])
            if not isinstance(options, list):
                options = []
            answer = str(question.get('answer', '') or '').strip()
            letters = ['A', 'B', 'C', 'D']
            for option_idx, option in enumerate(options[:4]):
                option_text = str(option or '').strip()
                is_correct = include_answers and option_text == answer and option_text != ''
                marker = 'âœ“' if is_correct else 'â€¢'
                letter = letters[option_idx] if option_idx < len(letters) else str(option_idx + 1)
                option_style = styles['pdfOptionCorrect'] if is_correct else styles['pdfOption']
                story.append(Paragraph(f"{marker} {letter}. {markdown_inline_to_pdf_html(option_text)}", option_style))

            explanation = str(question.get('explanation', '') or '').strip()
            if include_answers and explanation:
                story.append(Paragraph(f"<b>Explanation:</b> {markdown_inline_to_pdf_html(explanation)}", styles['pdfBody']))
            story.append(Spacer(1, 7))
    else:
        story.append(Paragraph('No practice questions available.', styles['pdfBody']))

    doc.build(story)
    pdf_buffer.seek(0)
    return pdf_buffer

def save_study_pack(job_id, job_data):
    try:
        notes_markdown = str(job_data.get('result', '') or '')
        max_notes_chars = 180000
        notes_truncated = len(notes_markdown) > max_notes_chars
        if notes_truncated:
            notes_markdown = notes_markdown[:max_notes_chars]

        doc_ref = study_repo.create_study_pack_doc_ref(db)
        now_ts = time.time()
        tzinfo, timezone_name = resolve_user_timezone(job_data.get('user_id', ''))
        local_title_time = datetime.fromtimestamp(now_ts, tz=timezone.utc).astimezone(tzinfo)
        doc_ref.set({
            'study_pack_id': doc_ref.id,
            'source_job_id': job_id,
            'uid': job_data.get('user_id', ''),
            'mode': job_data.get('mode', ''),
            'title': f"{job_data.get('mode', 'study-pack')} {local_title_time.strftime('%Y-%m-%d %H:%M')}",
            'title_timezone': timezone_name,
            'output_language': job_data.get('output_language', 'English'),
            'notes_markdown': notes_markdown,
            'notes_truncated': notes_truncated,
            'transcript_segments': job_data.get('transcript_segments', []),
            'notes_audio_map': job_data.get('notes_audio_map', []),
            'audio_storage_key': normalize_audio_storage_key(job_data.get('audio_storage_key', '')),
            'has_audio_sync': FEATURE_AUDIO_SECTION_SYNC and bool(job_data.get('audio_storage_key')) and bool(job_data.get('notes_audio_map', [])),
            'has_audio_playback': bool(job_data.get('audio_storage_key')),
            'flashcards': job_data.get('flashcards', []),
            'test_questions': job_data.get('test_questions', []),
            'flashcard_selection': job_data.get('flashcard_selection', '20'),
            'question_selection': job_data.get('question_selection', '10'),
            'study_features': job_data.get('study_features', 'none'),
            'interview_features': job_data.get('interview_features', []),
            'interview_summary': job_data.get('interview_summary'),
            'interview_sections': job_data.get('interview_sections'),
            'interview_combined': job_data.get('interview_combined'),
            'study_generation_error': job_data.get('study_generation_error'),
            'course': '',
            'subject': '',
            'semester': '',
            'block': '',
            'folder_id': '',
            'folder_name': '',
            'created_at': now_ts,
            'updated_at': now_ts,
        })
        job_data['study_pack_id'] = doc_ref.id
    except Exception as e:
        logger.error(f"âŒ Failed to save study pack for job {job_id}: {e}")

# =============================================
# AI PROCESSING FUNCTIONS
# =============================================

def process_lecture_notes(job_id, pdf_path, audio_path):
    gemini_files = []
    local_paths = [pdf_path, audio_path]
    set_fields = lambda **fields: update_job_fields(job_id, **fields)
    get_fields = lambda: get_job_snapshot(job_id) or {}
    tokens = TokenAccumulator()
    try:
        set_fields(status='processing', step=1, step_description='Extracting text from slides...')
        pdf_file = client.files.upload(file=pdf_path, config={'mime_type': 'application/pdf'})
        gemini_files.append(pdf_file)
        wait_for_file_processing(pdf_file)
        response = generate_with_policy(MODEL_SLIDES, [types.Content(role='user', parts=[types.Part.from_uri(file_uri=pdf_file.uri, mime_type='application/pdf'), types.Part.from_text(text=PROMPT_SLIDE_EXTRACTION)])])
        tokens.record('slide_extraction', response)
        slide_text = response.text
        set_fields(slide_text=slide_text, step=2, step_description='Transcribing audio...')
        output_language = get_fields().get('output_language', 'English')
        converted_audio_path, converted = convert_audio_to_mp3_with_ytdlp(audio_path)
        if converted and converted_audio_path not in local_paths:
            local_paths.append(converted_audio_path)
        set_fields(step_description='Optimizing audio for faster processing...')
        audio_mime_type = get_mime_type(converted_audio_path)
        audio_file = client.files.upload(file=converted_audio_path, config={'mime_type': audio_mime_type})
        gemini_files.append(audio_file)
        set_fields(step_description='Processing audio file (this may take a few minutes)...')
        wait_for_file_processing(audio_file)
        set_fields(step_description='Generating transcript...')
        if FEATURE_AUDIO_SECTION_SYNC:
            transcript, transcript_segments = transcribe_audio_with_timestamps(audio_file, audio_mime_type, output_language)
        else:
            transcript = transcribe_audio_plain(audio_file, audio_mime_type, output_language)
            transcript_segments = []
        set_fields(
            transcript=transcript,
            transcript_segments=transcript_segments,
            audio_storage_key=persist_audio_for_study_pack(job_id, converted_audio_path),
            step=3,
            step_description='Creating complete lecture notes...',
        )
        merge_transcript = format_transcript_with_timestamps(transcript_segments) if transcript_segments else transcript
        if FEATURE_AUDIO_SECTION_SYNC and transcript_segments:
            merge_prompt = PROMPT_MERGE_WITH_AUDIO_MARKERS.format(slide_text=slide_text, transcript=merge_transcript, output_language=output_language)
        else:
            merge_prompt = PROMPT_MERGE_TEMPLATE.format(slide_text=slide_text, transcript=transcript, output_language=output_language)
        response = generate_with_policy(MODEL_INTEGRATION, [types.Content(role='user', parts=[types.Part.from_text(text=merge_prompt)])])
        tokens.record('merge', response)
        merged_notes = response.text
        set_fields(
            result=merged_notes,
            notes_audio_map=parse_audio_markers_from_notes(merged_notes) if FEATURE_AUDIO_SECTION_SYNC else [],
        )

        job_data = get_fields()
        if job_data.get('study_features', 'none') != 'none':
            set_fields(step=4, step_description='Generating flashcards and practice test...')
            flashcards, test_questions, study_error = generate_study_materials(
                merged_notes,
                job_data.get('flashcard_selection', '20'),
                job_data.get('question_selection', '10'),
                job_data.get('study_features', 'none'),
                output_language
            )
            set_fields(
                flashcards=flashcards,
                test_questions=test_questions,
                study_generation_error=study_error,
            )
        else:
            set_fields(flashcards=[], test_questions=[], study_generation_error=None)
        job_data = get_fields()
        save_study_pack(job_id, job_data)
        final_snapshot = get_fields()
        set_fields(
            status='complete',
            step=final_snapshot.get('total_steps', 3),
            step_description='Complete!',
        )
    except Exception as e:
        logger.exception("Lecture-notes processing failed for job %s", job_id)
        set_fields(status='error', error=PROCESSING_PUBLIC_ERROR_MESSAGE)
        # Refund the credit since processing failed
        failed_job = get_fields()
        uid = failed_job.get('user_id')
        credit_type = failed_job.get('credit_deducted')
        refund_credit(uid, credit_type)
        failed_job = get_fields()
        add_job_credit_refund(failed_job, credit_type, 1)
        set_job(job_id, failed_job)
        set_fields(credit_refunded=True)
    finally:
        cleanup_files(local_paths, gemini_files)
        # Save token telemetry and log the job
        finished_at = time.time()
        set_fields(finished_at=finished_at, **tokens.as_dict())
        final_job = get_fields()
        save_job_log(job_id, final_job, finished_at)

def process_slides_only(job_id, pdf_path):
    gemini_files = []
    local_paths = [pdf_path]
    set_fields = lambda **fields: update_job_fields(job_id, **fields)
    get_fields = lambda: get_job_snapshot(job_id) or {}
    tokens = TokenAccumulator()
    try:
        set_fields(status='processing', step=1, step_description='Extracting text from slides...')
        pdf_file = client.files.upload(file=pdf_path, config={'mime_type': 'application/pdf'})
        gemini_files.append(pdf_file)
        wait_for_file_processing(pdf_file)
        response = generate_with_policy(MODEL_SLIDES, [types.Content(role='user', parts=[types.Part.from_uri(file_uri=pdf_file.uri, mime_type='application/pdf'), types.Part.from_text(text=PROMPT_SLIDE_EXTRACTION)])])
        tokens.record('slide_extraction', response)
        extracted_text = response.text
        set_fields(result=extracted_text)
        job_data = get_fields()
        if job_data.get('study_features', 'none') != 'none':
            set_fields(step=2, step_description='Generating flashcards and practice test...')
            flashcards, test_questions, study_error = generate_study_materials(
                extracted_text,
                job_data.get('flashcard_selection', '20'),
                job_data.get('question_selection', '10'),
                job_data.get('study_features', 'none'),
                job_data.get('output_language', 'English')
            )
            set_fields(
                flashcards=flashcards,
                test_questions=test_questions,
                study_generation_error=study_error,
            )
        else:
            set_fields(flashcards=[], test_questions=[], study_generation_error=None)
        job_data = get_fields()
        save_study_pack(job_id, job_data)
        final_snapshot = get_fields()
        set_fields(
            status='complete',
            step=final_snapshot.get('total_steps', 1),
            step_description='Complete!',
        )
    except Exception as e:
        logger.exception("Slides-only processing failed for job %s", job_id)
        set_fields(status='error', error=PROCESSING_PUBLIC_ERROR_MESSAGE)
        # Refund the credit since processing failed
        failed_job = get_fields()
        uid = failed_job.get('user_id')
        credit_type = failed_job.get('credit_deducted')
        refund_credit(uid, credit_type)
        failed_job = get_fields()
        add_job_credit_refund(failed_job, credit_type, 1)
        set_job(job_id, failed_job)
        set_fields(credit_refunded=True)
    finally:
        cleanup_files(local_paths, gemini_files)
        # Save token telemetry and log the job
        finished_at = time.time()
        set_fields(finished_at=finished_at, **tokens.as_dict())
        final_job = get_fields()
        save_job_log(job_id, final_job, finished_at)

def process_interview_transcription(job_id, audio_path):
    gemini_files = []
    local_paths = [audio_path]
    set_fields = lambda **fields: update_job_fields(job_id, **fields)
    get_fields = lambda: get_job_snapshot(job_id) or {}
    tokens = TokenAccumulator()
    try:
        set_fields(status='processing', step=1, step_description='Optimizing audio for faster processing...')
        output_language = get_fields().get('output_language', 'English')
        converted_audio_path, converted = convert_audio_to_mp3_with_ytdlp(audio_path)
        if converted and converted_audio_path not in local_paths:
            local_paths.append(converted_audio_path)
        set_fields(audio_storage_key=persist_audio_for_study_pack(job_id, converted_audio_path))
        audio_mime_type = get_mime_type(converted_audio_path)
        audio_file = client.files.upload(file=converted_audio_path, config={'mime_type': audio_mime_type})
        gemini_files.append(audio_file)
        set_fields(step_description='Processing audio file (this may take a few minutes)...')
        wait_for_file_processing(audio_file)
        set_fields(step_description='Generating transcript with timestamps...')
        interview_prompt = PROMPT_INTERVIEW_TRANSCRIPTION.format(output_language=output_language)
        response = generate_with_policy(MODEL_INTERVIEW, [types.Content(role='user', parts=[types.Part.from_uri(file_uri=audio_file.uri, mime_type=audio_mime_type), types.Part.from_text(text=interview_prompt)])])
        tokens.record('interview_transcription', response)
        transcript_text = response.text or ''
        set_fields(transcript=transcript_text, result=transcript_text)

        job_data = get_fields()
        selected_features = job_data.get('interview_features', [])
        if selected_features:
            set_fields(step=2, step_description='Creating interview summary and sections...')
            enhancement = generate_interview_enhancements(transcript_text, selected_features, output_language)
            set_fields(
                interview_summary=enhancement.get('summary'),
                interview_sections=enhancement.get('sections'),
                interview_combined=enhancement.get('combined'),
                interview_features_successful=enhancement.get('successful_features', []),
                study_generation_error=enhancement.get('error'),
            )

            failed_count = enhancement.get('failed_count', 0)
            if failed_count > 0:
                current_job = get_fields()
                uid = current_job.get('user_id')
                refund_slides_credits(uid, failed_count)
                current_job = get_fields()
                current_job['extra_slides_refunded'] = current_job.get('extra_slides_refunded', 0) + failed_count
                add_job_credit_refund(current_job, 'slides_credits', failed_count)
                set_job(job_id, current_job)

            if enhancement.get('summary') and enhancement.get('sections'):
                set_fields(result=enhancement.get('combined', transcript_text))
            elif enhancement.get('summary'):
                set_fields(result=enhancement.get('summary'))
            elif enhancement.get('sections'):
                set_fields(result=enhancement.get('sections'))

        job_data = get_fields()
        save_study_pack(job_id, job_data)
        final_snapshot = get_fields()
        set_fields(
            status='complete',
            step=final_snapshot.get('total_steps', 1),
            step_description='Complete!',
        )
    except Exception as e:
        logger.exception("Interview processing failed for job %s", job_id)
        set_fields(status='error', error=PROCESSING_PUBLIC_ERROR_MESSAGE)
        # Refund the credit since processing failed
        failed_job = get_fields()
        uid = failed_job.get('user_id')
        credit_type = failed_job.get('credit_deducted')
        refund_credit(uid, credit_type)
        failed_job = get_fields()
        add_job_credit_refund(failed_job, credit_type, 1)
        extra_spent = failed_job.get('interview_features_cost', 0)
        already_refunded = failed_job.get('extra_slides_refunded', 0)
        to_refund = max(0, extra_spent - already_refunded)
        if to_refund > 0:
            refund_slides_credits(uid, to_refund)
            failed_job['extra_slides_refunded'] = already_refunded + to_refund
            add_job_credit_refund(failed_job, 'slides_credits', to_refund)
        failed_job['credit_refunded'] = True
        set_job(job_id, failed_job)
    finally:
        cleanup_files(local_paths, gemini_files)
        # Save token telemetry and log the job
        finished_at = time.time()
        set_fields(finished_at=finished_at, **tokens.as_dict())
        final_job = get_fields()
        save_job_log(job_id, final_job, finished_at)

# =============================================
# ROUTES
# =============================================

@app.route('/')
def index():
    return render_template('landing.html')

@app.route('/dashboard')
def dashboard():
    return render_template(
        'index.html',
        sentry_frontend_dsn=SENTRY_FRONTEND_DSN,
        sentry_environment=SENTRY_ENVIRONMENT,
        sentry_release=SENTRY_RELEASE,
        index_js_asset=resolve_js_asset('js/index-app.js'),
    )

@app.route('/plan')
@app.route('/stats')
def plan_dashboard():
    return render_template('plan.html')

@app.route('/calendar')
def calendar_dashboard():
    return render_template('calendar.html')

@app.route('/features')
def features_page():
    return render_template('features.html')

@app.route('/admin')
def admin_dashboard():
    decoded_token = verify_admin_session_cookie(request)
    if not decoded_token:
        if ADMIN_PAGE_UNAUTHORIZED_MODE == '404':
            abort(404)
        return redirect('/dashboard')
    return render_template('admin.html', admin_js_asset=resolve_js_asset('js/admin.js'))

@app.route('/study')
def study_dashboard():
    return render_template('study.html', study_js_asset=resolve_js_asset('js/study.js'))

@app.route('/privacy')
def privacy_policy():
    return render_template(
        'privacy.html',
        legal_contact_email=LEGAL_CONTACT_EMAIL,
        last_updated='February 26, 2026',
    )

@app.route('/terms')
def terms_of_service():
    return render_template(
        'terms.html',
        legal_contact_email=LEGAL_CONTACT_EMAIL,
        last_updated='February 26, 2026',
    )

def create_admin_session_impl():
    from lecture_processor.services import auth_api_service

    return auth_api_service.create_admin_session(sys.modules[__name__], request)


def clear_admin_session_impl():
    from lecture_processor.services import auth_api_service

    return auth_api_service.clear_admin_session(sys.modules[__name__], request)

def verify_email_impl():
    from lecture_processor.services import auth_api_service

    return auth_api_service.verify_email(sys.modules[__name__], request)

def dev_sentry_test_impl():
    from lecture_processor.services import auth_api_service

    return auth_api_service.dev_sentry_test(sys.modules[__name__], request)

def ingest_analytics_event_impl():
    from lecture_processor.services import auth_api_service

    return auth_api_service.ingest_analytics_event(sys.modules[__name__], request)

def get_user_impl():
    from lecture_processor.services import auth_api_service

    return auth_api_service.get_user(sys.modules[__name__], request)

def get_user_preferences_impl():
    from lecture_processor.services import auth_api_service

    return auth_api_service.get_user_preferences(sys.modules[__name__], request)

def update_user_preferences_impl():
    from lecture_processor.services import auth_api_service

    return auth_api_service.update_user_preferences(sys.modules[__name__], request)

def export_account_data_impl():
    from lecture_processor.services import auth_api_service

    return auth_api_service.export_account_data(sys.modules[__name__], request)

def delete_account_data_impl():
    from lecture_processor.services import auth_api_service

    return auth_api_service.delete_account_data(sys.modules[__name__], request)

def get_study_progress_impl():
    from lecture_processor.services import study_api_service

    return study_api_service.get_study_progress(sys.modules[__name__], request)

def update_study_progress_impl():
    from lecture_processor.services import study_api_service

    return study_api_service.update_study_progress(sys.modules[__name__], request)

def get_study_progress_summary_impl():
    from lecture_processor.services import study_api_service

    return study_api_service.get_study_progress_summary(sys.modules[__name__], request)

# --- Stripe Routes ---

def get_config_impl():
    from lecture_processor.services import payments_api_service

    return payments_api_service.get_config(sys.modules[__name__])

def create_checkout_session_impl():
    from lecture_processor.services import payments_api_service

    return payments_api_service.create_checkout_session(sys.modules[__name__], request)

def confirm_checkout_session_impl():
    from lecture_processor.services import payments_api_service

    return payments_api_service.confirm_checkout_session(sys.modules[__name__], request)

def stripe_webhook_impl():
    from lecture_processor.services import payments_api_service

    return payments_api_service.stripe_webhook(sys.modules[__name__], request)

# --- Purchase History Route ---

def get_purchase_history_impl():
    from lecture_processor.services import payments_api_service

    return payments_api_service.get_purchase_history(sys.modules[__name__], request)

def get_study_packs_impl():
    from lecture_processor.services import study_api_service

    return study_api_service.get_study_packs(sys.modules[__name__], request)

def create_study_pack_impl():
    from lecture_processor.services import study_api_service

    return study_api_service.create_study_pack(sys.modules[__name__], request)

def get_study_pack_impl(pack_id):
    from lecture_processor.services import study_api_service

    return study_api_service.get_study_pack(sys.modules[__name__], request, pack_id)

def update_study_pack_impl(pack_id):
    from lecture_processor.services import study_api_service

    return study_api_service.update_study_pack(sys.modules[__name__], request, pack_id)

def delete_study_pack_impl(pack_id):
    from lecture_processor.services import study_api_service

    return study_api_service.delete_study_pack(sys.modules[__name__], request, pack_id)

def get_study_folders_impl():
    from lecture_processor.services import study_api_service

    return study_api_service.get_study_folders(sys.modules[__name__], request)

def get_study_pack_audio_url_impl(pack_id):
    from lecture_processor.services import study_api_service

    return study_api_service.get_study_pack_audio_url(sys.modules[__name__], request, pack_id)

def stream_study_pack_audio_impl(pack_id):
    from lecture_processor.services import study_api_service

    return study_api_service.stream_study_pack_audio(sys.modules[__name__], request, pack_id)

def stream_audio_token_impl(token):
    from lecture_processor.services import study_api_service

    return study_api_service.stream_audio_token(sys.modules[__name__], token)

def create_study_folder_impl():
    from lecture_processor.services import study_api_service

    return study_api_service.create_study_folder(sys.modules[__name__], request)

def update_study_folder_impl(folder_id):
    from lecture_processor.services import study_api_service

    return study_api_service.update_study_folder(sys.modules[__name__], request, folder_id)

def delete_study_folder_impl(folder_id):
    from lecture_processor.services import study_api_service

    return study_api_service.delete_study_folder(sys.modules[__name__], request, folder_id)

def admin_overview_impl():
    from lecture_processor.services import admin_api_service

    return admin_api_service.admin_overview(sys.modules[__name__], request)

def admin_export_impl():
    from lecture_processor.services import admin_api_service

    return admin_api_service.admin_export(sys.modules[__name__], request)

def admin_prompts_impl():
    decoded_token = verify_firebase_token(request)
    if not decoded_token:
        return jsonify({'error': 'Unauthorized'}), 401
    if not is_admin_user(decoded_token):
        return jsonify({'error': 'Forbidden'}), 403
    fmt = request.args.get('format', 'json')
    if fmt == 'markdown':
        return jsonify({'markdown': get_prompt_inventory_markdown()})
    return jsonify({'prompts': get_prompt_inventory()})


def processing_averages_impl():
    """Return average processing duration per mode from completed job_logs.
    No auth required â€“ only aggregate stats are returned."""
    try:
        from collections import defaultdict
        docs = db.collection('job_logs') \
            .where('status', '==', 'complete') \
            .order_by('created_at', direction=firestore.Query.DESCENDING) \
            .limit(200) \
            .stream()
        by_mode = defaultdict(list)
        for doc in docs:
            job = doc.to_dict() or {}
            mode = job.get('mode', 'unknown')
            duration = job.get('duration_seconds')
            if isinstance(duration, (int, float)) and duration > 0:
                by_mode[mode].append(duration)
        averages = {}
        total_jobs = 0
        for mode, durations in by_mode.items():
            total_jobs += len(durations)
            avg = round(sum(durations) / len(durations), 1)
            averages[mode] = {
                'avg_seconds': avg,
                'job_count': len(durations),
                'min_seconds': round(min(durations), 1),
                'max_seconds': round(max(durations), 1),
            }
        resp = jsonify({'averages': averages, 'total_jobs': total_jobs})
        resp.headers['Cache-Control'] = 'public, max-age=300'
        return resp
    except Exception as e:
        return jsonify({'averages': {}, 'total_jobs': 0, 'error': str(e)})

# --- Upload & Processing Routes ---

def import_audio_from_url_impl():
    from lecture_processor.services import upload_api_service

    return upload_api_service.import_audio_from_url(sys.modules[__name__], request)

def release_imported_audio_impl():
    from lecture_processor.services import upload_api_service

    return upload_api_service.release_imported_audio(sys.modules[__name__], request)

def upload_files_impl():
    from lecture_processor.services import upload_api_service

    return upload_api_service.upload_files(sys.modules[__name__], request)

def get_status_impl(job_id):
    from lecture_processor.services import upload_api_service

    return upload_api_service.get_status(sys.modules[__name__], request, job_id)

def download_docx_impl(job_id):
    from lecture_processor.services import upload_api_service

    return upload_api_service.download_docx(sys.modules[__name__], request, job_id)

def download_flashcards_csv_impl(job_id):
    from lecture_processor.services import upload_api_service

    return upload_api_service.download_flashcards_csv(sys.modules[__name__], request, job_id)

def export_study_pack_flashcards_csv_impl(pack_id):
    from lecture_processor.services import study_api_service

    return study_api_service.export_study_pack_flashcards_csv(sys.modules[__name__], request, pack_id)

def export_study_pack_notes_impl(pack_id):
    from lecture_processor.services import study_api_service

    return study_api_service.export_study_pack_notes(sys.modules[__name__], request, pack_id)

def export_study_pack_pdf_impl(pack_id):
    from lecture_processor.services import study_api_service

    return study_api_service.export_study_pack_pdf(sys.modules[__name__], request, pack_id)

# Register extracted non-admin API blueprints (Batch R3).
for _blueprint in (auth_bp, account_bp, study_bp, upload_bp, admin_bp, payments_bp):
    if _blueprint.name not in app.blueprints:
        app.register_blueprint(_blueprint)

recover_stale_runtime_jobs()

# =============================================
# HEALTH CHECK (Issue 38)
# =============================================
@app.route('/healthz')
def healthz():
    return jsonify({'status': 'ok'}), 200

if __name__ == '__main__':
    app.run(debug=True, port=5000)
