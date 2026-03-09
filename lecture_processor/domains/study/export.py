import html
import io
import re
from datetime import datetime
from html.parser import HTMLParser

from docx import Document
from docx.shared import Pt

from lecture_processor.domains.shared import sanitize_csv_row
from lecture_processor.runtime.container import get_runtime

REPORTLAB_AVAILABLE = True

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        ListFlowable,
        ListItem,
        PageBreak,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )
except Exception:
    REPORTLAB_AVAILABLE = False


ANNOTATED_NOTES_HIGHLIGHT_COLORS = {
    'yellow': '#FEF08A',
    'green': '#BBF7D0',
    'blue': '#BFDBFE',
    'pink': '#FBCFE8',
}
ANNOTATED_NOTES_CONTAINER_TAGS = {'div', 'section', 'article', 'main', 'aside'}
ANNOTATED_NOTES_INLINE_TAGS = {'span', 'strong', 'b', 'em', 'i', 'u', 'mark', 'a', 'code', 'small', 'sup', 'sub'}
ANNOTATED_NOTES_BLOCK_TAGS = {'p', 'ul', 'ol', 'li', 'blockquote', 'pre', 'hr', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6'}


class _AnnotatedHtmlNode:
    __slots__ = ('tag', 'attrs', 'children')

    def __init__(self, tag, attrs=None):
        self.tag = str(tag or '').lower()
        self.attrs = dict(attrs or {})
        self.children = []


class _AnnotatedHtmlTreeParser(HTMLParser):
    VOID_TAGS = {'br', 'hr'}

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.root = _AnnotatedHtmlNode('root')
        self.stack = [self.root]

    def handle_starttag(self, tag, attrs):
        node = _AnnotatedHtmlNode(tag, attrs)
        self.stack[-1].children.append(node)
        if node.tag not in self.VOID_TAGS:
            self.stack.append(node)

    def handle_startendtag(self, tag, attrs):
        node = _AnnotatedHtmlNode(tag, attrs)
        self.stack[-1].children.append(node)

    def handle_endtag(self, tag):
        target = str(tag or '').lower()
        for index in range(len(self.stack) - 1, 0, -1):
            if self.stack[index].tag == target:
                del self.stack[index:]
                break

    def handle_data(self, data):
        if data:
            self.stack[-1].children.append(str(data))


def _resolve_runtime(runtime=None):
    if runtime is not None:
        return runtime
    return get_runtime()


def markdown_to_docx(markdown_text, title='Document', runtime=None):
    _ = runtime
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    lines = str(markdown_text or '').split('\n')
    i = 0
    is_transcript = any(
        (
            len(line.strip()) > 3
            and line.strip()[0].isdigit()
            and (':' in line.strip()[:6])
            and (' - ' in line)
        )
        for line in lines[:20]
    )

    def add_inline_markdown_runs(paragraph, text):
        raw = str(text or '')
        parts = re.split('(\\*\\*.+?\\*\\*|__.+?__|\\*.+?\\*|_.+?_)', raw)
        for part in parts:
            if not part:
                continue
            if (part.startswith('**') and part.endswith('**') and (len(part) >= 4)) or (
                part.startswith('__') and part.endswith('__') and (len(part) >= 4)
            ):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
                continue
            if (part.startswith('*') and part.endswith('*') and (len(part) >= 3)) or (
                part.startswith('_') and part.endswith('_') and (len(part) >= 3)
            ):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
                continue
            paragraph.add_run(part.replace('**', '').replace('__', ''))

    def pick_list_style(kind, level):
        safe_level = max(1, min(int(level or 1), 3))
        if kind == 'number':
            preferred = ['List Number', 'List Number 2', 'List Number 3'][safe_level - 1]
            fallback = 'List Number'
        else:
            preferred = ['List Bullet', 'List Bullet 2', 'List Bullet 3'][safe_level - 1]
            fallback = 'List Bullet'
        for candidate in (preferred, fallback):
            try:
                _ = doc.styles[candidate]
                return candidate
            except KeyError:
                continue
        return ''

    def parse_list_line(raw_line):
        line_value = str(raw_line or '').replace('\t', '    ')
        if not line_value.strip():
            return None

        bullet_match = re.match('^(\\s*)[-*•]\\s+(.*)$', line_value)
        if bullet_match:
            indent_spaces = len(bullet_match.group(1))
            content = bullet_match.group(2).strip()
            if not content:
                return None
            extra_depth = 0
            while True:
                nested_bullet = re.match('^[-*•]\\s+(.*)$', content)
                if nested_bullet:
                    content = nested_bullet.group(1).strip()
                    extra_depth += 1
                    if not content:
                        return None
                    continue
                nested_number = re.match('^(\\d+[\\.\\)])\\s+(.*)$', content)
                if nested_number:
                    content = nested_number.group(2).strip()
                    if not content:
                        return None
                    return ('number', indent_spaces // 2 + 1 + extra_depth, content)
                break
            nested_number = re.match('^(\\d+[\\.\\)])\\s+(.*)$', content)
            kind = 'number' if nested_number else 'bullet'
            item_text = nested_number.group(2).strip() if nested_number else content
            return (kind, indent_spaces // 2 + 1 + extra_depth, item_text)

        number_match = re.match('^(\\s*)(\\d+[\\.\\)])\\s+(.*)$', line_value)
        if number_match:
            indent_spaces = len(number_match.group(1))
            content = number_match.group(3).strip()
            if not content:
                return None
            extra_depth = 0
            while True:
                nested_number = re.match('^(\\d+[\\.\\)])\\s+(.*)$', content)
                if nested_number:
                    content = nested_number.group(2).strip()
                    extra_depth += 1
                    if not content:
                        return None
                    continue
                nested_bullet = re.match('^[-*•]\\s+(.*)$', content)
                if nested_bullet:
                    content = nested_bullet.group(1).strip()
                    if not content:
                        return None
                    return ('bullet', indent_spaces // 2 + 1 + extra_depth, content)
                break
            return ('number', indent_spaces // 2 + 1 + extra_depth, content)

        return None

    while i < len(lines):
        raw_line = lines[i]
        line = raw_line.strip()
        list_info = parse_list_line(raw_line)

        if not line:
            i += 1
            continue

        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif list_info:
            list_kind, list_level, list_text = list_info
            list_style = pick_list_style(list_kind, list_level)
            paragraph = doc.add_paragraph(style=list_style) if list_style else doc.add_paragraph()
            add_inline_markdown_runs(paragraph, list_text)
        elif is_transcript and len(line) > 3 and line[0].isdigit() and (':' in line[:6]):
            paragraph = doc.add_paragraph()
            add_inline_markdown_runs(paragraph, line)
        else:
            paragraph_lines = [line]
            while i + 1 < len(lines):
                next_raw = lines[i + 1]
                next_line = next_raw.strip()
                if next_line and (not next_line.startswith('#')) and (not parse_list_line(next_raw)):
                    paragraph_lines.append(next_line)
                    i += 1
                else:
                    break
            paragraph_text = ' '.join(paragraph_lines)
            paragraph = doc.add_paragraph()
            add_inline_markdown_runs(paragraph, paragraph_text)
        i += 1

    return doc


def sanitize_export_filename(value, fallback='study-pack'):
    raw = str(value or '').strip().lower()
    if not raw:
        return str(fallback or 'study-pack')
    safe = re.sub(r'[^a-z0-9._-]+', '-', raw)
    safe = re.sub(r'-{2,}', '-', safe).strip('._-')
    return safe[:80] or str(fallback or 'study-pack')


def build_flashcards_csv_bytes(pack, runtime=None):
    resolved_runtime = _resolve_runtime(runtime)
    flashcards = pack.get('flashcards', []) if isinstance(pack.get('flashcards', []), list) else []
    if not flashcards:
        return None
    output = resolved_runtime.io.StringIO()
    writer = resolved_runtime.csv.writer(output)
    writer.writerow(['question', 'answer'])
    for card in flashcards:
        writer.writerow(sanitize_csv_row([card.get('front', ''), card.get('back', '')]))
    return output.getvalue().encode('utf-8')


def build_practice_test_csv_bytes(pack, runtime=None):
    resolved_runtime = _resolve_runtime(runtime)
    questions = pack.get('test_questions', []) if isinstance(pack.get('test_questions', []), list) else []
    if not questions:
        return None
    output = resolved_runtime.io.StringIO()
    writer = resolved_runtime.csv.writer(output)
    writer.writerow(['question', 'option_a', 'option_b', 'option_c', 'option_d', 'answer', 'explanation'])
    for question in questions:
        options = question.get('options', []) if isinstance(question.get('options', []), list) else []
        padded = (options + ['', '', '', ''])[:4]
        writer.writerow(sanitize_csv_row([
            question.get('question', ''),
            padded[0],
            padded[1],
            padded[2],
            padded[3],
            question.get('answer', ''),
            question.get('explanation', ''),
        ]))
    return output.getvalue().encode('utf-8')


def build_notes_docx_bytes(pack, runtime=None):
    notes_markdown = str(pack.get('notes_markdown', '') or '').strip()
    if not notes_markdown:
        return None
    title = str(pack.get('title', 'Lecture Notes') or 'Lecture Notes').strip()
    docx = markdown_to_docx(notes_markdown, title=title, runtime=runtime)
    buffer = io.BytesIO()
    docx.save(buffer)
    buffer.seek(0)
    return buffer.read()


def build_notes_pdf_bytes(pack, include_answers=True, runtime=None):
    notes_markdown = str(pack.get('notes_markdown', '') or '').strip()
    if not notes_markdown:
        return None
    pdf_buffer = build_study_pack_pdf(pack, include_answers=include_answers, runtime=runtime)
    pdf_buffer.seek(0)
    return pdf_buffer.read()


def _parse_annotated_notes_html(raw_html):
    parser = _AnnotatedHtmlTreeParser()
    parser.feed(str(raw_html or ''))
    parser.close()
    return parser.root


def _annotated_inline_has_text(html_value):
    text_only = re.sub(r'<[^>]+>', '', str(html_value or ''))
    return bool(text_only.replace('&nbsp;', ' ').strip())


def _collapse_annotated_whitespace(text):
    return re.sub(r'\s+', ' ', str(text or '').replace('\xa0', ' '))


def _render_annotated_inline_html(nodes):
    parts = []
    for node in nodes or []:
        if isinstance(node, str):
            text_value = _collapse_annotated_whitespace(node)
            if text_value:
                parts.append(html.escape(text_value))
            continue

        tag = node.tag
        if tag == 'br':
            parts.append('<br/>')
            continue

        inner_html = _render_annotated_inline_html(node.children)
        if tag in {'strong', 'b'} and inner_html:
            parts.append(f'<b>{inner_html}</b>')
            continue
        if tag in {'em', 'i'} and inner_html:
            parts.append(f'<i>{inner_html}</i>')
            continue
        if tag == 'u' and inner_html:
            parts.append(f'<u>{inner_html}</u>')
            continue
        if tag == 'mark' and inner_html:
            highlight_color = ANNOTATED_NOTES_HIGHLIGHT_COLORS.get(
                str(node.attrs.get('data-hl', '') or '').strip().lower(),
                ANNOTATED_NOTES_HIGHLIGHT_COLORS['yellow'],
            )
            parts.append(f'<font backColor="{highlight_color}">{inner_html}</font>')
            continue
        if tag == 'code' and inner_html:
            parts.append(f'<font name="Courier">{inner_html}</font>')
            continue
        if tag == 'a' and inner_html:
            parts.append(f'<font color="#4338CA"><u>{inner_html}</u></font>')
            continue
        parts.append(inner_html)

    html_value = ''.join(parts)
    html_value = re.sub(r'(?:\s*<br/>\s*){3,}', '<br/><br/>', html_value)
    return html_value.strip()


def _annotated_node_has_block_children(node):
    return any(
        isinstance(child, _AnnotatedHtmlNode)
        and (child.tag in ANNOTATED_NOTES_CONTAINER_TAGS or child.tag in ANNOTATED_NOTES_BLOCK_TAGS)
        for child in (node.children or [])
    )


def _annotated_heading_style_name(tag_name):
    if tag_name in {'h1', 'h2'}:
        return 'annotatedH1'
    if tag_name in {'h3', 'h4'}:
        return 'annotatedH2'
    return 'annotatedH3'


def _annotated_list_style_name(level):
    safe_level = max(0, min(int(level or 0), 2))
    return f'annotatedList{safe_level + 1}'


def _annotated_list_continue_style_name(level):
    safe_level = max(0, min(int(level or 0), 2))
    return f'annotatedListContinue{safe_level + 1}'


def _iter_annotated_nested_lists(node):
    for child in node.children or []:
        if isinstance(child, _AnnotatedHtmlNode) and child.tag in {'ul', 'ol'}:
            yield child


def _render_annotated_list_item_html(node):
    content_nodes = []
    for child in node.children or []:
        if isinstance(child, _AnnotatedHtmlNode) and child.tag in {'ul', 'ol'}:
            continue
        if isinstance(child, _AnnotatedHtmlNode) and child.tag in ANNOTATED_NOTES_BLOCK_TAGS.union(ANNOTATED_NOTES_CONTAINER_TAGS):
            if content_nodes:
                content_nodes.append(_AnnotatedHtmlNode('br'))
            content_nodes.extend(child.children or [])
            continue
        content_nodes.append(child)
    return _render_annotated_inline_html(content_nodes)


def _append_annotated_html_blocks(story, nodes, styles, default_style_name='annotatedBody', list_level=0):
    inline_buffer = []

    def flush_inline_buffer():
        nonlocal inline_buffer
        html_value = _render_annotated_inline_html(inline_buffer)
        inline_buffer = []
        if _annotated_inline_has_text(html_value):
            story.append(Paragraph(html_value, styles[default_style_name]))

    for node in nodes or []:
        if isinstance(node, str):
            inline_buffer.append(node)
            continue

        tag = node.tag
        if tag in ANNOTATED_NOTES_INLINE_TAGS or tag == 'br':
            inline_buffer.append(node)
            continue

        if tag in ANNOTATED_NOTES_CONTAINER_TAGS:
            if _annotated_node_has_block_children(node):
                flush_inline_buffer()
                _append_annotated_html_blocks(story, node.children, styles, default_style_name=default_style_name, list_level=list_level)
            else:
                inline_buffer.append(node)
            continue

        flush_inline_buffer()

        if tag == 'p':
            paragraph_html = _render_annotated_inline_html(node.children)
            if _annotated_inline_has_text(paragraph_html):
                story.append(Paragraph(paragraph_html, styles[default_style_name]))
            continue

        if tag in {'h1', 'h2', 'h3', 'h4', 'h5', 'h6'}:
            heading_html = _render_annotated_inline_html(node.children)
            if _annotated_inline_has_text(heading_html):
                story.append(Paragraph(heading_html, styles[_annotated_heading_style_name(tag)]))
            continue

        if tag in {'ul', 'ol'}:
            _append_annotated_html_list(story, node, styles, level=list_level, ordered=(tag == 'ol'))
            continue

        if tag == 'blockquote':
            _append_annotated_html_blocks(story, node.children, styles, default_style_name='annotatedQuote', list_level=list_level)
            continue

        if tag == 'pre':
            pre_text = html.escape(''.join(child if isinstance(child, str) else _render_annotated_inline_html(child.children) for child in node.children or []))
            pre_text = pre_text.replace('\n', '<br/>')
            if _annotated_inline_has_text(pre_text):
                story.append(Paragraph(pre_text, styles['annotatedCode']))
            continue

        if tag == 'hr':
            divider = Table([['']], colWidths=[160 * mm], rowHeights=[0.6 * mm], hAlign='LEFT')
            divider.setStyle(
                TableStyle(
                    [
                        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#E5E7EB')),
                        ('LEFTPADDING', (0, 0), (-1, -1), 0),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                        ('TOPPADDING', (0, 0), (-1, -1), 0),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
                    ]
                )
            )
            story.append(divider)
            continue

        _append_annotated_html_blocks(story, node.children, styles, default_style_name=default_style_name, list_level=list_level)

    flush_inline_buffer()


def _append_annotated_html_list(story, list_node, styles, level=0, ordered=False):
    item_index = 1
    for child in list_node.children or []:
        if not isinstance(child, _AnnotatedHtmlNode):
            continue
        if child.tag != 'li':
            _append_annotated_html_blocks(story, [child], styles, default_style_name='annotatedBody', list_level=level)
            continue

        item_html = _render_annotated_list_item_html(child)
        if _annotated_inline_has_text(item_html):
            story.append(
                Paragraph(
                    item_html,
                    styles[_annotated_list_style_name(level)],
                    bulletText=f'{item_index}.' if ordered else '•',
                )
            )

        nested_lists = list(_iter_annotated_nested_lists(child))
        if nested_lists:
            for nested in nested_lists:
                _append_annotated_html_list(story, nested, styles, level=level + 1, ordered=(nested.tag == 'ol'))
        item_index += 1


def build_annotated_notes_pdf(title, annotated_html, runtime=None):
    _ = runtime
    if not REPORTLAB_AVAILABLE:
        raise RuntimeError("PDF export requires the optional 'reportlab' dependency. Install it with: pip install reportlab==4.2.5")

    safe_title = str(title or 'Annotated Notes').strip() or 'Annotated Notes'
    html_root = _parse_annotated_notes_html(annotated_html)

    pdf_buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        pdf_buffer,
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=16 * mm,
        title=safe_title,
    )

    base_styles = getSampleStyleSheet()
    styles = {
        'annotatedTitle': ParagraphStyle('AnnotatedPdfTitle', parent=base_styles['Heading1'], fontName='Helvetica-Bold', fontSize=18, leading=22, textColor=colors.HexColor('#0F172A'), spaceAfter=2),
        'annotatedMeta': ParagraphStyle('AnnotatedPdfMeta', parent=base_styles['BodyText'], fontName='Helvetica', fontSize=9.5, leading=12, textColor=colors.HexColor('#6B7280'), spaceAfter=10),
        'annotatedBody': ParagraphStyle('AnnotatedPdfBody', parent=base_styles['BodyText'], fontName='Helvetica', fontSize=11, leading=15.2, textColor=colors.HexColor('#111827'), spaceAfter=7),
        'annotatedH1': ParagraphStyle('AnnotatedPdfH1', parent=base_styles['Heading2'], fontName='Helvetica-Bold', fontSize=14.5, leading=18, textColor=colors.HexColor('#111827'), spaceBefore=10, spaceAfter=5),
        'annotatedH2': ParagraphStyle('AnnotatedPdfH2', parent=base_styles['Heading3'], fontName='Helvetica-Bold', fontSize=12.4, leading=16, textColor=colors.HexColor('#1F2937'), spaceBefore=8, spaceAfter=4),
        'annotatedH3': ParagraphStyle('AnnotatedPdfH3', parent=base_styles['Heading4'], fontName='Helvetica-Bold', fontSize=11.4, leading=14.5, textColor=colors.HexColor('#334155'), spaceBefore=7, spaceAfter=3),
        'annotatedQuote': ParagraphStyle('AnnotatedPdfQuote', parent=base_styles['BodyText'], fontName='Helvetica', fontSize=10.5, leading=14.5, textColor=colors.HexColor('#374151'), leftIndent=10, borderPadding=8, borderWidth=0, borderLeftWidth=2, borderColor=colors.HexColor('#CBD5E1'), spaceAfter=7),
        'annotatedCode': ParagraphStyle('AnnotatedPdfCode', parent=base_styles['Code'], fontName='Courier', fontSize=9.5, leading=12.5, textColor=colors.HexColor('#111827'), backColor=colors.HexColor('#F8FAFC'), borderPadding=8, spaceAfter=7),
        'annotatedList1': ParagraphStyle('AnnotatedPdfList1', parent=base_styles['BodyText'], fontName='Helvetica', fontSize=11, leading=15.2, textColor=colors.HexColor('#111827'), leftIndent=18, bulletIndent=4, spaceBefore=0, spaceAfter=4),
        'annotatedList2': ParagraphStyle('AnnotatedPdfList2', parent=base_styles['BodyText'], fontName='Helvetica', fontSize=10.7, leading=14.7, textColor=colors.HexColor('#111827'), leftIndent=30, bulletIndent=16, spaceBefore=0, spaceAfter=3),
        'annotatedList3': ParagraphStyle('AnnotatedPdfList3', parent=base_styles['BodyText'], fontName='Helvetica', fontSize=10.4, leading=14.2, textColor=colors.HexColor('#111827'), leftIndent=42, bulletIndent=28, spaceBefore=0, spaceAfter=3),
    }

    story = [
        Paragraph(markdown_inline_to_pdf_html(safe_title), styles['annotatedTitle']),
        Paragraph('Annotated notes export', styles['annotatedMeta']),
    ]

    header_divider = Table([['']], colWidths=[doc.width], rowHeights=[0.8 * mm], hAlign='LEFT')
    header_divider.setStyle(
        TableStyle(
            [
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#4F46E5')),
                ('LEFTPADDING', (0, 0), (-1, -1), 0),
                ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                ('TOPPADDING', (0, 0), (-1, -1), 0),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
            ]
        )
    )
    story.append(header_divider)
    story.append(Spacer(1, 9))

    _append_annotated_html_blocks(story, html_root.children, styles)
    if len(story) <= 4:
        story.append(Paragraph('No annotated notes available.', styles['annotatedBody']))

    def _draw_page_chrome(canvas, pdf_doc):
        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor('#E5E7EB'))
        canvas.setLineWidth(0.5)
        top_rule_y = pdf_doc.pagesize[1] - 11 * mm
        canvas.line(pdf_doc.leftMargin, top_rule_y, pdf_doc.pagesize[0] - pdf_doc.rightMargin, top_rule_y)
        canvas.setFont('Helvetica', 8.5)
        canvas.setFillColor(colors.HexColor('#64748B'))
        canvas.drawString(pdf_doc.leftMargin, 9 * mm, safe_title[:72])
        canvas.drawRightString(pdf_doc.pagesize[0] - pdf_doc.rightMargin, 9 * mm, f'Page {pdf_doc.page}')
        canvas.restoreState()

    doc.build(story, onFirstPage=_draw_page_chrome, onLaterPages=_draw_page_chrome)
    pdf_buffer.seek(0)
    return pdf_buffer


def normalize_exam_date(raw_value, runtime=None):
    _ = runtime
    exam_date = str(raw_value or '').strip()
    if not exam_date:
        return ''
    try:
        return datetime.strptime(exam_date, '%Y-%m-%d').strftime('%Y-%m-%d')
    except ValueError:
        raise ValueError('Exam date must use YYYY-MM-DD format')


def markdown_inline_to_pdf_html(text, runtime=None):
    _ = runtime
    safe_text = html.escape(str(text or ''))
    safe_text = re.sub('\\*\\*(.+?)\\*\\*', '<b>\\1</b>', safe_text)
    safe_text = re.sub('\\*(.+?)\\*', '<i>\\1</i>', safe_text)
    return safe_text


def append_notes_markdown_to_story(story, notes_markdown, styles, runtime=None):
    resolved_runtime = _resolve_runtime(runtime)
    lines = str(notes_markdown or '').splitlines()
    bullet_items = []

    def flush_bullets():
        nonlocal bullet_items
        if not bullet_items:
            return
        list_flow = ListFlowable(
            [ListItem(Paragraph(item, styles['pdfBody']), leftIndent=6) for item in bullet_items],
            bulletType='bullet',
            leftIndent=14,
            bulletFontSize=8,
            bulletOffsetY=1,
        )
        story.append(list_flow)
        story.append(Spacer(1, 4))
        bullet_items = []

    for raw_line in lines:
        line = raw_line.strip()
        line = re.sub('^[-*•]\\s+(\\d+[\\.\\)]\\s+)', '\\1', line)
        if not line:
            flush_bullets()
            story.append(Spacer(1, 4))
            continue

        heading_level = 0
        if line.startswith('### '):
            heading_level = 3
        elif line.startswith('## '):
            heading_level = 2
        elif line.startswith('# '):
            heading_level = 1

        if heading_level:
            flush_bullets()
            heading_text = markdown_inline_to_pdf_html(line[heading_level + 1 :], runtime=resolved_runtime)
            heading_style = styles['pdfH1'] if heading_level == 1 else styles['pdfH2'] if heading_level == 2 else styles['pdfH3']
            story.append(Paragraph(heading_text, heading_style))
            story.append(Spacer(1, 3))
            continue

        if line.startswith('- ') or line.startswith('* '):
            bullet_items.append(markdown_inline_to_pdf_html(line[2:].strip(), runtime=resolved_runtime))
            continue

        numbered_match = re.match('^(\\d+)\\.\\s+(.*)$', line)
        if numbered_match:
            flush_bullets()
            text_html = markdown_inline_to_pdf_html(numbered_match.group(2), runtime=resolved_runtime)
            story.append(Paragraph(f"{numbered_match.group(1)}. {text_html}", styles['pdfBody']))
            story.append(Spacer(1, 2))
            continue

        flush_bullets()
        story.append(Paragraph(markdown_inline_to_pdf_html(line, runtime=resolved_runtime), styles['pdfBody']))
        story.append(Spacer(1, 2))

    flush_bullets()


def build_study_pack_pdf(pack, include_answers=True, runtime=None):
    _ = runtime
    if not REPORTLAB_AVAILABLE:
        raise RuntimeError("PDF export requires the optional 'reportlab' dependency. Install it with: pip install reportlab==4.2.5")

    pdf_buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        pdf_buffer,
        pagesize=A4,
        leftMargin=16 * mm,
        rightMargin=16 * mm,
        topMargin=14 * mm,
        bottomMargin=14 * mm,
        title=str(pack.get('title', 'Study Pack')).strip() or 'Study Pack',
    )

    base_styles = getSampleStyleSheet()
    styles = {
        'pdfTitle': ParagraphStyle('PdfTitle', parent=base_styles['Heading1'], fontName='Helvetica-Bold', fontSize=17, leading=21, spaceAfter=6, textColor=colors.HexColor('#111827')),
        'pdfMeta': ParagraphStyle('PdfMeta', parent=base_styles['BodyText'], fontName='Helvetica', fontSize=9.5, leading=12.5, textColor=colors.HexColor('#4B5563')),
        'pdfSection': ParagraphStyle('PdfSection', parent=base_styles['Heading2'], fontName='Helvetica-Bold', fontSize=12.5, leading=16, spaceBefore=6, spaceAfter=6, textColor=colors.HexColor('#111827')),
        'pdfH1': ParagraphStyle('PdfH1', parent=base_styles['Heading2'], fontName='Helvetica-Bold', fontSize=12, leading=15, textColor=colors.HexColor('#1F2937')),
        'pdfH2': ParagraphStyle('PdfH2', parent=base_styles['Heading3'], fontName='Helvetica-Bold', fontSize=11, leading=14, textColor=colors.HexColor('#1F2937')),
        'pdfH3': ParagraphStyle('PdfH3', parent=base_styles['Heading4'], fontName='Helvetica-Bold', fontSize=10, leading=13, textColor=colors.HexColor('#374151')),
        'pdfBody': ParagraphStyle('PdfBody', parent=base_styles['BodyText'], fontName='Helvetica', fontSize=9.5, leading=13, textColor=colors.HexColor('#111827')),
        'pdfQuestion': ParagraphStyle('PdfQuestion', parent=base_styles['BodyText'], fontName='Helvetica-Bold', fontSize=10, leading=13.5, textColor=colors.HexColor('#111827')),
        'pdfOption': ParagraphStyle('PdfOption', parent=base_styles['BodyText'], fontName='Helvetica', fontSize=9.5, leading=12.5, leftIndent=10, textColor=colors.HexColor('#1F2937')),
        'pdfOptionCorrect': ParagraphStyle('PdfOptionCorrect', parent=base_styles['BodyText'], fontName='Helvetica-Bold', fontSize=9.5, leading=12.5, leftIndent=10, textColor=colors.HexColor('#065F46')),
    }

    pack_title = str(pack.get('title', 'Study Pack')).strip() or 'Study Pack'
    story = [Paragraph(markdown_inline_to_pdf_html(pack_title), styles['pdfTitle'])]

    mode = str(pack.get('mode', '') or '').strip() or 'Unknown'
    output_language = str(pack.get('output_language', '') or '').strip() or 'Unknown'
    course = str(pack.get('course', '') or '').strip() or '-'
    subject = str(pack.get('subject', '') or '').strip() or '-'
    semester = str(pack.get('semester', '') or '').strip() or '-'
    block = str(pack.get('block', '') or '').strip() or '-'

    created_at = pack.get('created_at', 0)
    created_text = '-'
    try:
        if created_at:
            created_text = datetime.fromtimestamp(float(created_at)).strftime('%Y-%m-%d %H:%M')
    except Exception:
        created_text = '-'

    metadata_rows = [
        [Paragraph('<b>Mode</b>', styles['pdfMeta']), Paragraph(markdown_inline_to_pdf_html(mode), styles['pdfMeta'])],
        [Paragraph('<b>Language</b>', styles['pdfMeta']), Paragraph(markdown_inline_to_pdf_html(output_language), styles['pdfMeta'])],
        [Paragraph('<b>Course</b>', styles['pdfMeta']), Paragraph(markdown_inline_to_pdf_html(course), styles['pdfMeta'])],
        [Paragraph('<b>Subject</b>', styles['pdfMeta']), Paragraph(markdown_inline_to_pdf_html(subject), styles['pdfMeta'])],
        [Paragraph('<b>Semester / Block</b>', styles['pdfMeta']), Paragraph(markdown_inline_to_pdf_html(f'{semester} / {block}'), styles['pdfMeta'])],
        [Paragraph('<b>Created</b>', styles['pdfMeta']), Paragraph(markdown_inline_to_pdf_html(created_text), styles['pdfMeta'])],
    ]

    metadata_table = Table(metadata_rows, colWidths=[36 * mm, 145 * mm], hAlign='LEFT')
    metadata_table.setStyle(
        TableStyle(
            [
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('LEFTPADDING', (0, 0), (-1, -1), 4),
                ('RIGHTPADDING', (0, 0), (-1, -1), 4),
                ('TOPPADDING', (0, 0), (-1, -1), 2),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
                ('GRID', (0, 0), (-1, -1), 0.35, colors.HexColor('#E5E7EB')),
                ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#F9FAFB')),
            ]
        )
    )

    story.append(metadata_table)
    story.append(Spacer(1, 10))

    story.append(Paragraph('Integrated Notes', styles['pdfSection']))
    notes_markdown = str(pack.get('notes_markdown', '') or '').strip()
    if notes_markdown:
        append_notes_markdown_to_story(story, notes_markdown, styles)
    else:
        story.append(Paragraph('No integrated notes available.', styles['pdfBody']))

    story.append(Spacer(1, 10))
    story.append(Paragraph('Flashcards', styles['pdfSection']))
    flashcards = pack.get('flashcards', []) if isinstance(pack.get('flashcards', []), list) else []
    if flashcards:
        card_rows = [[Paragraph('<b>Front</b>', styles['pdfMeta']), Paragraph('<b>Back</b>', styles['pdfMeta'])]]
        for card in flashcards:
            card_rows.append(
                [
                    Paragraph(markdown_inline_to_pdf_html(str(card.get('front', '') or '')), styles['pdfBody']),
                    Paragraph(markdown_inline_to_pdf_html(str(card.get('back', '') or '')), styles['pdfBody']),
                ]
            )
        flashcard_table = Table(card_rows, colWidths=[84 * mm, 97 * mm], repeatRows=1, hAlign='LEFT')
        flashcard_table.setStyle(
            TableStyle(
                [
                    ('GRID', (0, 0), (-1, -1), 0.35, colors.HexColor('#D1D5DB')),
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#F3F4F6')),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 5),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 5),
                    ('TOPPADDING', (0, 0), (-1, -1), 4),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                ]
            )
        )
        story.append(flashcard_table)
    else:
        story.append(Paragraph('No flashcards available.', styles['pdfBody']))

    story.append(PageBreak())
    practice_title = 'Practice Questions'
    if not include_answers:
        practice_title += ' (Without Answers)'
    story.append(Paragraph(practice_title, styles['pdfSection']))

    questions = pack.get('test_questions', []) if isinstance(pack.get('test_questions', []), list) else []
    if questions:
        for idx, question in enumerate(questions, 1):
            question_text = str(question.get('question', '') or '').strip() or f'Question {idx}'
            story.append(Paragraph(f'{idx}. {markdown_inline_to_pdf_html(question_text)}', styles['pdfQuestion']))
            options = question.get('options', [])
            if not isinstance(options, list):
                options = []
            answer = str(question.get('answer', '') or '').strip()
            letters = ['A', 'B', 'C', 'D']
            for option_idx, option in enumerate(options[:4]):
                option_text = str(option or '').strip()
                is_correct = include_answers and option_text == answer and (option_text != '')
                marker = '✓' if is_correct else '•'
                letter = letters[option_idx] if option_idx < len(letters) else str(option_idx + 1)
                option_style = styles['pdfOptionCorrect'] if is_correct else styles['pdfOption']
                story.append(Paragraph(f'{marker} {letter}. {markdown_inline_to_pdf_html(option_text)}', option_style))
            explanation = str(question.get('explanation', '') or '').strip()
            if include_answers and explanation:
                story.append(Paragraph(f'<b>Explanation:</b> {markdown_inline_to_pdf_html(explanation)}', styles['pdfBody']))
            story.append(Spacer(1, 7))
    else:
        story.append(Paragraph('No practice questions available.', styles['pdfBody']))

    doc.build(story)
    pdf_buffer.seek(0)
    return pdf_buffer
