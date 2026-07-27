import io
import json
import zipfile
from pathlib import Path

import pytest
from docx import Document

from lecture_processor.physio_companion import CompanionConfig, create_companion_app
from lecture_processor.physio_companion.sources import (
    SourceConflict,
    SourceManager,
    SourceManagerError,
    SourceTooLarge,
    classify_source,
)


@pytest.fixture()
def manager(tmp_path):
    vault = tmp_path / "vault"
    support = tmp_path / "support"
    vault.mkdir()
    return SourceManager(
        support / "source-manifest.json",
        support / "Sources",
        vault,
        max_source_bytes=1024,
    )


def test_filename_classification_and_explicit_category_override(manager):
    assert classify_source("KNGF Richtlijn knie.pdf") == "guidelines"
    assert classify_source("KNGF_test_richtlijn.txt") == "guidelines"
    assert classify_source("Semester 2 samenvatting.docx") == "semester-summaries"
    assert classify_source("Anatomie spieren.xlsx") == "anatomy"
    assert classify_source("Craft schouder.md") == "craft"
    assert classify_source("Hoorcollege pijn.txt") == "lectures"
    assert classify_source("Kenhub boek.pdf") == "books"
    assert classify_source("los bestand.pdf") == "other"

    imported = manager.import_stream(io.BytesIO(b"inhoud"), "KNGF Richtlijn.txt", category="craft")
    assert imported["source"]["category"] == "craft"
    assert Path(imported["source"]["local_path"]).parent.name == "craft"


def test_import_copies_deduplicates_and_creates_draft_source_note(manager, tmp_path):
    original = tmp_path / "College anatomie.txt"
    original.write_text("onbewerkte inhoud die niet in de kennisindex mag", encoding="utf-8")

    with original.open("rb") as source:
        created = manager.import_stream(source, original.name)
    record = created["source"]
    managed_path = Path(record["local_path"])
    note_path = manager.vault_path / record["knowledge_note_path"]

    assert created["deduplicated"] is False
    assert original.exists()
    assert managed_path.read_bytes() == original.read_bytes()
    assert manager.sources_root in managed_path.parents
    assert record["extraction_status"] == "text-ready"
    assert "curation_status: draft" in note_path.read_text(encoding="utf-8")
    assert f"id: {record['knowledge_note_id']}" in note_path.read_text(encoding="utf-8")
    assert "trust_tier: semester_summary" in note_path.read_text(encoding="utf-8")
    assert "onbewerkte inhoud" not in note_path.read_text(encoding="utf-8")
    duplicate = manager.import_stream(io.BytesIO(original.read_bytes()), "andere-naam.txt")
    assert duplicate["deduplicated"] is True
    assert duplicate["source"]["id"] == record["id"]
    assert manager.list_sources()["total"] == 1


def test_review_update_move_and_delete_are_managed_only(manager):
    created = manager.import_stream(io.BytesIO(b"pdf"), "los.pdf")
    source_id = created["source"]["id"]

    activated = manager.set_review_status(source_id, "activate")
    note_path = manager.vault_path / activated["knowledge_note_path"]
    assert activated["review_status"] == "active"
    assert "curation_status: reviewed" in note_path.read_text(encoding="utf-8")
    assert "last_reviewed: null" not in note_path.read_text(encoding="utf-8")

    updated = manager.update(source_id, {"category": "books", "title": "Handboek", "trust_tier": "evidence-publication"})
    assert Path(updated["local_path"]).parent.name == "books"
    assert updated["trust_tier"] == 400
    assert updated["trust_label"] == "evidence_publication"

    assert manager.delete(source_id) is True
    assert not Path(updated["local_path"]).exists()
    assert not note_path.exists()
    assert manager.get(source_id) is None

    external = manager.sources_root.parent / "extern.pdf"
    external.write_bytes(b"extern")
    manifest = {"schema_version": 1, "sources": [{
        "id": "external-source", "sha256": "x", "local_path": str(external),
        "source_type": "guideline", "privacy_class": "private-local",
    }]}
    manager.manifest_path.write_text(json.dumps(manifest), encoding="utf-8")
    listed = manager.list_sources()["sources"][0]
    assert listed["review_status"] == "pending"
    assert listed["managed"] is False
    with pytest.raises(SourceConflict):
        manager.delete("external-source")
    assert external.exists()


def test_legacy_source_is_active_only_when_a_reviewed_note_cites_it(manager):
    external = manager.sources_root.parent / "richtlijn.pdf"
    external.write_bytes(b"richtlijn")
    manager.manifest_path.write_text(json.dumps({"sources": [{
        "id": "src-reviewed-guideline", "sha256": "reviewed", "local_path": str(external),
        "source_type": "guideline", "privacy_class": "private-local",
    }]}), encoding="utf-8")
    note = manager.vault_path / "portaal.md"
    note.write_text(
        "---\nid: portal\ntype: region\nregions: [knie]\nsystems: [musculoskeletal]\n"
        "clinical_use: [navigation]\naliases: []\ncuration_status: reviewed\n"
        "source_refs: [\"src-reviewed-guideline#page=2\"]\ntrust_tier: guideline\n"
        "last_reviewed: 2026-07-19\n---\n# Knie\n",
        encoding="utf-8",
    )

    listed = manager.list_sources()["sources"][0]
    assert listed["review_status"] == "active"
    assert listed["regions"] == ["knie"]

    manually_unlinked = manager.update("src-reviewed-guideline", {"regions": []})
    assert manually_unlinked["regions"] == []

    note.write_text(note.read_text(encoding="utf-8").replace("reviewed", "draft", 1), encoding="utf-8")
    assert manager.list_sources()["sources"][0]["review_status"] == "pending"


def test_active_docx_source_is_extracted_ranked_and_available_as_virtual_note(tmp_path):
    vault = tmp_path / "vault"
    support = tmp_path / "support"
    vault.mkdir()
    source_path = tmp_path / "Alle pijncolleges.docx"
    document = Document()
    document.add_heading("Pijn en neurofysiologie", level=1)
    document.add_paragraph("De fasen zijn transductie, conductie, transmissie, modulatie en perceptie.")
    document.add_heading("Conductie", level=2)
    document.add_paragraph("Na transductie volgt conductie: het signaal wordt naar het ruggenmerg geleid.")
    document.save(source_path)
    manifest = support / "source-manifest.json"
    support.mkdir()
    manifest.write_text(json.dumps({"sources": [{
        "id": "src-pain", "path": str(source_path), "local_path": str(source_path),
        "source_type": "semester-summary", "review_status": "active", "regions": ["pijn"],
    }]}), encoding="utf-8")
    source_manager = SourceManager(manifest, support / "Sources", vault, max_source_bytes=50_000_000)

    context = source_manager.deep_context("Wat komt er na transductie?", region="pijn", max_chars=20_000)

    assert context[0]["note_id"] == "source--src-pain"
    assert "conductie" in context[0]["body"].casefold()
    assert context[0]["allowed_anchors"]
    assert source_manager.get("src-pain")["extraction_status"] == "indexed"
    virtual = source_manager.virtual_note("source--src-pain")
    assert virtual["title"] == "Alle pijncolleges.docx"
    assert virtual["source_uri"].endswith("/src-pain")
    assert "Na transductie volgt conductie" in virtual["body"]


def test_craft_virtual_note_maps_relative_images_to_manifest_preview_ids(tmp_path):
    vault = tmp_path / "vault"
    support = tmp_path / "support"
    assets = tmp_path / "Craft.assets"
    vault.mkdir()
    support.mkdir()
    assets.mkdir()
    note_path = tmp_path / "Craft.md"
    image_path = assets / "Image.png"
    second_image_path = assets / "Image (2).png"
    note_path.write_text(
        "# Craft\n\n![Foto](Craft.assets/Image.png)\n\n![Tweede](Craft.assets/Image%20(2).png)\n",
        encoding="utf-8",
    )
    image_path.write_bytes(b"image")
    second_image_path.write_bytes(b"second")
    (support / "source-manifest.json").write_text(json.dumps({"sources": [
        {"id": "src-note", "local_path": str(note_path), "source_type": "craft-note"},
        {"id": "src-image", "local_path": str(image_path), "mime_type": "image/png"},
        {"id": "src-second", "local_path": str(second_image_path), "mime_type": "image/png"},
    ]}), encoding="utf-8")
    source_manager = SourceManager(support / "source-manifest.json", support / "Sources", vault, max_source_bytes=1024)

    virtual = source_manager.virtual_note("source--src-note")

    assert virtual["inline_assets"] == {
        "Craft.assets/Image.png": "src-image",
        "Craft.assets/Image%20(2).png": "src-second",
    }
    assert source_manager.preview_file("src-image")["path"] == image_path.resolve()


def test_auto_triage_classifies_only_clear_sources_and_prefers_docx_pairs(tmp_path):
    vault = tmp_path / "vault"
    support = tmp_path / "support"
    vault.mkdir()
    support.mkdir()
    files = {
        "parkinson": tmp_path / "KNGF-richtlijn Parkinson.pdf",
        "meniscus": tmp_path / "meniscectomie-praktijkrichtlijn.pdf",
        "peripheral_docx": tmp_path / "blok 1.4" / "college week 1.docx",
        "peripheral_txt": tmp_path / "blok 1.4" / "college week 1.txt",
        "feedback": tmp_path / "Feedback studentmentoren blok 2.1.md",
        "unclear": tmp_path / "los bestand.pdf",
    }
    for path in files.values():
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_bytes(b"bron")
    records = [
        {"id": f"src-{name}", "local_path": str(path), "source_type": "guideline" if name in {"parkinson", "meniscus"} else "lecture"}
        for name, path in files.items()
    ]
    (support / "source-manifest.json").write_text(json.dumps({"sources": records}), encoding="utf-8")
    source_manager = SourceManager(support / "source-manifest.json", support / "Sources", vault, max_source_bytes=1024)

    result = source_manager.auto_triage()

    assert result["changed"] == 3
    assert source_manager.get("src-parkinson")["regions"] == ["centraal-neurologisch"]
    assert source_manager.get("src-meniscus")["regions"] == ["knie"]
    assert source_manager.get("src-peripheral_docx")["regions"] == ["perifeer-neurologisch"]
    assert source_manager.get("src-peripheral_txt")["review_status"] == "pending"
    assert source_manager.get("src-feedback")["review_status"] == "pending"
    assert source_manager.get("src-unclear")["review_status"] == "pending"


def test_delete_keeps_file_if_atomic_manifest_write_fails(manager, monkeypatch):
    created = manager.import_stream(io.BytesIO(b"keep me"), "bron.txt")["source"]
    managed_path = Path(created["local_path"])

    def fail_write(*_args, **_kwargs):
        raise OSError("disk full")

    monkeypatch.setattr(manager, "_write_records", fail_write)
    with pytest.raises(OSError):
        manager.delete(created["id"])
    assert managed_path.exists()
    assert manager.get(created["id"]) is not None


def test_upload_rejects_unsafe_type_category_and_size(manager):
    with pytest.raises(SourceManagerError):
        manager.import_stream(io.BytesIO(b"x"), "malware.exe")
    with pytest.raises(SourceManagerError):
        manager.import_stream(io.BytesIO(b"x"), "ok.txt", category="../../outside")
    with pytest.raises(SourceTooLarge):
        manager.import_stream(io.BytesIO(b"x" * 1025), "groot.txt")
    assert not list(manager.sources_root.rglob(".upload-*"))


def test_docx_extraction_rejects_zip_bomb_compression_ratio(tmp_path):
    from lecture_processor.physio_companion.sources import SourceManager, SourceManagerError

    path = tmp_path / "compressed.docx"
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("word/document.xml", b"0" * (1024 * 1024))
    with pytest.raises(SourceManagerError, match="compressieverhouding"):
        SourceManager._extract_docx(path)


def test_pdf_extraction_rejects_excessive_page_count(tmp_path, monkeypatch):
    from lecture_processor.physio_companion import sources as source_module

    class _Reader:
        pages = [object()] * (source_module.MAX_PDF_PAGES + 1)

    monkeypatch.setattr(source_module, "PdfReader", lambda _path: _Reader())
    with pytest.raises(source_module.SourceManagerError, match="meer dan"):
        source_module.SourceManager._extract_pdf(tmp_path / "huge.pdf")


@pytest.fixture()
def source_client(tmp_path):
    vault = tmp_path / "vault"
    support = tmp_path / "support"
    vault.mkdir()
    config = CompanionConfig(
        vault_path=vault,
        app_support_path=support,
        secret_key="sources-test",
        max_source_bytes=2048,
    )
    app = create_companion_app(config)
    app.config.update(TESTING=True)
    with app.test_client() as client:
        assert client.post(
            "/owner-session", json={"owner_token": config.owner_token}
        ).status_code == 200
        yield client, app.extensions["physio_companion"]
    app.extensions["physio_companion"].close()


def _csrf(client):
    return client.get("/api/local/physio/csrf").get_json()["csrf_token"]


def test_source_manager_api_crud_review_and_csrf(source_client):
    client, _service = source_client
    endpoint = "/api/local/physio/sources-manager"
    without_csrf = client.post(
        endpoint + "/upload",
        data={"file": (io.BytesIO(b"bron"), "KNGF richtlijn.txt")},
    )
    assert without_csrf.status_code == 403
    headers = {"X-CSRF-Token": _csrf(client)}
    uploaded = client.post(
        endpoint + "/upload",
        data={"file": (io.BytesIO(b"bron"), "KNGF richtlijn.txt")},
        headers=headers,
    )
    assert uploaded.status_code == 201
    source = uploaded.get_json()["source"]
    assert source["category"] == "guidelines"
    assert source["managed"] is True
    assert source["obsidian_uri"].startswith("obsidian://open?")

    listing = client.get(endpoint + "?managed=true&category=guidelines").get_json()
    assert listing["total"] == 1
    assert {item["id"] for item in listing["categories"]} >= {"guidelines", "books", "other"}
    patched = client.patch(
        endpoint + "/" + source["id"],
        json={"title": "Nieuwe titel", "trust_tier": "guideline"},
        headers=headers,
    )
    assert patched.get_json()["title"] == "Nieuwe titel"
    rejected = client.post(
        endpoint + "/" + source["id"] + "/review",
        json={"action": "reject"},
        headers=headers,
    )
    assert rejected.get_json()["review_status"] == "rejected"
    assert client.get("/api/local/physio/media/" + source["id"]).status_code == 404
    activated = client.post(
        endpoint + "/" + source["id"] + "/review",
        json={"action": "activate"},
        headers=headers,
    )
    assert activated.get_json()["review_status"] == "active"
    assert client.get("/api/local/physio/media/" + source["id"]).status_code == 200
    virtual = client.get("/api/local/physio/notes/source--" + source["id"])
    assert virtual.status_code == 200
    assert virtual.get_json()["source_uri"].endswith("/" + source["id"])
    search = client.get("/api/local/physio/search?q=Nieuwe").get_json()["results"]
    assert search[0]["note_id"] == source["knowledge_note_id"]
    assert client.get(endpoint + "/" + source["id"]).status_code == 200
    assert client.delete(endpoint + "/" + source["id"], headers=headers).status_code == 204
    assert client.get(endpoint + "/" + source["id"]).status_code == 404


def test_source_manager_preview_is_available_while_source_is_pending(source_client):
    client, _service = source_client
    endpoint = "/api/local/physio/sources-manager"
    uploaded = client.post(
        endpoint + "/upload",
        data={"file": (io.BytesIO(b"not-a-real-png-but-streamable"), "schouder.png")},
        headers={"X-CSRF-Token": _csrf(client)},
    ).get_json()["source"]

    preview = client.get(endpoint + "/" + uploaded["id"] + "/preview", headers={"Range": "bytes=0-3"})

    assert uploaded["review_status"] == "pending"
    assert preview.status_code == 206
    assert preview.data == b"not-"
    assert preview.headers["Content-Type"].startswith("image/png")


def test_source_manager_auto_triage_endpoint_requires_csrf(source_client, tmp_path):
    client, service = source_client
    source = tmp_path / "Parkinson richtlijn.pdf"
    source.write_bytes(b"pdf")
    service.sources.manifest_path.write_text(json.dumps({"sources": [{
        "id": "src-parkinson", "local_path": str(source), "source_type": "guideline"
    }]}), encoding="utf-8")
    endpoint = "/api/local/physio/sources-manager/auto-triage"

    assert client.post(endpoint).status_code == 403
    response = client.post(endpoint, json={}, headers={"X-CSRF-Token": _csrf(client)})

    assert response.status_code == 200
    assert response.get_json()["changed"] == 1
    assert service.sources.get("src-parkinson")["regions"] == ["centraal-neurologisch"]
