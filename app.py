import os
import uuid
import threading
import time
import io
import json
import csv
import re
import shutil
import subprocess
import html
import warnings
from datetime import datetime, timedelta

# Keep startup clean in local dev environments.
warnings.filterwarnings(
    "ignore",
    message=r"urllib3 v2 only supports OpenSSL 1\.1\.1\+.*"
)
warnings.filterwarnings(
    "ignore",
    message=r"You are using a Python version 3\.9 past its end of life.*",
    category=FutureWarning
)
warnings.filterwarnings(
    "ignore",
    message=r"You are using a non-supported Python version \(3\.9.*\).*",
    category=FutureWarning
)

import stripe
from flask import Flask, request, jsonify, render_template, send_file
from google import genai
from google.genai import types
from dotenv import load_dotenv
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPORTLAB_AVAILABLE = True
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, ListFlowable, ListItem, PageBreak
except Exception:
    REPORTLAB_AVAILABLE = False
import firebase_admin
from firebase_admin import credentials, auth, firestore

load_dotenv()
app = Flask(__name__)

UPLOAD_FOLDER = 'uploads'
ALLOWED_PDF_EXTENSIONS = {'pdf'}
ALLOWED_AUDIO_EXTENSIONS = {'mp3', 'm4a', 'wav', 'aac', 'ogg', 'flac'}
MAX_CONTENT_LENGTH = 500 * 1024 * 1024

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
client = genai.Client(api_key=os.getenv('GEMINI_API_KEY'))

# --- Firebase Setup ---
if os.path.exists('firebase-credentials.json'):
    cred = credentials.Certificate('firebase-credentials.json')
else:
    firebase_creds = json.loads(os.getenv('FIREBASE_CREDENTIALS', '{}'))
    cred = credentials.Certificate(firebase_creds)

firebase_admin.initialize_app(cred)
db = firestore.client()

# --- Stripe Setup ---
stripe.api_key = os.getenv('STRIPE_SECRET_KEY')
STRIPE_PUBLISHABLE_KEY = os.getenv('STRIPE_PUBLISHABLE_KEY', '')
STRIPE_WEBHOOK_SECRET = os.getenv('STRIPE_WEBHOOK_SECRET', '')
ADMIN_EMAILS = {email.strip().lower() for email in os.getenv('ADMIN_EMAILS', '').split(',') if email.strip()}
ADMIN_UIDS = {uid.strip() for uid in os.getenv('ADMIN_UIDS', '').split(',') if uid.strip()}

# --- In-Memory Storage (jobs only — credits are in Firestore now) ---
jobs = {}
AUDIO_STREAM_TOKEN_TTL_SECONDS = 3600
AUDIO_STREAM_TOKENS = {}
FEATURE_AUDIO_SECTION_SYNC = os.getenv('FEATURE_AUDIO_SECTION_SYNC', '0').strip().lower() in {'1', 'true', 'yes', 'on'}

# --- Credit Bundles (what users can buy) ---
CREDIT_BUNDLES = {
    'lecture_5': {
        'name': 'Lecture Notes — 5 Pack',
        'description': '5 standard lecture credits',
        'credits': {'lecture_credits_standard': 5},
        'price_cents': 999,
        'currency': 'eur',
    },
    'lecture_10': {
        'name': 'Lecture Notes — 10 Pack',
        'description': '10 standard lecture credits (best value)',
        'credits': {'lecture_credits_standard': 10},
        'price_cents': 1699,
        'currency': 'eur',
    },
    'slides_10': {
        'name': 'Slides Extraction — 10 Pack',
        'description': '10 slides extraction credits',
        'credits': {'slides_credits': 10},
        'price_cents': 499,
        'currency': 'eur',
    },
    'slides_25': {
        'name': 'Slides Extraction — 25 Pack',
        'description': '25 slides extraction credits (best value)',
        'credits': {'slides_credits': 25},
        'price_cents': 999,
        'currency': 'eur',
    },
    'interview_3': {
        'name': 'Interview Transcription — 3 Pack',
        'description': '3 interview transcription credits',
        'credits': {'interview_credits_short': 3},
        'price_cents': 799,
        'currency': 'eur',
    },
    'interview_8': {
        'name': 'Interview Transcription — 8 Pack',
        'description': '8 interview transcription credits (best value)',
        'credits': {'interview_credits_short': 8},
        'price_cents': 1799,
        'currency': 'eur',
    },
}

# --- Email Allowlist ---
ALLOWED_EMAIL_DOMAINS = {
    'gmail.com', 'googlemail.com', 'outlook.com', 'hotmail.com', 'live.com', 'icloud.com', 'yahoo.com', 'yahoo.nl',
    'student.tudelft.nl', 'tudelft.nl', 'uva.nl', 'student.uva.nl', 'vu.nl', 'student.vu.nl', 'eur.nl', 'student.eur.nl',
    'uu.nl', 'students.uu.nl', 'ru.nl', 'student.ru.nl', 'utwente.nl', 'student.utwente.nl', 'tue.nl', 'student.tue.nl',
    'maastrichtuniversity.nl', 'student.maastrichtuniversity.nl', 'leidenuniv.nl', 'student.leidenuniv.nl', 'rug.nl',
    'student.rug.nl', 'tilburguniversity.edu', 'uvt.nl', 'han.nl', 'student.han.nl', 'hva.nl', 'student.hva.nl', 'hr.nl',
    'student.hr.nl', 'fontys.nl', 'student.fontys.nl', 'saxion.nl', 'student.saxion.nl', 'avans.nl', 'student.avans.nl',
    'inholland.nl', 'student.inholland.nl', 'hanze.nl', 'student.hanze.nl', 'zuyd.nl', 'student.zuyd.nl', 'hu.nl',
    'student.hu.nl', 'windesheim.nl', 'student.windesheim.nl', 'nhlstenden.com', 'student.nhlstenden.com',
}

ALLOWED_EMAIL_PATTERNS = ['.edu', '.ac.uk', '.edu.nl']

def is_email_allowed(email):
    if not email: return False
    email = email.lower()
    domain = email.split('@')[-1] if '@' in email else ''
    if domain in ALLOWED_EMAIL_DOMAINS: return True
    for pattern in ALLOWED_EMAIL_PATTERNS:
        if domain.endswith(pattern): return True
    return False

# --- AI Model Config ---
MODEL_SLIDES = 'gemini-2.5-flash-lite'
MODEL_AUDIO = 'gemini-3-flash-preview'
MODEL_INTEGRATION = 'gemini-2.5-pro'
MODEL_INTERVIEW = 'gemini-2.5-pro'
MODEL_STUDY = 'gemini-2.5-flash-lite'

FREE_LECTURE_CREDITS = 1
FREE_SLIDES_CREDITS = 2
FREE_INTERVIEW_CREDITS = 0

OUTPUT_LANGUAGE_MAP = {
    'dutch': 'Dutch',
    'english': 'English',
    'spanish': 'Spanish',
    'french': 'French',
    'german': 'German',
    'chinese': 'Chinese',
}

# --- Prompts ---
PROMPT_SLIDE_EXTRACTION = """Extraheer alle tekst van de slides uit het bijgevoegde PDF-bestand en identificeer de functie van visuele elementen.
Instructies:
1. Geef per slide duidelijk aan welk slide-nummer het betreft (bv. "Slide 1:").
2. Neem de titel van de slide over.
3. Neem alle tekstuele inhoud (bullet points, paragrafen) van de slide over.
4. Identificeer waar afbeeldingen of tabellen staan. Gebruik strikte criteria:
   - Informatief: Gebruik de placeholder ALLEEN als de afbeelding tekst, data, grafieken, diagrammen, flowcharts, of een specifiek wetenschappelijk/technisch diagram bevat dat cruciaal is voor begrip van de slide. Formaat: [Informatieve Afbeelding/Tabel: Geef een neutrale beschrijving van wat zichtbaar is of het onderwerp]
   - Decoratief: Gebruik de placeholder voor ALLE foto's van mensen, landschappen, bedrijfslogo's, universiteitslogo's, achtergrondillustraties, stockfoto's, of sfeerbeelden. Bij twijfel, classificeer het als decoratief! Formaat: [Decoratieve Afbeelding]
5. Laat de zin "Share Your talent move the world" weg, indien aanwezig.
6. Lever de output als platte tekst, zonder specifieke Word-opmaak anders dan de slide-indicatie en de placeholders."""

PROMPT_AUDIO_TRANSCRIPTION = """Maak een nauwkeurig en 'schoon' transcript van het bijgevoegde audiobestand.
Instructies:
1. Transcribeer de gesproken tekst zo letterlijk mogelijk.
2. Verwijder stopwoorden en aarzelingen (zoals "eh," "uhm," "nou ja," "weet je wel") om de leesbaarheid te verhogen, maar behoud de volledige inhoudelijke boodschap. Verander geen zinsconstructies.
3. Gebruik geen tijdcodes.
4. Gebruik alinea's om langere spreekbeurten op te delen.
5. Schrijf de uiteindelijke output volledig in deze taal: {output_language}."""

PROMPT_AUDIO_TRANSCRIPTION_TIMESTAMPED = """Maak een nauwkeurig transcript met tijdsegmenten van het bijgevoegde audiobestand.

Geef ALLEEN geldige JSON terug, zonder markdown of extra tekst, in exact dit formaat:
{{
  "transcript_segments": [
    {{
      "start_ms": 0,
      "end_ms": 10000,
      "text": "..."
    }}
  ],
  "full_transcript": "..."
}}

Regels:
- Gebruik natuurlijke segmenten van ongeveer 5-25 seconden.
- start_ms en end_ms zijn milliseconden vanaf het begin.
- Verwijder stopwoorden en aarzelingen om de leesbaarheid te verbeteren zonder inhoud te verliezen.
- full_transcript bevat de volledige transcriptie als doorlopende tekst.
- Schrijf tekstinhoud volledig in deze taal: {output_language}."""

PROMPT_INTERVIEW_TRANSCRIPTION = """Transcribe this interview in the format: timecode (mm:ss) - speaker - caption.
Rules:
- Use speaker A, speaker B, etc. to identify speakers.
- Keep timestamps in each line.
- Write the output fully in this language: {output_language}."""

PROMPT_INTERVIEW_SUMMARY = """You are an expert interviewer analyst.
Create a concise summary of this interview.
Rules:
- Maximum one page equivalent (about 400-600 words).
- Focus only on the most important points, commitments, and conclusions.
- Use short headings and bullet points where useful.
- Do not invent information outside the transcript.
- Write the output fully in this language: {output_language}.
Transcript:
{transcript}
"""

PROMPT_INTERVIEW_SECTIONED = """You are an expert transcript editor.
Rewrite this interview transcript into a structured version with clear headings.
Rules:
- Keep timestamps and speaker labels from the source where possible.
- Split content into relevant sections (for example: Introduction, Background, Key Discussion, Decisions, Next Steps).
- Use meaningful heading titles based on actual content.
- Do not invent information outside the transcript.
- Write the output fully in this language: {output_language}.
Transcript:
{transcript}
"""

PROMPT_MERGE_TEMPLATE = """Maak één complete, consistente en studieklare uitwerking van het college op basis van slide-tekst en audio-transcript.

DOEL:
- Lever een volledig naslagdocument op (geen samenvatting).
- Maak de tekst direct bruikbaar voor studenten bij voorbereiding op toets/tentamen.
- Integreer alle relevante inhoud uit beide bronnen in één samenhangende tekst.

OUTPUTVORM (VERPLICHT):
1. Start direct met inhoud in Markdown (geen inleidende assistent-zin).
2. Eerste regel is een titel met `#`.
3. Gebruik daarna `##` en `###` met duidelijke, logische opbouw.
4. Gebruik geen transcriptvorm met sprekers, dialooglabels of vraag-antwoord stijl.
5. Lever alleen de uiteindelijke tekst; geen toelichting op je werkwijze.

VERBODEN OPENINGEN:
- "Hier is de uitwerking"
- "Absoluut"
- "Onderstaand"
- "In dit document"
- "Hieronder volgt"

INHOUDELIJKE REGELS:
1. Integratie:
   - Gebruik de slide-volgorde als ruggengraat.
   - Verwerk audio-uitleg op de logisch juiste plaats.
   - Behoud inhoudelijke details die didactische waarde hebben.
2. Redactie:
   - Verwijder conversatie-ruis (opstartzinnen, klasinteractie, herhalingen zonder inhoud).
   - Zet spreektaal en klasdialoog om naar vloeiende, doorlopende leertekst.
3. Structuur:
   - Per onderwerp: korte definitie/afbakening -> uitleg/mechanisme -> klinische of praktische relevantie.
   - Gebruik bullets alleen waar dat de scanbaarheid verbetert.
   - Behoud casussen/opdrachten als aparte secties als ze in de input staan.
4. Visual placeholders:
   - Behoud alleen `[Informatieve Afbeelding/Tabel: ...]` op de juiste plek.
   - Laat decoratieve placeholders weg.
5. Taal:
   - Schrijf volledig in: {output_language}.
   - Houd toon professioneel, neutraal en didactisch.

SOFT FIDELITY (BALANS TUSSEN BETROUWBAARHEID EN LEESBAARHEID):
- Baseren op slide-tekst + transcript als primaire waarheid.
- Toegestaan:
  - Korte verbindingszinnen voor leesbaarheid.
  - Voorzichtige herformulering/duiding van impliciete verbanden die direct uit de input volgen.
- Niet toegestaan:
  - Nieuwe cijfers, richtlijnen, bronnen, diagnoses of behandelclaims die niet uit de input komen.
  - Nieuwe medische feiten die niet in slide of transcript te herleiden zijn.
- Bij twijfel: laat weg of formuleer neutraal zonder extra claim.

AFSLUITING (VERPLICHT):
- Voeg een laatste sectie toe: `## Kernpunten voor tentamen`.
- Geef 8-15 concrete bullets met de belangrijkste leerpunten.

EINDCONTROLE VOOR UITVOER:
- Staat er nog een meta-inleiding? Verwijderen.
- Staat er nog letterlijke klasdialoog? Herschrijven.
- Zijn de hoofdonderwerpen uit zowel slides als transcript afgedekt? Zo niet: aanvullen.

INPUT SLIDE-TEKST:
{slide_text}

INPUT AUDIO-TRANSCRIPT:
{transcript}"""

PROMPT_MERGE_WITH_AUDIO_MARKERS = """Creëer een volledige, integrale en goed leesbare uitwerking van een college door slide-tekst en audio-transcript te combineren.

BELANGRIJK - AUDIO MARKERS:
Voor elke hoofdsectie gebruik je direct onder de kop exact dit marker-formaat:
<!-- audio:START_MS-END_MS -->
waar START_MS en END_MS de relevante tijdrange uit het transcript met tijdsegmenten aangeven.

Regels:
1. Niet samenvatten, maar compleet uitschrijven.
2. Gebruik koppen en subkoppen voor structuur.
3. Verwijder alleen irrelevante spreektaal; behoud alle inhoudelijke uitleg.
4. Gebruik geen labels zoals "Audio:" of "Slide:".
5. Schrijf volledig in deze taal: {output_language}.

Input slide-tekst:
{slide_text}

Input audio-transcript met tijdsegmenten:
{transcript}"""

PROMPT_STUDY_TEMPLATE = """You are an expert university professor creating study materials. I will provide you with the complete text of a lecture or slide deck.

Your task is to generate {flashcard_amount} flashcards and {question_amount} multiple-choice test questions based strictly on the provided text. Do not invent outside information.
Write all generated output fully in this language: {output_language}.

RULES FOR FLASHCARDS:
- The 'front' should be a clear term or concept.
- The 'back' should be a concise, accurate definition/explanation.

RULES FOR TEST QUESTIONS:
- Create challenging, university-level multiple-choice questions.
- Provide exactly 4 options (A, B, C, D) as an array of strings.
- Provide the correct answer (must match one option exactly).
- Provide a brief 'explanation' of WHY the answer is correct.

REQUIRED OUTPUT FORMAT:
You must respond with strictly valid JSON matching this structure:
{{
  "flashcards": [{{"front": "string", "back": "string"}}],
  "test_questions": [{{"question": "string", "options": ["string", "string", "string", "string"], "answer": "string", "explanation": "string"}}]
}}

LECTURE TEXT:
{source_text}
"""

# =============================================
# FIRESTORE USER FUNCTIONS
# =============================================

def get_or_create_user(uid, email):
    """Get a user from Firestore, or create them with free credits if they don't exist."""
    user_ref = db.collection('users').document(uid)
    user_doc = user_ref.get()

    if user_doc.exists:
        user_data = user_doc.to_dict()
        # Update email if it changed
        if user_data.get('email') != email and email:
            user_ref.update({'email': email})
            user_data['email'] = email
        return user_data
    else:
        # New user — create with free credits
        user_data = {
            'uid': uid,
            'email': email,
            'lecture_credits_standard': FREE_LECTURE_CREDITS,
            'lecture_credits_extended': 0,
            'slides_credits': FREE_SLIDES_CREDITS,
            'interview_credits_short': FREE_INTERVIEW_CREDITS,
            'interview_credits_medium': 0,
            'interview_credits_long': 0,
            'total_processed': 0,
            'created_at': time.time(),
        }
        user_ref.set(user_data)
        print(f"New user created: {uid} ({email})")
        return user_data

def grant_credits_to_user(uid, bundle_id):
    """Grant credits from a purchased bundle to a user in Firestore."""
    bundle = CREDIT_BUNDLES.get(bundle_id)
    if not bundle:
        print(f"Warning: Unknown bundle_id '{bundle_id}' in grant_credits_to_user")
        return False

    user_ref = db.collection('users').document(uid)
    user_doc = user_ref.get()

    if not user_doc.exists:
        # User not in Firestore yet — create with defaults first
        user_data = {
            'uid': uid,
            'email': '',
            'lecture_credits_standard': FREE_LECTURE_CREDITS,
            'lecture_credits_extended': 0,
            'slides_credits': FREE_SLIDES_CREDITS,
            'interview_credits_short': FREE_INTERVIEW_CREDITS,
            'interview_credits_medium': 0,
            'interview_credits_long': 0,
            'total_processed': 0,
            'created_at': time.time(),
        }
        user_ref.set(user_data)

    # Add the purchased credits
    for credit_key, credit_amount in bundle['credits'].items():
        user_ref.update({credit_key: firestore.Increment(credit_amount)})
        print(f"Granted {credit_amount} '{credit_key}' credits to user {uid}.")

    return True

def deduct_credit(uid, credit_type_primary, credit_type_fallback=None):
    """Deduct one credit from the user. Returns the credit type that was deducted, or None if no credits."""
    user_ref = db.collection('users').document(uid)
    user_doc = user_ref.get()

    if not user_doc.exists:
        return None

    user_data = user_doc.to_dict()

    if user_data.get(credit_type_primary, 0) > 0:
        user_ref.update({
            credit_type_primary: firestore.Increment(-1),
            'total_processed': firestore.Increment(1),
        })
        return credit_type_primary
    elif credit_type_fallback and user_data.get(credit_type_fallback, 0) > 0:
        user_ref.update({
            credit_type_fallback: firestore.Increment(-1),
            'total_processed': firestore.Increment(1),
        })
        return credit_type_fallback
    else:
        return None

def deduct_interview_credit(uid):
    """Deduct one interview credit, checking short -> medium -> long. Returns the credit type deducted, or None."""
    user_ref = db.collection('users').document(uid)
    user_doc = user_ref.get()

    if not user_doc.exists:
        return None

    user_data = user_doc.to_dict()

    if user_data.get('interview_credits_short', 0) > 0:
        user_ref.update({
            'interview_credits_short': firestore.Increment(-1),
            'total_processed': firestore.Increment(1),
        })
        return 'interview_credits_short'
    elif user_data.get('interview_credits_medium', 0) > 0:
        user_ref.update({
            'interview_credits_medium': firestore.Increment(-1),
            'total_processed': firestore.Increment(1),
        })
        return 'interview_credits_medium'
    elif user_data.get('interview_credits_long', 0) > 0:
        user_ref.update({
            'interview_credits_long': firestore.Increment(-1),
            'total_processed': firestore.Increment(1),
        })
        return 'interview_credits_long'
    else:
        return None

def refund_credit(uid, credit_type):
    """Refund one credit back to the user after a failed processing job."""
    if not uid or not credit_type:
        return
    try:
        user_ref = db.collection('users').document(uid)
        user_ref.update({
            credit_type: firestore.Increment(1),
            'total_processed': firestore.Increment(-1),
        })
        print(f"✅ Refunded 1 '{credit_type}' credit to user {uid} due to processing failure.")
    except Exception as e:
        print(f"❌ Failed to refund credit '{credit_type}' to user {uid}: {e}")

def save_purchase_record(uid, bundle_id, stripe_session_id):
    """Save a purchase record to Firestore for purchase history."""
    bundle = CREDIT_BUNDLES.get(bundle_id)
    if not bundle:
        return
    try:
        db.collection('purchases').add({
            'uid': uid,
            'bundle_id': bundle_id,
            'bundle_name': bundle['name'],
            'price_cents': bundle['price_cents'],
            'currency': bundle['currency'],
            'credits': bundle['credits'],
            'stripe_session_id': stripe_session_id,
            'created_at': time.time(),
        })
        print(f"📝 Saved purchase record for user {uid}: {bundle['name']}")
    except Exception as e:
        print(f"❌ Failed to save purchase record for user {uid}: {e}")

def save_job_log(job_id, job_data, finished_at):
    """Save a processing job log to Firestore for analytics."""
    try:
        started_at = job_data.get('started_at', 0)
        duration = round(finished_at - started_at, 1) if started_at else 0
        db.collection('job_logs').document(job_id).set({
            'job_id': job_id,
            'uid': job_data.get('user_id', ''),
            'email': job_data.get('user_email', ''),
            'mode': job_data.get('mode', ''),
            'status': job_data.get('status', ''),
            'credit_deducted': job_data.get('credit_deducted', ''),
            'credit_refunded': job_data.get('credit_refunded', False),
            'error_message': job_data.get('error', ''),
            'started_at': started_at,
            'finished_at': finished_at,
            'duration_seconds': duration,
        })
        print(f"📊 Logged job {job_id}: mode={job_data.get('mode')}, status={job_data.get('status')}, duration={duration}s")
    except Exception as e:
        print(f"❌ Failed to log job {job_id}: {e}")

# =============================================
# HELPER FUNCTIONS
# =============================================

def verify_firebase_token(request):
    auth_header = request.headers.get('Authorization', '')
    if not auth_header.startswith('Bearer '): return None
    token = auth_header.split('Bearer ')[1]
    try:
        return auth.verify_id_token(token)
    except Exception as e:
        print(f"Token verification failed: {e}")
        return None

def is_admin_user(decoded_token):
    if not decoded_token:
        return False
    uid = decoded_token.get('uid', '')
    email = decoded_token.get('email', '').lower()
    return uid in ADMIN_UIDS or email in ADMIN_EMAILS

def get_admin_window(window_key):
    windows = {
        '24h': 24 * 60 * 60,
        '7d': 7 * 24 * 60 * 60,
        '30d': 30 * 24 * 60 * 60,
    }
    safe_key = window_key if window_key in windows else '7d'
    return safe_key, windows[safe_key]

def get_timestamp(value):
    return value if isinstance(value, (int, float)) else 0

def build_time_buckets(window_key, now_ts):
    labels = []
    keys = []
    if window_key == '24h':
        now_dt = datetime.utcfromtimestamp(now_ts).replace(minute=0, second=0, microsecond=0)
        start_dt = now_dt - timedelta(hours=23)
        for i in range(24):
            current = start_dt + timedelta(hours=i)
            labels.append(current.strftime('%H:%M'))
            keys.append(current.strftime('%Y-%m-%d %H:00'))
        granularity = 'hour'
    else:
        days = 7 if window_key == '7d' else 30
        now_dt = datetime.utcfromtimestamp(now_ts).replace(hour=0, minute=0, second=0, microsecond=0)
        start_dt = now_dt - timedelta(days=days - 1)
        for i in range(days):
            current = start_dt + timedelta(days=i)
            labels.append(current.strftime('%d %b'))
            keys.append(current.strftime('%Y-%m-%d'))
        granularity = 'day'
    return labels, keys, granularity

def get_bucket_key(timestamp, window_key):
    dt = datetime.utcfromtimestamp(timestamp)
    if window_key == '24h':
        return dt.replace(minute=0, second=0, microsecond=0).strftime('%Y-%m-%d %H:00')
    return dt.replace(hour=0, minute=0, second=0, microsecond=0).strftime('%Y-%m-%d')

def parse_requested_amount(raw_value, allowed, default):
    value = str(raw_value or default).strip().lower()
    return value if value in allowed else default

def parse_study_features(raw_value):
    value = str(raw_value or 'none').strip().lower()
    return value if value in {'none', 'flashcards', 'test', 'both'} else 'none'

def parse_output_language(raw_value, custom_value=''):
    key = str(raw_value or 'english').strip().lower()
    if key in OUTPUT_LANGUAGE_MAP:
        return OUTPUT_LANGUAGE_MAP[key]
    if key == 'other':
        custom = str(custom_value or '').strip()
        if custom:
            return custom[:40]
    return 'English'

def parse_interview_features(raw_value):
    value = str(raw_value or 'none').strip().lower()
    if value in {'none', ''}:
        return []
    if value == 'both':
        return ['summary', 'sections']
    parts = [part.strip() for part in value.split(',') if part.strip()]
    features = []
    for part in parts:
        if part in {'summary', 'sections'} and part not in features:
            features.append(part)
    return features

def deduct_slides_credits(uid, amount):
    if amount <= 0:
        return True
    user_ref = db.collection('users').document(uid)
    user_doc = user_ref.get()
    if not user_doc.exists:
        return False
    user_data = user_doc.to_dict()
    current = user_data.get('slides_credits', 0)
    if current < amount:
        return False
    user_ref.update({'slides_credits': firestore.Increment(-amount)})
    return True

def refund_slides_credits(uid, amount):
    if not uid or amount <= 0:
        return
    try:
        db.collection('users').document(uid).update({'slides_credits': firestore.Increment(amount)})
        print(f"✅ Refunded {amount} slides credits to user {uid}.")
    except Exception as e:
        print(f"❌ Failed to refund {amount} slides credits to user {uid}: {e}")

def generate_with_optional_thinking(model, prompt_text, max_output_tokens=65536, thinking_budget=256):
    base_config = {'max_output_tokens': max_output_tokens}
    try:
        if hasattr(types, 'ThinkingConfig'):
            base_config['thinking_config'] = types.ThinkingConfig(thinking_budget=thinking_budget)
    except Exception:
        pass
    try:
        config = types.GenerateContentConfig(**base_config)
    except Exception:
        config = types.GenerateContentConfig(max_output_tokens=max_output_tokens)
    return client.models.generate_content(
        model=model,
        contents=[types.Content(role='user', parts=[types.Part.from_text(text=prompt_text)])],
        config=config
    )

def convert_audio_to_mp3_with_ytdlp(local_audio_path):
    if not local_audio_path or local_audio_path.lower().endswith('.mp3'):
        return local_audio_path, False
    ytdlp_bin = shutil.which('yt-dlp')
    base_no_ext = os.path.splitext(local_audio_path)[0]
    output_path = f"{base_no_ext}_converted.mp3"

    if ytdlp_bin:
        try:
            source = f"file://{os.path.abspath(local_audio_path)}"
            command = [
                ytdlp_bin,
                '--no-playlist',
                '--extract-audio',
                '--audio-format', 'mp3',
                '--audio-quality', '5',
                '--output', f"{base_no_ext}_converted.%(ext)s",
                source
            ]
            result = subprocess.run(command, check=False, capture_output=True, text=True)
            if result.returncode == 0 and os.path.exists(output_path):
                return output_path, True
            print(f"⚠️ yt-dlp conversion failed: {(result.stderr or '').strip()[:300]}")
        except Exception as e:
            print(f"⚠️ yt-dlp conversion exception: {e}")
    else:
        print("⚠️ yt-dlp not found, skipping yt-dlp conversion.")

    # Fallback conversion for local files when yt-dlp is unavailable or fails.
    ffmpeg_bin = shutil.which('ffmpeg')
    if not ffmpeg_bin:
        return local_audio_path, False
    try:
        command = [
            ffmpeg_bin,
            '-y',
            '-i', local_audio_path,
            '-vn',
            '-codec:a', 'libmp3lame',
            '-q:a', '5',
            output_path
        ]
        result = subprocess.run(command, check=False, capture_output=True, text=True)
        if result.returncode == 0 and os.path.exists(output_path):
            return output_path, True
        print(f"⚠️ ffmpeg conversion failed: {(result.stderr or '').strip()[:300]}")
    except Exception as e:
        print(f"⚠️ ffmpeg conversion exception: {e}")
    return local_audio_path, False

def resolve_auto_amount(kind, source_text):
    word_count = len((source_text or '').split())
    if kind == 'flashcards':
        if word_count < 1200:
            return 10
        if word_count < 2600:
            return 20
        return 30
    if word_count < 1200:
        return 5
    if word_count < 2600:
        return 10
    return 15

def resolve_study_amounts(flashcard_selection, question_selection, source_text):
    flashcard_amount = resolve_auto_amount('flashcards', source_text) if flashcard_selection == 'auto' else int(flashcard_selection)
    question_amount = resolve_auto_amount('questions', source_text) if question_selection == 'auto' else int(question_selection)
    return flashcard_amount, question_amount

def extract_json_payload(raw_text):
    if not raw_text:
        return None
    text = raw_text.strip()
    if text.startswith('```'):
        lines = text.splitlines()
        if len(lines) >= 3 and lines[0].startswith('```') and lines[-1].strip() == '```':
            text = '\n'.join(lines[1:-1]).strip()
    start = text.find('{')
    if start == -1:
        return None
    decoder = json.JSONDecoder()
    try:
        parsed, _ = decoder.raw_decode(text[start:])
        return parsed
    except json.JSONDecodeError:
        end = text.rfind('}')
        if end == -1 or end <= start:
            return None
        try:
            return json.loads(text[start:end + 1])
        except json.JSONDecodeError:
            return None

def sanitize_flashcards(items, max_items):
    if not isinstance(items, list):
        return []
    cleaned = []
    seen = set()
    for item in items:
        if not isinstance(item, dict):
            continue
        front = str(item.get('front', '')).strip()
        back = str(item.get('back', '')).strip()
        if not front or not back:
            continue
        key = (front.lower(), back.lower())
        if key in seen:
            continue
        seen.add(key)
        cleaned.append({'front': front, 'back': back})
        if len(cleaned) >= max_items:
            break
    return cleaned

def sanitize_questions(items, max_items):
    if not isinstance(items, list):
        return []
    cleaned = []
    seen = set()
    for item in items:
        if not isinstance(item, dict):
            continue
        question = str(item.get('question', '')).strip()
        options = item.get('options', [])
        answer = str(item.get('answer', '')).strip()
        explanation = str(item.get('explanation', '')).strip()
        if not question or not isinstance(options, list) or len(options) != 4 or not answer:
            continue
        option_strings = [str(option).strip() for option in options]
        if any(not option for option in option_strings):
            continue
        if len(set(option_strings)) != 4:
            continue
        if answer not in option_strings:
            continue
        dedupe_key = question.lower()
        if dedupe_key in seen:
            continue
        seen.add(dedupe_key)
        cleaned.append({
            'question': question,
            'options': option_strings,
            'answer': answer,
            'explanation': explanation,
        })
        if len(cleaned) >= max_items:
            break
    return cleaned

def generate_study_materials(source_text, flashcard_selection, question_selection, study_features='both', output_language='English'):
    if study_features == 'none':
        return [], [], None
    flashcard_amount, question_amount = resolve_study_amounts(flashcard_selection, question_selection, source_text)
    if study_features == 'flashcards':
        question_amount = 0
    elif study_features == 'test':
        flashcard_amount = 0
    prompt = PROMPT_STUDY_TEMPLATE.format(
        flashcard_amount=flashcard_amount,
        question_amount=question_amount,
        output_language=output_language,
        source_text=source_text[:120000],
    )
    try:
        response = client.models.generate_content(
            model=MODEL_STUDY,
            contents=[types.Content(role='user', parts=[types.Part.from_text(text=prompt)])],
            config=types.GenerateContentConfig(max_output_tokens=32768)
        )
        parsed = extract_json_payload(response.text)
        if not isinstance(parsed, dict):
            return [], [], 'Study materials JSON parsing failed.'
        flashcards = sanitize_flashcards(parsed.get('flashcards', []), flashcard_amount)
        test_questions = sanitize_questions(parsed.get('test_questions', []), question_amount)
        if not flashcards and not test_questions and study_features != 'none':
            return [], [], 'Study materials were empty after validation.'
        return flashcards, test_questions, None
    except Exception as e:
        return [], [], f'Study materials generation failed: {e}'

def generate_interview_enhancements(transcript_text, selected_features, output_language='English'):
    summary_text = None
    sectioned_text = None
    errors = []
    for feature in selected_features:
        try:
            if feature == 'summary':
                prompt = PROMPT_INTERVIEW_SUMMARY.format(transcript=transcript_text[:120000], output_language=output_language)
                response = generate_with_optional_thinking(MODEL_STUDY, prompt, max_output_tokens=8192, thinking_budget=384)
                summary_text = (response.text or '').strip()
                if not summary_text:
                    errors.append('Summary generation returned empty output.')
            elif feature == 'sections':
                prompt = PROMPT_INTERVIEW_SECTIONED.format(transcript=transcript_text[:120000], output_language=output_language)
                response = generate_with_optional_thinking(MODEL_STUDY, prompt, max_output_tokens=32768, thinking_budget=384)
                sectioned_text = (response.text or '').strip()
                if not sectioned_text:
                    errors.append('Sectioned transcript generation returned empty output.')
        except Exception as e:
            errors.append(f"{feature} generation failed: {e}")

    successful = []
    if summary_text:
        successful.append('summary')
    if sectioned_text:
        successful.append('sections')

    combined_text = None
    if summary_text and sectioned_text:
        combined_text = f"# Interview Summary\n\n{summary_text}\n\n# Structured Interview Transcript\n\n{sectioned_text}"

    failed_count = max(0, len(selected_features) - len(successful))
    return {
        'summary': summary_text,
        'sections': sectioned_text,
        'combined': combined_text,
        'successful_features': successful,
        'failed_count': failed_count,
        'error': '; '.join(errors) if errors else None,
    }

def allowed_file(filename, allowed_extensions):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions

def get_mime_type(filename):
    ext = filename.rsplit('.', 1)[1].lower()
    mime_types = {'pdf': 'application/pdf', 'mp3': 'audio/mpeg', 'm4a': 'audio/mp4', 'wav': 'audio/wav', 'aac': 'audio/aac', 'ogg': 'audio/ogg', 'flac': 'audio/flac'}
    return mime_types.get(ext, 'application/octet-stream')

def wait_for_file_processing(uploaded_file):
    max_wait_time = 300
    wait_interval = 5
    total_waited = 0
    while total_waited < max_wait_time:
        file_info = client.files.get(name=uploaded_file.name)
        if file_info.state.name == 'ACTIVE': return True
        elif file_info.state.name == 'FAILED': raise Exception(f"File processing failed: {uploaded_file.name}")
        time.sleep(wait_interval)
        total_waited += wait_interval
    raise Exception(f"File processing timed out after {max_wait_time} seconds")

def cleanup_files(local_paths, gemini_files):
    for path in local_paths:
        try:
            if os.path.exists(path): os.remove(path)
        except Exception as e: print(f"Warning: Could not delete local file {path}: {e}")
    for gemini_file in gemini_files:
        try:
            client.files.delete(name=gemini_file.name)
        except Exception as e: print(f"Warning: Could not delete Gemini file {gemini_file.name}: {e}")

def parse_audio_markers_from_notes(notes_markdown):
    if not notes_markdown:
        return []
    pattern = re.compile(r'#{1,3}\s+(.+?)\s*\n\s*<!--\s*audio:(\d+)-(\d+)\s*-->', re.MULTILINE)
    notes_audio_map = []
    section_index = 0
    for match in pattern.finditer(notes_markdown):
        try:
            start_ms = int(match.group(2))
            end_ms = int(match.group(3))
        except Exception:
            continue
        notes_audio_map.append({
            'section_index': section_index,
            'section_title': match.group(1).strip(),
            'start_ms': max(0, start_ms),
            'end_ms': max(start_ms, end_ms),
        })
        section_index += 1
    return notes_audio_map

def format_transcript_with_timestamps(segments):
    if not isinstance(segments, list):
        return ''
    lines = []
    for seg in segments:
        if not isinstance(seg, dict):
            continue
        text = str(seg.get('text', '') or '').strip()
        if not text:
            continue
        start_ms = int(seg.get('start_ms', 0) or 0)
        end_ms = int(seg.get('end_ms', start_ms) or start_ms)
        lines.append(f"[{start_ms}-{end_ms}] {text}")
    return '\n'.join(lines)

def transcribe_audio_plain(audio_file, audio_mime_type, output_language='English'):
    output_language = OUTPUT_LANGUAGE_MAP.get(str(output_language).lower(), str(output_language))
    prompt = PROMPT_AUDIO_TRANSCRIPTION.format(output_language=output_language)
    response = client.models.generate_content(
        model=MODEL_AUDIO,
        contents=[types.Content(role='user', parts=[
            types.Part.from_uri(file_uri=audio_file.uri, mime_type=audio_mime_type),
            types.Part.from_text(text=prompt)
        ])],
        config=types.GenerateContentConfig(max_output_tokens=65536)
    )
    return (getattr(response, 'text', '') or '').strip()

def transcribe_audio_with_timestamps(audio_file, audio_mime_type, output_language='English'):
    output_language = OUTPUT_LANGUAGE_MAP.get(str(output_language).lower(), str(output_language))
    prompt = PROMPT_AUDIO_TRANSCRIPTION_TIMESTAMPED.format(output_language=output_language)
    try:
        response = client.models.generate_content(
            model=MODEL_AUDIO,
            contents=[types.Content(role='user', parts=[
                types.Part.from_uri(file_uri=audio_file.uri, mime_type=audio_mime_type),
                types.Part.from_text(text=prompt)
            ])],
            config=types.GenerateContentConfig(max_output_tokens=65536)
        )
        parsed = extract_json_payload(getattr(response, 'text', '') or '')
        if not isinstance(parsed, dict):
            raise ValueError('Timestamped transcription JSON not found')
        raw_segments = parsed.get('transcript_segments', [])
        full_transcript = str(parsed.get('full_transcript', '') or '').strip()
        clean_segments = []
        if isinstance(raw_segments, list):
            for seg in raw_segments:
                if not isinstance(seg, dict):
                    continue
                text = str(seg.get('text', '') or '').strip()
                if not text:
                    continue
                try:
                    start_ms = int(seg.get('start_ms', 0) or 0)
                    end_ms = int(seg.get('end_ms', start_ms) or start_ms)
                except Exception:
                    continue
                clean_segments.append({
                    'start_ms': max(0, start_ms),
                    'end_ms': max(start_ms, end_ms),
                    'text': text,
                })
        if not full_transcript and clean_segments:
            full_transcript = '\n'.join([s['text'] for s in clean_segments]).strip()
        if not full_transcript:
            raise ValueError('Empty transcript')
        return full_transcript, clean_segments
    except Exception as e:
        print(f"⚠️ Timestamp transcription failed, falling back to plain transcript: {e}")
        fallback_prompt = PROMPT_AUDIO_TRANSCRIPTION.format(output_language=output_language)
        fallback_response = client.models.generate_content(
            model=MODEL_AUDIO,
            contents=[types.Content(role='user', parts=[
                types.Part.from_uri(file_uri=audio_file.uri, mime_type=audio_mime_type),
                types.Part.from_text(text=fallback_prompt)
            ])],
            config=types.GenerateContentConfig(max_output_tokens=65536)
        )
        return (getattr(fallback_response, 'text', '') or '').strip(), []

def persist_audio_for_study_pack(job_id, audio_source_path):
    if not audio_source_path or not os.path.exists(audio_source_path):
        return ''
    ext = os.path.splitext(audio_source_path)[1].lower() or '.mp3'
    audio_dir = os.path.join(UPLOAD_FOLDER, 'study_audio')
    os.makedirs(audio_dir, exist_ok=True)
    target_path = os.path.join(audio_dir, f"{job_id}{ext}")
    try:
        shutil.copy2(audio_source_path, target_path)
        return target_path
    except Exception as e:
        print(f"⚠️ Could not persist audio for study pack {job_id}: {e}")
        return ''

def markdown_to_docx(markdown_text, title="Document"):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    lines = markdown_text.split('\n')
    i = 0
    is_transcript = any(len(line.strip()) > 3 and line.strip()[0].isdigit() and ':' in line.strip()[:6] and ' - ' in line for line in lines[:20])

    def add_inline_markdown_runs(paragraph, text):
        raw = str(text or '')
        # Supports inline emphasis markers used in generated output.
        parts = re.split(r'(\*\*.+?\*\*|__.+?__|\*.+?\*|_.+?_)', raw)
        for part in parts:
            if not part:
                continue
            if (part.startswith('**') and part.endswith('**') and len(part) >= 4) or (part.startswith('__') and part.endswith('__') and len(part) >= 4):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
                continue
            if (part.startswith('*') and part.endswith('*') and len(part) >= 3) or (part.startswith('_') and part.endswith('_') and len(part) >= 3):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
                continue
            paragraph.add_run(part.replace('**', '').replace('__', ''))
    
    while i < len(lines):
        line = lines[i].strip()
        numbered_match = re.match(r'^\d+\.\s+(.*)$', line)
        if not line:
            i += 1
            continue
        if line.startswith('### '): doc.add_heading(line[4:], level=3)
        elif line.startswith('## '): doc.add_heading(line[3:], level=2)
        elif line.startswith('# '): doc.add_heading(line[2:], level=1)
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            add_inline_markdown_runs(p, line[2:])
        elif numbered_match:
            p = doc.add_paragraph(style='List Number')
            add_inline_markdown_runs(p, numbered_match.group(1))
        elif is_transcript and len(line) > 3 and line[0].isdigit() and ':' in line[:6]:
            p = doc.add_paragraph()
            add_inline_markdown_runs(p, line)
        else:
            paragraph_lines = [line]
            while i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                if (
                    next_line
                    and not next_line.startswith('#')
                    and not next_line.startswith('- ')
                    and not next_line.startswith('* ')
                    and not re.match(r'^\d+\.\s+', next_line)
                ):
                    paragraph_lines.append(next_line)
                    i += 1
                else:
                    break
            paragraph_text = ' '.join(paragraph_lines)
            p = doc.add_paragraph()
            add_inline_markdown_runs(p, paragraph_text)
        i += 1
    return doc

def normalize_exam_date(raw_value):
    exam_date = str(raw_value or '').strip()
    if not exam_date:
        return ''
    try:
        return datetime.strptime(exam_date, '%Y-%m-%d').strftime('%Y-%m-%d')
    except ValueError:
        raise ValueError('Exam date must use YYYY-MM-DD format')

def markdown_inline_to_pdf_html(text):
    safe_text = html.escape(str(text or ''))
    safe_text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', safe_text)
    safe_text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', safe_text)
    return safe_text

def append_notes_markdown_to_story(story, notes_markdown, styles):
    lines = str(notes_markdown or '').splitlines()
    bullet_items = []

    def flush_bullets():
        nonlocal bullet_items
        if not bullet_items:
            return
        list_flow = ListFlowable(
            [ListItem(Paragraph(item, styles['pdfBody']), leftIndent=6) for item in bullet_items],
            bulletType='bullet',
            leftIndent=14,
            bulletFontSize=8,
            bulletOffsetY=1
        )
        story.append(list_flow)
        story.append(Spacer(1, 4))
        bullet_items = []

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            flush_bullets()
            story.append(Spacer(1, 4))
            continue

        heading_level = 0
        if line.startswith('### '):
            heading_level = 3
        elif line.startswith('## '):
            heading_level = 2
        elif line.startswith('# '):
            heading_level = 1

        if heading_level:
            flush_bullets()
            heading_text = markdown_inline_to_pdf_html(line[heading_level + 1:])
            heading_style = styles['pdfH1'] if heading_level == 1 else styles['pdfH2'] if heading_level == 2 else styles['pdfH3']
            story.append(Paragraph(heading_text, heading_style))
            story.append(Spacer(1, 3))
            continue

        if line.startswith('- ') or line.startswith('* '):
            bullet_items.append(markdown_inline_to_pdf_html(line[2:].strip()))
            continue

        numbered_match = re.match(r'^(\d+)\.\s+(.*)$', line)
        if numbered_match:
            flush_bullets()
            text_html = markdown_inline_to_pdf_html(numbered_match.group(2))
            story.append(Paragraph(f"{numbered_match.group(1)}. {text_html}", styles['pdfBody']))
            story.append(Spacer(1, 2))
            continue

        flush_bullets()
        story.append(Paragraph(markdown_inline_to_pdf_html(line), styles['pdfBody']))
        story.append(Spacer(1, 2))

    flush_bullets()

def build_study_pack_pdf(pack, include_answers=True):
    if not REPORTLAB_AVAILABLE:
        raise RuntimeError(
            "PDF export requires the optional 'reportlab' dependency. "
            "Install it with: pip install reportlab==4.2.5"
        )

    pdf_buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        pdf_buffer,
        pagesize=A4,
        leftMargin=16 * mm,
        rightMargin=16 * mm,
        topMargin=14 * mm,
        bottomMargin=14 * mm,
        title=str(pack.get('title', 'Study Pack')).strip() or 'Study Pack'
    )

    base_styles = getSampleStyleSheet()
    styles = {
        'pdfTitle': ParagraphStyle(
            'PdfTitle',
            parent=base_styles['Heading1'],
            fontName='Helvetica-Bold',
            fontSize=17,
            leading=21,
            spaceAfter=6,
            textColor=colors.HexColor('#111827')
        ),
        'pdfMeta': ParagraphStyle(
            'PdfMeta',
            parent=base_styles['BodyText'],
            fontName='Helvetica',
            fontSize=9.5,
            leading=12.5,
            textColor=colors.HexColor('#4B5563')
        ),
        'pdfSection': ParagraphStyle(
            'PdfSection',
            parent=base_styles['Heading2'],
            fontName='Helvetica-Bold',
            fontSize=12.5,
            leading=16,
            spaceBefore=6,
            spaceAfter=6,
            textColor=colors.HexColor('#111827')
        ),
        'pdfH1': ParagraphStyle(
            'PdfH1',
            parent=base_styles['Heading2'],
            fontName='Helvetica-Bold',
            fontSize=12,
            leading=15,
            textColor=colors.HexColor('#1F2937')
        ),
        'pdfH2': ParagraphStyle(
            'PdfH2',
            parent=base_styles['Heading3'],
            fontName='Helvetica-Bold',
            fontSize=11,
            leading=14,
            textColor=colors.HexColor('#1F2937')
        ),
        'pdfH3': ParagraphStyle(
            'PdfH3',
            parent=base_styles['Heading4'],
            fontName='Helvetica-Bold',
            fontSize=10,
            leading=13,
            textColor=colors.HexColor('#374151')
        ),
        'pdfBody': ParagraphStyle(
            'PdfBody',
            parent=base_styles['BodyText'],
            fontName='Helvetica',
            fontSize=9.5,
            leading=13,
            textColor=colors.HexColor('#111827')
        ),
        'pdfQuestion': ParagraphStyle(
            'PdfQuestion',
            parent=base_styles['BodyText'],
            fontName='Helvetica-Bold',
            fontSize=10,
            leading=13.5,
            textColor=colors.HexColor('#111827')
        ),
        'pdfOption': ParagraphStyle(
            'PdfOption',
            parent=base_styles['BodyText'],
            fontName='Helvetica',
            fontSize=9.5,
            leading=12.5,
            leftIndent=10,
            textColor=colors.HexColor('#1F2937')
        ),
        'pdfOptionCorrect': ParagraphStyle(
            'PdfOptionCorrect',
            parent=base_styles['BodyText'],
            fontName='Helvetica-Bold',
            fontSize=9.5,
            leading=12.5,
            leftIndent=10,
            textColor=colors.HexColor('#065F46')
        ),
    }

    pack_title = str(pack.get('title', 'Study Pack')).strip() or 'Study Pack'
    story = [Paragraph(markdown_inline_to_pdf_html(pack_title), styles['pdfTitle'])]

    mode = str(pack.get('mode', '') or '').strip() or 'Unknown'
    output_language = str(pack.get('output_language', '') or '').strip() or 'Unknown'
    course = str(pack.get('course', '') or '').strip() or '-'
    subject = str(pack.get('subject', '') or '').strip() or '-'
    semester = str(pack.get('semester', '') or '').strip() or '-'
    block = str(pack.get('block', '') or '').strip() or '-'
    created_at = pack.get('created_at', 0)
    created_text = '-'
    try:
        if created_at:
            created_text = datetime.fromtimestamp(float(created_at)).strftime('%Y-%m-%d %H:%M')
    except Exception:
        created_text = '-'

    metadata_rows = [
        [Paragraph('<b>Mode</b>', styles['pdfMeta']), Paragraph(markdown_inline_to_pdf_html(mode), styles['pdfMeta'])],
        [Paragraph('<b>Language</b>', styles['pdfMeta']), Paragraph(markdown_inline_to_pdf_html(output_language), styles['pdfMeta'])],
        [Paragraph('<b>Course</b>', styles['pdfMeta']), Paragraph(markdown_inline_to_pdf_html(course), styles['pdfMeta'])],
        [Paragraph('<b>Subject</b>', styles['pdfMeta']), Paragraph(markdown_inline_to_pdf_html(subject), styles['pdfMeta'])],
        [Paragraph('<b>Semester / Block</b>', styles['pdfMeta']), Paragraph(markdown_inline_to_pdf_html(f"{semester} / {block}"), styles['pdfMeta'])],
        [Paragraph('<b>Created</b>', styles['pdfMeta']), Paragraph(markdown_inline_to_pdf_html(created_text), styles['pdfMeta'])],
    ]
    metadata_table = Table(metadata_rows, colWidths=[36 * mm, 145 * mm], hAlign='LEFT')
    metadata_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('TOPPADDING', (0, 0), (-1, -1), 2),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
        ('GRID', (0, 0), (-1, -1), 0.35, colors.HexColor('#E5E7EB')),
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#F9FAFB')),
    ]))
    story.append(metadata_table)
    story.append(Spacer(1, 10))

    story.append(Paragraph('Integrated Notes', styles['pdfSection']))
    notes_markdown = str(pack.get('notes_markdown', '') or '').strip()
    if notes_markdown:
        append_notes_markdown_to_story(story, notes_markdown, styles)
    else:
        story.append(Paragraph('No integrated notes available.', styles['pdfBody']))
    story.append(Spacer(1, 10))

    story.append(Paragraph('Flashcards', styles['pdfSection']))
    flashcards = pack.get('flashcards', []) if isinstance(pack.get('flashcards', []), list) else []
    if flashcards:
        card_rows = [[Paragraph('<b>Front</b>', styles['pdfMeta']), Paragraph('<b>Back</b>', styles['pdfMeta'])]]
        for card in flashcards:
            card_rows.append([
                Paragraph(markdown_inline_to_pdf_html(str(card.get('front', '') or '')), styles['pdfBody']),
                Paragraph(markdown_inline_to_pdf_html(str(card.get('back', '') or '')), styles['pdfBody']),
            ])
        flashcard_table = Table(card_rows, colWidths=[84 * mm, 97 * mm], repeatRows=1, hAlign='LEFT')
        flashcard_table.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 0.35, colors.HexColor('#D1D5DB')),
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#F3F4F6')),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 5),
            ('RIGHTPADDING', (0, 0), (-1, -1), 5),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ]))
        story.append(flashcard_table)
    else:
        story.append(Paragraph('No flashcards available.', styles['pdfBody']))

    story.append(PageBreak())
    practice_title = 'Practice Questions'
    if not include_answers:
        practice_title += ' (Without Answers)'
    story.append(Paragraph(practice_title, styles['pdfSection']))
    questions = pack.get('test_questions', []) if isinstance(pack.get('test_questions', []), list) else []
    if questions:
        for idx, question in enumerate(questions, 1):
            question_text = str(question.get('question', '') or '').strip() or f'Question {idx}'
            story.append(Paragraph(f"{idx}. {markdown_inline_to_pdf_html(question_text)}", styles['pdfQuestion']))

            options = question.get('options', [])
            if not isinstance(options, list):
                options = []
            answer = str(question.get('answer', '') or '').strip()
            letters = ['A', 'B', 'C', 'D']
            for option_idx, option in enumerate(options[:4]):
                option_text = str(option or '').strip()
                is_correct = include_answers and option_text == answer and option_text != ''
                marker = '✓' if is_correct else '•'
                letter = letters[option_idx] if option_idx < len(letters) else str(option_idx + 1)
                option_style = styles['pdfOptionCorrect'] if is_correct else styles['pdfOption']
                story.append(Paragraph(f"{marker} {letter}. {markdown_inline_to_pdf_html(option_text)}", option_style))

            explanation = str(question.get('explanation', '') or '').strip()
            if include_answers and explanation:
                story.append(Paragraph(f"<b>Explanation:</b> {markdown_inline_to_pdf_html(explanation)}", styles['pdfBody']))
            story.append(Spacer(1, 7))
    else:
        story.append(Paragraph('No practice questions available.', styles['pdfBody']))

    doc.build(story)
    pdf_buffer.seek(0)
    return pdf_buffer

def save_study_pack(job_id, job_data):
    try:
        notes_markdown = str(job_data.get('result', '') or '')
        max_notes_chars = 180000
        notes_truncated = len(notes_markdown) > max_notes_chars
        if notes_truncated:
            notes_markdown = notes_markdown[:max_notes_chars]

        doc_ref = db.collection('study_packs').document()
        now_ts = time.time()
        doc_ref.set({
            'study_pack_id': doc_ref.id,
            'source_job_id': job_id,
            'uid': job_data.get('user_id', ''),
            'email': job_data.get('user_email', ''),
            'mode': job_data.get('mode', ''),
            'title': f"{job_data.get('mode', 'study-pack')} {datetime.utcfromtimestamp(now_ts).strftime('%Y-%m-%d %H:%M')}",
            'output_language': job_data.get('output_language', 'English'),
            'notes_markdown': notes_markdown,
            'notes_truncated': notes_truncated,
            'transcript_segments': job_data.get('transcript_segments', []),
            'notes_audio_map': job_data.get('notes_audio_map', []),
            'audio_storage_path': job_data.get('audio_storage_path', ''),
            'has_audio_sync': FEATURE_AUDIO_SECTION_SYNC and bool(job_data.get('audio_storage_path')) and bool(job_data.get('notes_audio_map', [])),
            'has_audio_playback': bool(job_data.get('audio_storage_path')),
            'flashcards': job_data.get('flashcards', []),
            'test_questions': job_data.get('test_questions', []),
            'flashcard_selection': job_data.get('flashcard_selection', '20'),
            'question_selection': job_data.get('question_selection', '10'),
            'study_features': job_data.get('study_features', 'none'),
            'interview_features': job_data.get('interview_features', []),
            'interview_summary': job_data.get('interview_summary'),
            'interview_sections': job_data.get('interview_sections'),
            'interview_combined': job_data.get('interview_combined'),
            'study_generation_error': job_data.get('study_generation_error'),
            'course': '',
            'subject': '',
            'semester': '',
            'block': '',
            'folder_id': '',
            'folder_name': '',
            'created_at': now_ts,
            'updated_at': now_ts,
        })
        job_data['study_pack_id'] = doc_ref.id
    except Exception as e:
        print(f"❌ Failed to save study pack for job {job_id}: {e}")

# =============================================
# AI PROCESSING FUNCTIONS
# =============================================

def process_lecture_notes(job_id, pdf_path, audio_path):
    gemini_files = []
    local_paths = [pdf_path, audio_path]
    try:
        jobs[job_id]['status'] = 'processing'
        jobs[job_id]['step'] = 1
        jobs[job_id]['step_description'] = 'Extracting text from slides...'
        pdf_file = client.files.upload(file=pdf_path, config={'mime_type': 'application/pdf'})
        gemini_files.append(pdf_file)
        wait_for_file_processing(pdf_file)
        response = client.models.generate_content(model=MODEL_SLIDES, contents=[types.Content(role='user', parts=[types.Part.from_uri(file_uri=pdf_file.uri, mime_type='application/pdf'), types.Part.from_text(text=PROMPT_SLIDE_EXTRACTION)])], config=types.GenerateContentConfig(max_output_tokens=65536))
        slide_text = response.text
        jobs[job_id]['slide_text'] = slide_text
        
        jobs[job_id]['step'] = 2
        jobs[job_id]['step_description'] = 'Transcribing audio...'
        output_language = jobs[job_id].get('output_language', 'English')
        converted_audio_path, converted = convert_audio_to_mp3_with_ytdlp(audio_path)
        if converted and converted_audio_path not in local_paths:
            local_paths.append(converted_audio_path)
        jobs[job_id]['step_description'] = 'Optimizing audio for faster processing...'
        audio_mime_type = get_mime_type(converted_audio_path)
        audio_file = client.files.upload(file=converted_audio_path, config={'mime_type': audio_mime_type})
        gemini_files.append(audio_file)
        jobs[job_id]['step_description'] = 'Processing audio file (this may take a few minutes)...'
        wait_for_file_processing(audio_file)
        jobs[job_id]['step_description'] = 'Generating transcript...'
        if FEATURE_AUDIO_SECTION_SYNC:
            transcript, transcript_segments = transcribe_audio_with_timestamps(audio_file, audio_mime_type, output_language)
        else:
            transcript = transcribe_audio_plain(audio_file, audio_mime_type, output_language)
            transcript_segments = []
        jobs[job_id]['transcript'] = transcript
        jobs[job_id]['transcript_segments'] = transcript_segments
        jobs[job_id]['audio_storage_path'] = persist_audio_for_study_pack(job_id, converted_audio_path)
        
        jobs[job_id]['step'] = 3
        jobs[job_id]['step_description'] = 'Creating complete lecture notes...'
        merge_transcript = format_transcript_with_timestamps(transcript_segments) if transcript_segments else transcript
        if FEATURE_AUDIO_SECTION_SYNC and transcript_segments:
            merge_prompt = PROMPT_MERGE_WITH_AUDIO_MARKERS.format(slide_text=slide_text, transcript=merge_transcript, output_language=output_language)
        else:
            merge_prompt = PROMPT_MERGE_TEMPLATE.format(slide_text=slide_text, transcript=transcript, output_language=output_language)
        response = client.models.generate_content(model=MODEL_INTEGRATION, contents=[types.Content(role='user', parts=[types.Part.from_text(text=merge_prompt)])], config=types.GenerateContentConfig(max_output_tokens=65536))
        merged_notes = response.text
        jobs[job_id]['result'] = merged_notes
        jobs[job_id]['notes_audio_map'] = parse_audio_markers_from_notes(merged_notes) if FEATURE_AUDIO_SECTION_SYNC else []

        if jobs[job_id].get('study_features', 'none') != 'none':
            jobs[job_id]['step'] = 4
            jobs[job_id]['step_description'] = 'Generating flashcards and practice test...'
            flashcards, test_questions, study_error = generate_study_materials(
                merged_notes,
                jobs[job_id].get('flashcard_selection', '20'),
                jobs[job_id].get('question_selection', '10'),
                jobs[job_id].get('study_features', 'none'),
                output_language
            )
            jobs[job_id]['flashcards'] = flashcards
            jobs[job_id]['test_questions'] = test_questions
            jobs[job_id]['study_generation_error'] = study_error
        else:
            jobs[job_id]['flashcards'] = []
            jobs[job_id]['test_questions'] = []
            jobs[job_id]['study_generation_error'] = None
        save_study_pack(job_id, jobs[job_id])

        jobs[job_id]['status'] = 'complete'
        jobs[job_id]['step'] = jobs[job_id].get('total_steps', 3)
        jobs[job_id]['step_description'] = 'Complete!'
    except Exception as e:
        jobs[job_id]['status'] = 'error'
        jobs[job_id]['error'] = str(e)
        # Refund the credit since processing failed
        uid = jobs[job_id].get('user_id')
        credit_type = jobs[job_id].get('credit_deducted')
        refund_credit(uid, credit_type)
        jobs[job_id]['credit_refunded'] = True
    finally:
        cleanup_files(local_paths, gemini_files)
        # Log the job to Firestore
        save_job_log(job_id, jobs[job_id], time.time())

def process_slides_only(job_id, pdf_path):
    gemini_files = []
    local_paths = [pdf_path]
    try:
        jobs[job_id]['status'] = 'processing'
        jobs[job_id]['step'] = 1
        jobs[job_id]['step_description'] = 'Extracting text from slides...'
        pdf_file = client.files.upload(file=pdf_path, config={'mime_type': 'application/pdf'})
        gemini_files.append(pdf_file)
        wait_for_file_processing(pdf_file)
        response = client.models.generate_content(model=MODEL_SLIDES, contents=[types.Content(role='user', parts=[types.Part.from_uri(file_uri=pdf_file.uri, mime_type='application/pdf'), types.Part.from_text(text=PROMPT_SLIDE_EXTRACTION)])], config=types.GenerateContentConfig(max_output_tokens=65536))
        extracted_text = response.text
        jobs[job_id]['result'] = extracted_text
        if jobs[job_id].get('study_features', 'none') != 'none':
            jobs[job_id]['step'] = 2
            jobs[job_id]['step_description'] = 'Generating flashcards and practice test...'
            flashcards, test_questions, study_error = generate_study_materials(
                extracted_text,
                jobs[job_id].get('flashcard_selection', '20'),
                jobs[job_id].get('question_selection', '10'),
                jobs[job_id].get('study_features', 'none'),
                jobs[job_id].get('output_language', 'English')
            )
            jobs[job_id]['flashcards'] = flashcards
            jobs[job_id]['test_questions'] = test_questions
            jobs[job_id]['study_generation_error'] = study_error
        else:
            jobs[job_id]['flashcards'] = []
            jobs[job_id]['test_questions'] = []
            jobs[job_id]['study_generation_error'] = None
        save_study_pack(job_id, jobs[job_id])
        jobs[job_id]['status'] = 'complete'
        jobs[job_id]['step'] = jobs[job_id].get('total_steps', 1)
        jobs[job_id]['step_description'] = 'Complete!'
    except Exception as e:
        jobs[job_id]['status'] = 'error'
        jobs[job_id]['error'] = str(e)
        # Refund the credit since processing failed
        uid = jobs[job_id].get('user_id')
        credit_type = jobs[job_id].get('credit_deducted')
        refund_credit(uid, credit_type)
        jobs[job_id]['credit_refunded'] = True
    finally:
        cleanup_files(local_paths, gemini_files)
        # Log the job to Firestore
        save_job_log(job_id, jobs[job_id], time.time())

def process_interview_transcription(job_id, audio_path):
    gemini_files = []
    local_paths = [audio_path]
    try:
        jobs[job_id]['status'] = 'processing'
        jobs[job_id]['step'] = 1
        jobs[job_id]['step_description'] = 'Optimizing audio for faster processing...'
        output_language = jobs[job_id].get('output_language', 'English')
        converted_audio_path, converted = convert_audio_to_mp3_with_ytdlp(audio_path)
        if converted and converted_audio_path not in local_paths:
            local_paths.append(converted_audio_path)
        audio_mime_type = get_mime_type(converted_audio_path)
        audio_file = client.files.upload(file=converted_audio_path, config={'mime_type': audio_mime_type})
        gemini_files.append(audio_file)
        jobs[job_id]['step_description'] = 'Processing audio file (this may take a few minutes)...'
        wait_for_file_processing(audio_file)
        jobs[job_id]['step_description'] = 'Generating transcript with timestamps...'
        interview_prompt = PROMPT_INTERVIEW_TRANSCRIPTION.format(output_language=output_language)
        response = client.models.generate_content(model=MODEL_INTERVIEW, contents=[types.Content(role='user', parts=[types.Part.from_uri(file_uri=audio_file.uri, mime_type=audio_mime_type), types.Part.from_text(text=interview_prompt)])], config=types.GenerateContentConfig(max_output_tokens=65536))
        transcript_text = response.text or ''
        jobs[job_id]['transcript'] = transcript_text
        jobs[job_id]['result'] = transcript_text

        selected_features = jobs[job_id].get('interview_features', [])
        if selected_features:
            jobs[job_id]['step'] = 2
            jobs[job_id]['step_description'] = 'Creating interview summary and sections...'
            enhancement = generate_interview_enhancements(transcript_text, selected_features, output_language)
            jobs[job_id]['interview_summary'] = enhancement.get('summary')
            jobs[job_id]['interview_sections'] = enhancement.get('sections')
            jobs[job_id]['interview_combined'] = enhancement.get('combined')
            jobs[job_id]['interview_features_successful'] = enhancement.get('successful_features', [])
            jobs[job_id]['study_generation_error'] = enhancement.get('error')

            failed_count = enhancement.get('failed_count', 0)
            if failed_count > 0:
                uid = jobs[job_id].get('user_id')
                refund_slides_credits(uid, failed_count)
                jobs[job_id]['extra_slides_refunded'] = jobs[job_id].get('extra_slides_refunded', 0) + failed_count

            if enhancement.get('summary') and enhancement.get('sections'):
                jobs[job_id]['result'] = enhancement.get('combined', transcript_text)
            elif enhancement.get('summary'):
                jobs[job_id]['result'] = enhancement.get('summary')
            elif enhancement.get('sections'):
                jobs[job_id]['result'] = enhancement.get('sections')

        jobs[job_id]['status'] = 'complete'
        jobs[job_id]['step'] = jobs[job_id].get('total_steps', 1)
        jobs[job_id]['step_description'] = 'Complete!'
    except Exception as e:
        jobs[job_id]['status'] = 'error'
        jobs[job_id]['error'] = str(e)
        # Refund the credit since processing failed
        uid = jobs[job_id].get('user_id')
        credit_type = jobs[job_id].get('credit_deducted')
        refund_credit(uid, credit_type)
        extra_spent = jobs[job_id].get('interview_features_cost', 0)
        already_refunded = jobs[job_id].get('extra_slides_refunded', 0)
        to_refund = max(0, extra_spent - already_refunded)
        if to_refund > 0:
            refund_slides_credits(uid, to_refund)
            jobs[job_id]['extra_slides_refunded'] = already_refunded + to_refund
        jobs[job_id]['credit_refunded'] = True
    finally:
        cleanup_files(local_paths, gemini_files)
        # Log the job to Firestore
        save_job_log(job_id, jobs[job_id], time.time())

# =============================================
# ROUTES
# =============================================

@app.route('/')
def index():
    return render_template('landing.html')

@app.route('/dashboard')
def dashboard():
    return render_template('index.html')

@app.route('/plan')
@app.route('/stats')
def plan_dashboard():
    return render_template('plan.html')

@app.route('/calendar')
def calendar_dashboard():
    return render_template('calendar.html')

@app.route('/features')
def features_page():
    return render_template('features.html')

@app.route('/admin')
def admin_dashboard():
    return render_template('admin.html')

@app.route('/study')
def study_dashboard():
    return render_template('study.html')

@app.route('/api/verify-email', methods=['POST'])
def verify_email():
    email = request.get_json().get('email', '')
    if is_email_allowed(email):
        return jsonify({'allowed': True})
    return jsonify({'allowed': False, 'message': 'Please use your university email or a major email provider (Gmail, Outlook, iCloud, Yahoo).'})

@app.route('/api/auth/user', methods=['GET'])
def get_user():
    decoded_token = verify_firebase_token(request)
    if not decoded_token: return jsonify({'error': 'Unauthorized'}), 401
    uid = decoded_token['uid']
    email = decoded_token.get('email', '')
    if not is_email_allowed(email): return jsonify({'error': 'Email not allowed', 'message': 'Please use your university email.'}), 403
    user = get_or_create_user(uid, email)
    return jsonify({
        'uid': user['uid'], 'email': user['email'],
        'credits': {
            'lecture_standard': user.get('lecture_credits_standard', 0),
            'lecture_extended': user.get('lecture_credits_extended', 0),
            'slides': user.get('slides_credits', 0),
            'interview_short': user.get('interview_credits_short', 0),
            'interview_medium': user.get('interview_credits_medium', 0),
            'interview_long': user.get('interview_credits_long', 0),
        },
        'total_processed': user.get('total_processed', 0)
    })

# --- Stripe Routes ---

@app.route('/api/config', methods=['GET'])
def get_config():
    return jsonify({
        'stripe_publishable_key': STRIPE_PUBLISHABLE_KEY,
        'bundles': {
            bundle_id: {
                'name': bundle['name'],
                'description': bundle['description'],
                'price_cents': bundle['price_cents'],
                'currency': bundle['currency'],
                'credits': bundle['credits'],
            }
            for bundle_id, bundle in CREDIT_BUNDLES.items()
        }
    })

@app.route('/api/create-checkout-session', methods=['POST'])
def create_checkout_session():
    decoded_token = verify_firebase_token(request)
    if not decoded_token:
        return jsonify({'error': 'Please sign in to continue'}), 401

    uid = decoded_token['uid']
    email = decoded_token.get('email', '')

    data = request.get_json()
    bundle_id = data.get('bundle_id', '')

    if bundle_id not in CREDIT_BUNDLES:
        return jsonify({'error': 'Invalid bundle selected'}), 400

    bundle = CREDIT_BUNDLES[bundle_id]

    try:
        checkout_session = stripe.checkout.Session.create(
            payment_method_types=['card', 'ideal'],
            line_items=[{
                'price_data': {
                    'currency': bundle['currency'],
                    'product_data': {
                        'name': bundle['name'],
                        'description': bundle['description'],
                    },
                    'unit_amount': bundle['price_cents'],
                },
                'quantity': 1,
            }],
            mode='payment',
            success_url=request.host_url.rstrip('/') + '?payment=success',
            cancel_url=request.host_url.rstrip('/') + '?payment=cancelled',
            customer_email=email,
            metadata={
                'uid': uid,
                'bundle_id': bundle_id,
            },
        )
        return jsonify({'checkout_url': checkout_session.url})
    except Exception as e:
        print(f"Stripe checkout error: {e}")
        return jsonify({'error': 'Could not create checkout session. Please try again.'}), 500

@app.route('/api/stripe-webhook', methods=['POST'])
def stripe_webhook():
    payload = request.data
    sig_header = request.headers.get('Stripe-Signature', '')

    if STRIPE_WEBHOOK_SECRET:
        try:
            event = stripe.Webhook.construct_event(
                payload, sig_header, STRIPE_WEBHOOK_SECRET
            )
        except ValueError:
            print("Stripe webhook: Invalid payload")
            return 'Invalid payload', 400
        except Exception as e:
            print(f"Stripe webhook signature verification failed: {e}")
            return 'Invalid signature', 400
    else:
        try:
            event = json.loads(payload)
        except json.JSONDecodeError:
            return 'Invalid payload', 400

    if event.get('type') == 'checkout.session.completed':
        session = event['data']['object']
        metadata = session.get('metadata', {})
        uid = metadata.get('uid', '')
        bundle_id = metadata.get('bundle_id', '')
        stripe_session_id = session.get('id', '')

        if uid and bundle_id:
            success = grant_credits_to_user(uid, bundle_id)
            if success:
                print(f"✅ Payment successful! Granted bundle '{bundle_id}' to user '{uid}'")
                # Save purchase record for history
                save_purchase_record(uid, bundle_id, stripe_session_id)
            else:
                print(f"❌ Failed to grant bundle '{bundle_id}' to user '{uid}'")
        else:
            print(f"⚠️ Webhook received but missing metadata. uid='{uid}', bundle_id='{bundle_id}'")

    return '', 200

# --- Purchase History Route ---

@app.route('/api/purchase-history', methods=['GET'])
def get_purchase_history():
    decoded_token = verify_firebase_token(request)
    if not decoded_token:
        return jsonify({'error': 'Unauthorized'}), 401

    uid = decoded_token['uid']

    try:
        purchases_ref = db.collection('purchases').where('uid', '==', uid).order_by('created_at', direction=firestore.Query.DESCENDING).limit(50)
        purchases = []
        for doc in purchases_ref.stream():
            p = doc.to_dict()
            purchases.append({
                'id': doc.id,
                'bundle_name': p.get('bundle_name', 'Unknown'),
                'price_cents': p.get('price_cents', 0),
                'currency': p.get('currency', 'eur'),
                'credits': p.get('credits', {}),
                'created_at': p.get('created_at', 0),
            })
        return jsonify({'purchases': purchases})
    except Exception as e:
        print(f"Error fetching purchase history: {e}")
        return jsonify({'purchases': []})

@app.route('/api/study-packs', methods=['GET'])
def get_study_packs():
    decoded_token = verify_firebase_token(request)
    if not decoded_token:
        return jsonify({'error': 'Unauthorized'}), 401

    uid = decoded_token['uid']
    try:
        study_docs = list(db.collection('study_packs').where('uid', '==', uid).limit(200).stream())
        packs = []
        for doc in study_docs:
            pack = doc.to_dict()
            packs.append({
                'study_pack_id': doc.id,
                'title': pack.get('title', ''),
                'mode': pack.get('mode', ''),
                'flashcards_count': len(pack.get('flashcards', [])),
                'test_questions_count': len(pack.get('test_questions', [])),
                'course': pack.get('course', ''),
                'subject': pack.get('subject', ''),
                'semester': pack.get('semester', ''),
                'block': pack.get('block', ''),
                'folder_id': pack.get('folder_id', ''),
                'folder_name': pack.get('folder_name', ''),
                'created_at': pack.get('created_at', 0),
            })
        packs.sort(key=lambda p: p.get('created_at', 0), reverse=True)
        return jsonify({'study_packs': packs[:50]})
    except Exception as e:
        print(f"Error fetching study packs: {e}")
        return jsonify({'study_packs': []})

@app.route('/api/study-packs', methods=['POST'])
def create_study_pack():
    decoded_token = verify_firebase_token(request)
    if not decoded_token:
        return jsonify({'error': 'Unauthorized'}), 401

    uid = decoded_token['uid']
    payload = request.get_json() or {}

    title = str(payload.get('title', '')).strip()[:120]
    if not title:
        title = f"Untitled pack {datetime.utcnow().strftime('%Y-%m-%d %H:%M')}"

    try:
        now_ts = time.time()
        folder_id = str(payload.get('folder_id', '')).strip()
        folder_name = ''
        if folder_id:
            folder_doc = db.collection('study_folders').document(folder_id).get()
            if not folder_doc.exists:
                return jsonify({'error': 'Folder not found'}), 404
            folder_data = folder_doc.to_dict()
            if folder_data.get('uid', '') != uid:
                return jsonify({'error': 'Forbidden'}), 403
            folder_name = folder_data.get('name', '')
        else:
            folder_id = ''

        flashcards = sanitize_flashcards(payload.get('flashcards', []), 500)
        test_questions = sanitize_questions(payload.get('test_questions', []), 500)
        notes_markdown = str(payload.get('notes_markdown', '')).strip()
        notes_audio_map = parse_audio_markers_from_notes(notes_markdown) if FEATURE_AUDIO_SECTION_SYNC else []

        doc_ref = db.collection('study_packs').document()
        doc_ref.set({
            'study_pack_id': doc_ref.id,
            'source_job_id': '',
            'uid': uid,
            'email': decoded_token.get('email', ''),
            'mode': 'manual',
            'title': title,
            'output_language': str(payload.get('output_language', 'English')).strip()[:64] or 'English',
            'notes_markdown': notes_markdown,
            'notes_truncated': False,
            'transcript_segments': [],
            'notes_audio_map': notes_audio_map,
            'audio_storage_path': '',
            'has_audio_sync': False,
            'has_audio_playback': False,
            'flashcards': flashcards,
            'test_questions': test_questions,
            'flashcard_selection': 'manual',
            'question_selection': 'manual',
            'study_features': 'both',
            'interview_features': [],
            'interview_summary': None,
            'interview_sections': None,
            'interview_combined': None,
            'study_generation_error': None,
            'course': str(payload.get('course', '')).strip()[:120],
            'subject': str(payload.get('subject', '')).strip()[:120],
            'semester': str(payload.get('semester', '')).strip()[:120],
            'block': str(payload.get('block', '')).strip()[:120],
            'folder_id': folder_id,
            'folder_name': folder_name,
            'created_at': now_ts,
            'updated_at': now_ts,
        })

        return jsonify({'ok': True, 'study_pack_id': doc_ref.id})
    except Exception as e:
        print(f"Error creating study pack: {e}")
        return jsonify({'error': 'Could not create study pack'}), 500

@app.route('/api/study-packs/<pack_id>', methods=['GET'])
def get_study_pack(pack_id):
    decoded_token = verify_firebase_token(request)
    if not decoded_token:
        return jsonify({'error': 'Unauthorized'}), 401

    uid = decoded_token['uid']
    try:
        doc = db.collection('study_packs').document(pack_id).get()
        if not doc.exists:
            return jsonify({'error': 'Study pack not found'}), 404
        pack = doc.to_dict()
        if pack.get('uid', '') != uid:
            return jsonify({'error': 'Forbidden'}), 403
        has_audio_playback = bool(pack.get('has_audio_playback', False) or pack.get('audio_storage_path', ''))
        has_audio_sync = FEATURE_AUDIO_SECTION_SYNC and bool(pack.get('has_audio_sync', False))
        notes_audio_map = pack.get('notes_audio_map', []) if has_audio_sync else []
        return jsonify({
            'study_pack_id': pack_id,
            'title': pack.get('title', ''),
            'mode': pack.get('mode', ''),
            'output_language': pack.get('output_language', 'English'),
            'notes_markdown': pack.get('notes_markdown', ''),
            'transcript_segments': pack.get('transcript_segments', []),
            'notes_audio_map': notes_audio_map,
            'has_audio_sync': has_audio_sync,
            'has_audio_playback': has_audio_playback,
            'flashcards': pack.get('flashcards', []),
            'test_questions': pack.get('test_questions', []),
            'interview_summary': pack.get('interview_summary'),
            'interview_sections': pack.get('interview_sections'),
            'interview_combined': pack.get('interview_combined'),
            'study_features': pack.get('study_features', 'none'),
            'interview_features': pack.get('interview_features', []),
            'course': pack.get('course', ''),
            'subject': pack.get('subject', ''),
            'semester': pack.get('semester', ''),
            'block': pack.get('block', ''),
            'folder_id': pack.get('folder_id', ''),
            'folder_name': pack.get('folder_name', ''),
            'created_at': pack.get('created_at', 0),
        })
    except Exception as e:
        print(f"Error fetching study pack {pack_id}: {e}")
        return jsonify({'error': 'Could not fetch study pack'}), 500

@app.route('/api/study-packs/<pack_id>', methods=['PATCH'])
def update_study_pack(pack_id):
    decoded_token = verify_firebase_token(request)
    if not decoded_token:
        return jsonify({'error': 'Unauthorized'}), 401
    uid = decoded_token['uid']
    payload = request.get_json() or {}

    try:
        pack_ref = db.collection('study_packs').document(pack_id)
        doc = pack_ref.get()
        if not doc.exists:
            return jsonify({'error': 'Study pack not found'}), 404
        pack = doc.to_dict()
        if pack.get('uid', '') != uid:
            return jsonify({'error': 'Forbidden'}), 403

        updates = {'updated_at': time.time()}
        if 'title' in payload:
            updates['title'] = str(payload.get('title', '')).strip()[:120]
        if 'course' in payload:
            updates['course'] = str(payload.get('course', '')).strip()[:120]
        if 'subject' in payload:
            updates['subject'] = str(payload.get('subject', '')).strip()[:120]
        if 'semester' in payload:
            updates['semester'] = str(payload.get('semester', '')).strip()[:120]
        if 'block' in payload:
            updates['block'] = str(payload.get('block', '')).strip()[:120]
        if 'folder_id' in payload:
            folder_id = str(payload.get('folder_id', '')).strip()
            updates['folder_id'] = ''
            updates['folder_name'] = ''
            if folder_id:
                folder_doc = db.collection('study_folders').document(folder_id).get()
                if not folder_doc.exists:
                    return jsonify({'error': 'Folder not found'}), 404
                folder_data = folder_doc.to_dict()
                if folder_data.get('uid', '') != uid:
                    return jsonify({'error': 'Forbidden'}), 403
                updates['folder_id'] = folder_id
                updates['folder_name'] = folder_data.get('name', '')

        if 'flashcards' in payload:
            updates['flashcards'] = sanitize_flashcards(payload.get('flashcards', []), 500)
        if 'test_questions' in payload:
            updates['test_questions'] = sanitize_questions(payload.get('test_questions', []), 500)
        if 'notes_markdown' in payload:
            updates['notes_markdown'] = str(payload.get('notes_markdown', ''))
            notes_audio_map = parse_audio_markers_from_notes(updates['notes_markdown']) if FEATURE_AUDIO_SECTION_SYNC else []
            updates['notes_audio_map'] = notes_audio_map
            updates['has_audio_sync'] = FEATURE_AUDIO_SECTION_SYNC and bool(pack.get('audio_storage_path', '')) and bool(notes_audio_map)
        updates['has_audio_playback'] = bool(pack.get('audio_storage_path', ''))

        pack_ref.update(updates)
        return jsonify({'ok': True})
    except Exception as e:
        print(f"Error updating study pack {pack_id}: {e}")
        return jsonify({'error': 'Could not update study pack'}), 500

@app.route('/api/study-packs/<pack_id>', methods=['DELETE'])
def delete_study_pack(pack_id):
    decoded_token = verify_firebase_token(request)
    if not decoded_token:
        return jsonify({'error': 'Unauthorized'}), 401
    uid = decoded_token['uid']
    try:
        pack_ref = db.collection('study_packs').document(pack_id)
        doc = pack_ref.get()
        if not doc.exists:
            return jsonify({'error': 'Study pack not found'}), 404
        pack = doc.to_dict()
        if pack.get('uid', '') != uid:
            return jsonify({'error': 'Forbidden'}), 403
        audio_storage_path = pack.get('audio_storage_path', '')
        if audio_storage_path and isinstance(audio_storage_path, str):
            try:
                if os.path.exists(audio_storage_path):
                    os.remove(audio_storage_path)
            except Exception as e:
                print(f"Warning: could not delete study-pack audio file {audio_storage_path}: {e}")
        pack_ref.delete()
        return jsonify({'ok': True})
    except Exception as e:
        print(f"Error deleting study pack {pack_id}: {e}")
        return jsonify({'error': 'Could not delete study pack'}), 500

@app.route('/api/study-folders', methods=['GET'])
def get_study_folders():
    decoded_token = verify_firebase_token(request)
    if not decoded_token:
        return jsonify({'error': 'Unauthorized'}), 401
    uid = decoded_token['uid']
    try:
        docs = list(db.collection('study_folders').where('uid', '==', uid).stream())
        folders = []
        for doc in docs:
            folder = doc.to_dict()
            folders.append({
                'folder_id': doc.id,
                'name': folder.get('name', ''),
                'course': folder.get('course', ''),
                'subject': folder.get('subject', ''),
                'semester': folder.get('semester', ''),
                'block': folder.get('block', ''),
                'exam_date': folder.get('exam_date', ''),
                'created_at': folder.get('created_at', 0),
            })
        folders.sort(key=lambda f: f.get('created_at', 0), reverse=True)
        return jsonify({'folders': folders})
    except Exception as e:
        print(f"Error fetching study folders: {e}")
        return jsonify({'folders': []})

@app.route('/api/study-packs/<pack_id>/audio-url', methods=['GET'])
def get_study_pack_audio_url(pack_id):
    decoded_token = verify_firebase_token(request)
    if not decoded_token:
        return jsonify({'error': 'Unauthorized'}), 401
    uid = decoded_token['uid']
    try:
        doc = db.collection('study_packs').document(pack_id).get()
        if not doc.exists:
            return jsonify({'error': 'Study pack not found'}), 404
        pack = doc.to_dict()
        if pack.get('uid', '') != uid:
            return jsonify({'error': 'Forbidden'}), 403
        audio_storage_path = str(pack.get('audio_storage_path', '') or '').strip()
        if not audio_storage_path:
            return jsonify({'error': 'No audio file for this study pack'}), 404
        if not os.path.exists(audio_storage_path):
            return jsonify({'error': 'Audio file not found'}), 404
        stream_token = str(uuid.uuid4())
        AUDIO_STREAM_TOKENS[stream_token] = {
            'path': audio_storage_path,
            'expires_at': time.time() + AUDIO_STREAM_TOKEN_TTL_SECONDS
        }
        return jsonify({'audio_url': f"/api/audio-stream/{stream_token}"})
    except Exception as e:
        print(f"Error generating study-pack audio URL {pack_id}: {e}")
        return jsonify({'error': 'Could not generate audio URL'}), 500

@app.route('/api/audio-stream/<token>', methods=['GET'])
def stream_audio_token(token):
    token_data = AUDIO_STREAM_TOKENS.get(token)
    if not token_data:
        return jsonify({'error': 'Invalid token'}), 404
    if time.time() > token_data.get('expires_at', 0):
        AUDIO_STREAM_TOKENS.pop(token, None)
        return jsonify({'error': 'Token expired'}), 410
    file_path = token_data.get('path', '')
    if not file_path or not os.path.exists(file_path):
        return jsonify({'error': 'Audio file not found'}), 404
    mime_type = get_mime_type(file_path)
    return send_file(file_path, mimetype=mime_type, conditional=True)

@app.route('/api/study-folders', methods=['POST'])
def create_study_folder():
    decoded_token = verify_firebase_token(request)
    if not decoded_token:
        return jsonify({'error': 'Unauthorized'}), 401
    uid = decoded_token['uid']
    payload = request.get_json() or {}
    name = str(payload.get('name', '')).strip()[:120]
    if not name:
        return jsonify({'error': 'Folder name is required'}), 400
    try:
        now_ts = time.time()
        try:
            exam_date = normalize_exam_date(payload.get('exam_date', ''))
        except ValueError as ve:
            return jsonify({'error': str(ve)}), 400
        doc_ref = db.collection('study_folders').document()
        doc_ref.set({
            'folder_id': doc_ref.id,
            'uid': uid,
            'name': name,
            'course': str(payload.get('course', '')).strip()[:120],
            'subject': str(payload.get('subject', '')).strip()[:120],
            'semester': str(payload.get('semester', '')).strip()[:120],
            'block': str(payload.get('block', '')).strip()[:120],
            'exam_date': exam_date,
            'created_at': now_ts,
            'updated_at': now_ts,
        })
        return jsonify({'ok': True, 'folder_id': doc_ref.id})
    except Exception as e:
        print(f"Error creating study folder: {e}")
        return jsonify({'error': 'Could not create folder'}), 500

@app.route('/api/study-folders/<folder_id>', methods=['PATCH'])
def update_study_folder(folder_id):
    decoded_token = verify_firebase_token(request)
    if not decoded_token:
        return jsonify({'error': 'Unauthorized'}), 401
    uid = decoded_token['uid']
    payload = request.get_json() or {}
    try:
        folder_ref = db.collection('study_folders').document(folder_id)
        doc = folder_ref.get()
        if not doc.exists:
            return jsonify({'error': 'Folder not found'}), 404
        folder = doc.to_dict()
        if folder.get('uid', '') != uid:
            return jsonify({'error': 'Forbidden'}), 403
        updates = {'updated_at': time.time()}
        if 'name' in payload:
            name = str(payload.get('name', '')).strip()[:120]
            if not name:
                return jsonify({'error': 'Folder name is required'}), 400
            updates['name'] = name
        for field in ['course', 'subject', 'semester', 'block']:
            if field in payload:
                updates[field] = str(payload.get(field, '')).strip()[:120]
        if 'exam_date' in payload:
            try:
                updates['exam_date'] = normalize_exam_date(payload.get('exam_date', ''))
            except ValueError as ve:
                return jsonify({'error': str(ve)}), 400
        folder_ref.update(updates)
        if 'name' in updates:
            packs = list(db.collection('study_packs').where('uid', '==', uid).where('folder_id', '==', folder_id).stream())
            for pack_doc in packs:
                pack_doc.reference.update({'folder_name': updates['name'], 'updated_at': time.time()})
        return jsonify({'ok': True})
    except Exception as e:
        print(f"Error updating folder {folder_id}: {e}")
        return jsonify({'error': 'Could not update folder'}), 500

@app.route('/api/study-folders/<folder_id>', methods=['DELETE'])
def delete_study_folder(folder_id):
    decoded_token = verify_firebase_token(request)
    if not decoded_token:
        return jsonify({'error': 'Unauthorized'}), 401
    uid = decoded_token['uid']
    try:
        folder_ref = db.collection('study_folders').document(folder_id)
        doc = folder_ref.get()
        if not doc.exists:
            return jsonify({'error': 'Folder not found'}), 404
        folder = doc.to_dict()
        if folder.get('uid', '') != uid:
            return jsonify({'error': 'Forbidden'}), 403
        folder_ref.delete()
        packs = list(db.collection('study_packs').where('uid', '==', uid).where('folder_id', '==', folder_id).stream())
        for pack_doc in packs:
            pack_doc.reference.update({'folder_id': '', 'folder_name': '', 'updated_at': time.time()})
        return jsonify({'ok': True})
    except Exception as e:
        print(f"Error deleting folder {folder_id}: {e}")
        return jsonify({'error': 'Could not delete folder'}), 500

@app.route('/api/admin/overview', methods=['GET'])
def get_admin_overview():
    decoded_token = verify_firebase_token(request)
    if not decoded_token:
        return jsonify({'error': 'Unauthorized'}), 401
    if not is_admin_user(decoded_token):
        return jsonify({'error': 'Forbidden'}), 403

    try:
        window_key, window_seconds = get_admin_window(request.args.get('window', '7d'))
        now_ts = time.time()
        window_start = now_ts - window_seconds

        users_docs = list(db.collection('users').stream())
        purchases_docs = list(db.collection('purchases').stream())
        jobs_docs = list(db.collection('job_logs').stream())

        total_users = len(users_docs)
        new_users = 0
        for doc in users_docs:
            created_at = get_timestamp(doc.to_dict().get('created_at'))
            if created_at >= window_start:
                new_users += 1
        total_processed = sum((doc.to_dict().get('total_processed', 0) or 0) for doc in users_docs)

        total_revenue_cents = 0
        purchase_count = 0
        filtered_purchases = []
        for doc in purchases_docs:
            purchase = doc.to_dict()
            created_at = get_timestamp(purchase.get('created_at'))
            if created_at < window_start:
                continue
            filtered_purchases.append(purchase)
            purchase_count += 1
            total_revenue_cents += purchase.get('price_cents', 0) or 0

        job_count = 0
        success_jobs = 0
        failed_jobs = 0
        refunded_jobs = 0
        durations = []
        filtered_jobs = []
        for doc in jobs_docs:
            job = doc.to_dict()
            finished_at = get_timestamp(job.get('finished_at'))
            if finished_at < window_start:
                continue
            filtered_jobs.append(job)
            job_count += 1
            status = job.get('status', '')
            if status == 'complete':
                success_jobs += 1
            elif status == 'error':
                failed_jobs += 1
            if job.get('credit_refunded'):
                refunded_jobs += 1
            duration = job.get('duration_seconds')
            if isinstance(duration, (int, float)):
                durations.append(duration)

        avg_duration_seconds = round(sum(durations) / len(durations), 1) if durations else 0

        mode_breakdown = {
            'lecture-notes': {'label': 'Lecture Notes', 'total': 0, 'complete': 0, 'error': 0},
            'slides-only': {'label': 'Slides Only', 'total': 0, 'complete': 0, 'error': 0},
            'interview': {'label': 'Interview Transcript', 'total': 0, 'complete': 0, 'error': 0},
            'other': {'label': 'Other', 'total': 0, 'complete': 0, 'error': 0},
        }
        for job in filtered_jobs:
            mode = job.get('mode', '')
            key = mode if mode in mode_breakdown else 'other'
            status = job.get('status', '')
            mode_breakdown[key]['total'] += 1
            if status == 'complete':
                mode_breakdown[key]['complete'] += 1
            elif status == 'error':
                mode_breakdown[key]['error'] += 1

        recent_jobs_sorted = sorted(
            filtered_jobs,
            key=lambda j: get_timestamp(j.get('finished_at')),
            reverse=True
        )[:20]
        recent_jobs = []
        for job in recent_jobs_sorted:
            recent_jobs.append({
                'job_id': job.get('job_id', ''),
                'email': job.get('email', ''),
                'mode': job.get('mode', ''),
                'status': job.get('status', ''),
                'duration_seconds': job.get('duration_seconds', 0),
                'credit_refunded': job.get('credit_refunded', False),
                'finished_at': job.get('finished_at', 0),
            })

        recent_purchases_sorted = sorted(
            filtered_purchases,
            key=lambda p: get_timestamp(p.get('created_at')),
            reverse=True
        )[:20]
        recent_purchases = []
        for purchase in recent_purchases_sorted:
            recent_purchases.append({
                'uid': purchase.get('uid', ''),
                'bundle_name': purchase.get('bundle_name', 'Unknown'),
                'price_cents': purchase.get('price_cents', 0),
                'currency': purchase.get('currency', 'eur'),
                'created_at': purchase.get('created_at', 0),
            })

        trend_labels, trend_keys, trend_granularity = build_time_buckets(window_key, now_ts)
        success_by_bucket = {key: {'complete': 0, 'error': 0} for key in trend_keys}
        revenue_by_bucket = {key: 0 for key in trend_keys}

        for job in filtered_jobs:
            timestamp = get_timestamp(job.get('finished_at'))
            bucket_key = get_bucket_key(timestamp, window_key)
            if bucket_key not in success_by_bucket:
                continue
            status = job.get('status', '')
            if status == 'complete':
                success_by_bucket[bucket_key]['complete'] += 1
            elif status == 'error':
                success_by_bucket[bucket_key]['error'] += 1

        for purchase in filtered_purchases:
            timestamp = get_timestamp(purchase.get('created_at'))
            bucket_key = get_bucket_key(timestamp, window_key)
            if bucket_key not in revenue_by_bucket:
                continue
            revenue_by_bucket[bucket_key] += purchase.get('price_cents', 0) or 0

        success_trend = []
        revenue_trend = []
        for key in trend_keys:
            complete_count = success_by_bucket[key]['complete']
            error_count = success_by_bucket[key]['error']
            total_count = complete_count + error_count
            success_rate = round((complete_count / total_count) * 100, 1) if total_count > 0 else 0
            success_trend.append(success_rate)
            revenue_trend.append(revenue_by_bucket[key])

        return jsonify({
            'window': {
                'key': window_key,
                'start': window_start,
                'end': now_ts,
            },
            'metrics': {
                'total_users': total_users,
                'new_users': new_users,
                'total_processed': total_processed,
                'total_revenue_cents': total_revenue_cents,
                'purchase_count': purchase_count,
                'job_count': job_count,
                'success_jobs': success_jobs,
                'failed_jobs': failed_jobs,
                'refunded_jobs': refunded_jobs,
                'avg_duration_seconds': avg_duration_seconds,
            },
            'trends': {
                'labels': trend_labels,
                'success_rate': success_trend,
                'revenue_cents': revenue_trend,
                'granularity': trend_granularity,
            },
            'mode_breakdown': mode_breakdown,
            'recent_jobs': recent_jobs,
            'recent_purchases': recent_purchases,
        })
    except Exception as e:
        print(f"Error fetching admin overview: {e}")
        return jsonify({'error': 'Could not fetch admin dashboard data'}), 500

@app.route('/api/admin/export', methods=['GET'])
def export_admin_csv():
    decoded_token = verify_firebase_token(request)
    if not decoded_token:
        return jsonify({'error': 'Unauthorized'}), 401
    if not is_admin_user(decoded_token):
        return jsonify({'error': 'Forbidden'}), 403

    export_type = request.args.get('type', 'jobs')
    if export_type not in {'jobs', 'purchases'}:
        return jsonify({'error': 'Invalid export type'}), 400

    window_key, window_seconds = get_admin_window(request.args.get('window', '7d'))
    now_ts = time.time()
    window_start = now_ts - window_seconds

    output = io.StringIO()
    writer = csv.writer(output)

    try:
        if export_type == 'jobs':
            writer.writerow([
                'job_id', 'uid', 'email', 'mode', 'status', 'credit_deducted',
                'credit_refunded', 'error_message', 'started_at', 'finished_at', 'duration_seconds'
            ])
            for doc in db.collection('job_logs').stream():
                job = doc.to_dict()
                finished_at = get_timestamp(job.get('finished_at'))
                if finished_at < window_start:
                    continue
                writer.writerow([
                    job.get('job_id', doc.id),
                    job.get('uid', ''),
                    job.get('email', ''),
                    job.get('mode', ''),
                    job.get('status', ''),
                    job.get('credit_deducted', ''),
                    job.get('credit_refunded', False),
                    job.get('error_message', ''),
                    job.get('started_at', 0),
                    job.get('finished_at', 0),
                    job.get('duration_seconds', 0),
                ])
        else:
            writer.writerow([
                'uid', 'bundle_id', 'bundle_name', 'price_cents', 'currency',
                'credits', 'stripe_session_id', 'created_at'
            ])
            for doc in db.collection('purchases').stream():
                purchase = doc.to_dict()
                created_at = get_timestamp(purchase.get('created_at'))
                if created_at < window_start:
                    continue
                writer.writerow([
                    purchase.get('uid', ''),
                    purchase.get('bundle_id', ''),
                    purchase.get('bundle_name', ''),
                    purchase.get('price_cents', 0),
                    purchase.get('currency', 'eur'),
                    json.dumps(purchase.get('credits', {}), ensure_ascii=True),
                    purchase.get('stripe_session_id', ''),
                    purchase.get('created_at', 0),
                ])

        filename = f"admin-{export_type}-{window_key}.csv"
        response = app.response_class(output.getvalue(), mimetype='text/csv')
        response.headers['Content-Disposition'] = f'attachment; filename={filename}'
        return response
    except Exception as e:
        print(f"Error exporting admin CSV ({export_type}): {e}")
        return jsonify({'error': 'Could not export CSV'}), 500

# --- Upload & Processing Routes ---

@app.route('/upload', methods=['POST'])
def upload_files():
    decoded_token = verify_firebase_token(request)
    if not decoded_token: return jsonify({'error': 'Please sign in to continue'}), 401
    uid = decoded_token['uid']
    email = decoded_token.get('email', '')
    if not is_email_allowed(email): return jsonify({'error': 'Email not allowed'}), 403
    user = get_or_create_user(uid, email)
    mode = request.form.get('mode', 'lecture-notes')
    flashcard_selection = parse_requested_amount(request.form.get('flashcard_amount', '20'), {'10', '20', '30', 'auto'}, '20')
    question_selection = parse_requested_amount(request.form.get('question_amount', '10'), {'5', '10', '15', 'auto'}, '10')
    output_language = parse_output_language(request.form.get('output_language', 'english'), request.form.get('output_language_custom', ''))
    study_features = parse_study_features(request.form.get('study_features', 'none'))
    interview_features = parse_interview_features(request.form.get('interview_features', 'none'))
    
    if mode == 'lecture-notes':
        total_lecture = user.get('lecture_credits_standard', 0) + user.get('lecture_credits_extended', 0)
        if total_lecture <= 0:
            return jsonify({'error': 'No lecture credits remaining. Please purchase more credits.'}), 402
        if 'pdf' not in request.files or 'audio' not in request.files: return jsonify({'error': 'Both PDF and audio files are required'}), 400
        pdf_file = request.files['pdf']
        audio_file = request.files['audio']
        if pdf_file.filename == '' or audio_file.filename == '': return jsonify({'error': 'Both files must be selected'}), 400
        if not allowed_file(pdf_file.filename, ALLOWED_PDF_EXTENSIONS): return jsonify({'error': 'Invalid PDF file'}), 400
        if not allowed_file(audio_file.filename, ALLOWED_AUDIO_EXTENSIONS): return jsonify({'error': 'Invalid audio file'}), 400
        job_id = str(uuid.uuid4())
        pdf_path = os.path.join(UPLOAD_FOLDER, f"{job_id}_{secure_filename(pdf_file.filename)}")
        audio_path = os.path.join(UPLOAD_FOLDER, f"{job_id}_{secure_filename(audio_file.filename)}")
        pdf_file.save(pdf_path)
        audio_file.save(audio_path)
        deducted = deduct_credit(uid, 'lecture_credits_standard', 'lecture_credits_extended')
        if not deducted:
            return jsonify({'error': 'No lecture credits remaining.'}), 402
        total_steps = 4 if study_features != 'none' else 3
        jobs[job_id] = {'status': 'starting', 'step': 0, 'step_description': 'Starting...', 'total_steps': total_steps, 'mode': 'lecture-notes', 'user_id': uid, 'user_email': email, 'credit_deducted': deducted, 'credit_refunded': False, 'started_at': time.time(), 'result': None, 'slide_text': None, 'transcript': None, 'flashcard_selection': flashcard_selection, 'question_selection': question_selection, 'study_features': study_features, 'output_language': output_language, 'flashcards': [], 'test_questions': [], 'study_generation_error': None, 'study_pack_id': None, 'error': None}
        thread = threading.Thread(target=process_lecture_notes, args=(job_id, pdf_path, audio_path))
        thread.start()
        
    elif mode == 'slides-only':
        if user.get('slides_credits', 0) <= 0:
            return jsonify({'error': 'No slides credits remaining. Please purchase more credits.'}), 402
        if 'pdf' not in request.files: return jsonify({'error': 'PDF file is required'}), 400
        pdf_file = request.files['pdf']
        if pdf_file.filename == '': return jsonify({'error': 'PDF file must be selected'}), 400
        if not allowed_file(pdf_file.filename, ALLOWED_PDF_EXTENSIONS): return jsonify({'error': 'Invalid PDF file'}), 400
        job_id = str(uuid.uuid4())
        pdf_path = os.path.join(UPLOAD_FOLDER, f"{job_id}_{secure_filename(pdf_file.filename)}")
        pdf_file.save(pdf_path)
        deducted = deduct_credit(uid, 'slides_credits')
        if not deducted:
            return jsonify({'error': 'No slides credits remaining.'}), 402
        total_steps = 2 if study_features != 'none' else 1
        jobs[job_id] = {'status': 'starting', 'step': 0, 'step_description': 'Starting...', 'total_steps': total_steps, 'mode': 'slides-only', 'user_id': uid, 'user_email': email, 'credit_deducted': deducted, 'credit_refunded': False, 'started_at': time.time(), 'result': None, 'flashcard_selection': flashcard_selection, 'question_selection': question_selection, 'study_features': study_features, 'output_language': output_language, 'flashcards': [], 'test_questions': [], 'study_generation_error': None, 'study_pack_id': None, 'error': None}
        thread = threading.Thread(target=process_slides_only, args=(job_id, pdf_path))
        thread.start()
        
    elif mode == 'interview':
        total_interview = user.get('interview_credits_short', 0) + user.get('interview_credits_medium', 0) + user.get('interview_credits_long', 0)
        if total_interview <= 0:
            return jsonify({'error': 'No interview credits remaining. Please purchase more credits.'}), 402
        if 'audio' not in request.files: return jsonify({'error': 'Audio file is required'}), 400
        audio_file = request.files['audio']
        if audio_file.filename == '': return jsonify({'error': 'Audio file must be selected'}), 400
        if not allowed_file(audio_file.filename, ALLOWED_AUDIO_EXTENSIONS): return jsonify({'error': 'Invalid audio file'}), 400
        job_id = str(uuid.uuid4())
        audio_path = os.path.join(UPLOAD_FOLDER, f"{job_id}_{secure_filename(audio_file.filename)}")
        audio_file.save(audio_path)
        deducted = deduct_interview_credit(uid)
        if not deducted:
            return jsonify({'error': 'No interview credits remaining.'}), 402
        interview_features_cost = len(interview_features)
        if interview_features_cost > 0:
            if user.get('slides_credits', 0) < interview_features_cost:
                refund_credit(uid, deducted)
                return jsonify({'error': f'Not enough slides credits for interview extras. You selected {interview_features_cost} option(s) and need {interview_features_cost} slides credits.'}), 402
            if not deduct_slides_credits(uid, interview_features_cost):
                refund_credit(uid, deducted)
                return jsonify({'error': 'Could not reserve slides credits for interview extras. Please try again.'}), 402
        total_steps = 2 if interview_features_cost > 0 else 1
        jobs[job_id] = {
            'status': 'starting',
            'step': 0,
            'step_description': 'Starting...',
            'total_steps': total_steps,
            'mode': 'interview',
            'user_id': uid,
            'user_email': email,
            'credit_deducted': deducted,
            'credit_refunded': False,
            'started_at': time.time(),
            'result': None,
            'transcript': None,
            'flashcards': [],
            'test_questions': [],
            'study_features': 'none',
            'output_language': output_language,
            'interview_features': interview_features,
            'interview_features_cost': interview_features_cost,
            'interview_features_successful': [],
            'interview_summary': None,
            'interview_sections': None,
            'interview_combined': None,
            'extra_slides_refunded': 0,
            'study_generation_error': None,
            'error': None
        }
        thread = threading.Thread(target=process_interview_transcription, args=(job_id, audio_path))
        thread.start()
        
    return jsonify({'job_id': job_id})

@app.route('/status/<job_id>')
def get_status(job_id):
    if job_id not in jobs: return jsonify({'error': 'Job not found'}), 404
    job = jobs[job_id]
    response = {'status': job['status'], 'step': job['step'], 'step_description': job['step_description'], 'total_steps': job.get('total_steps', 3), 'mode': job.get('mode', 'lecture-notes')}
    if job['status'] == 'complete':
        response['result'] = job['result']
        response['flashcards'] = job.get('flashcards', [])
        response['test_questions'] = job.get('test_questions', [])
        response['study_generation_error'] = job.get('study_generation_error')
        response['study_pack_id'] = job.get('study_pack_id')
        response['study_features'] = job.get('study_features', 'none')
        response['output_language'] = job.get('output_language', 'English')
        response['interview_features'] = job.get('interview_features', [])
        response['interview_features_successful'] = job.get('interview_features_successful', [])
        response['interview_summary'] = job.get('interview_summary')
        response['interview_sections'] = job.get('interview_sections')
        response['interview_combined'] = job.get('interview_combined')
        if job.get('mode') == 'lecture-notes':
            response['slide_text'] = job.get('slide_text')
            response['transcript'] = job.get('transcript')
        if job.get('mode') == 'interview':
            response['transcript'] = job.get('transcript')
    elif job['status'] == 'error':
        response['error'] = job['error']
        response['credit_refunded'] = job.get('credit_refunded', False)
    return jsonify(response)

@app.route('/download-docx/<job_id>')
def download_docx(job_id):
    if job_id not in jobs: return jsonify({'error': 'Job not found'}), 404
    job = jobs[job_id]
    if job['status'] != 'complete': return jsonify({'error': 'Job not complete'}), 400
    content_type = request.args.get('type', 'result')
    
    if content_type == 'slides' and job.get('slide_text'):
        content, filename, title = job['slide_text'], 'extracted-slides.docx', 'Extracted Slides'
    elif content_type == 'transcript' and job.get('transcript'):
        content, filename, title = job['transcript'], 'transcript.docx', 'Lecture Transcript'
    elif content_type == 'summary' and job.get('interview_summary'):
        content, filename, title = job['interview_summary'], 'interview-summary.docx', 'Interview Summary'
    elif content_type == 'sections' and job.get('interview_sections'):
        content, filename, title = job['interview_sections'], 'interview-structured.docx', 'Structured Interview Transcript'
    elif content_type == 'combined' and job.get('interview_combined'):
        content, filename, title = job['interview_combined'], 'interview-summary-structured.docx', 'Interview Summary + Structured Transcript'
    else:
        content = job['result']
        mode = job.get('mode', 'lecture-notes')
        if mode == 'lecture-notes': filename, title = 'lecture-notes.docx', 'Lecture Notes'
        elif mode == 'slides-only': filename, title = 'extracted-slides.docx', 'Extracted Slides'
        else: filename, title = 'interview-transcript.docx', 'Interview Transcript'
        
    doc = markdown_to_docx(content, title)
    docx_io = io.BytesIO()
    doc.save(docx_io)
    docx_io.seek(0)
    return send_file(docx_io, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', as_attachment=True, download_name=filename)

@app.route('/download-flashcards-csv/<job_id>')
def download_flashcards_csv(job_id):
    if job_id not in jobs:
        return jsonify({'error': 'Job not found'}), 404
    job = jobs[job_id]
    if job.get('status') != 'complete':
        return jsonify({'error': 'Job not complete'}), 400
    export_type = request.args.get('type', 'flashcards').strip().lower()

    output = io.StringIO()
    writer = csv.writer(output)
    if export_type == 'test':
        test_questions = job.get('test_questions', [])
        if not test_questions:
            return jsonify({'error': 'No practice questions available for this job'}), 400
        writer.writerow(['question', 'option_a', 'option_b', 'option_c', 'option_d', 'answer', 'explanation'])
        for q in test_questions:
            options = q.get('options', [])
            padded = (options + ['', '', '', ''])[:4]
            writer.writerow([
                q.get('question', ''),
                padded[0],
                padded[1],
                padded[2],
                padded[3],
                q.get('answer', ''),
                q.get('explanation', ''),
            ])
        filename = f'practice-test-{job_id}.csv'
    else:
        flashcards = job.get('flashcards', [])
        if not flashcards:
            return jsonify({'error': 'No flashcards available for this job'}), 400
        writer.writerow(['question', 'answer'])
        for card in flashcards:
            writer.writerow([card.get('front', ''), card.get('back', '')])
        filename = f'flashcards-{job_id}.csv'

    csv_bytes = io.BytesIO(output.getvalue().encode('utf-8'))
    csv_bytes.seek(0)
    return send_file(csv_bytes, mimetype='text/csv', as_attachment=True, download_name=filename)

@app.route('/api/study-packs/<pack_id>/export-flashcards-csv', methods=['GET'])
def export_study_pack_flashcards_csv(pack_id):
    decoded_token = verify_firebase_token(request)
    if not decoded_token:
        return jsonify({'error': 'Unauthorized'}), 401
    uid = decoded_token['uid']
    try:
        doc = db.collection('study_packs').document(pack_id).get()
        if not doc.exists:
            return jsonify({'error': 'Study pack not found'}), 404
        pack = doc.to_dict()
        if pack.get('uid', '') != uid:
            return jsonify({'error': 'Forbidden'}), 403
        export_type = request.args.get('type', 'flashcards').strip().lower()
        output = io.StringIO()
        writer = csv.writer(output)
        if export_type == 'test':
            test_questions = pack.get('test_questions', [])
            if not test_questions:
                return jsonify({'error': 'No practice questions available'}), 400
            writer.writerow(['question', 'option_a', 'option_b', 'option_c', 'option_d', 'answer', 'explanation'])
            for q in test_questions:
                options = q.get('options', [])
                padded = (options + ['', '', '', ''])[:4]
                writer.writerow([
                    q.get('question', ''),
                    padded[0],
                    padded[1],
                    padded[2],
                    padded[3],
                    q.get('answer', ''),
                    q.get('explanation', ''),
                ])
            filename = f'study-pack-{pack_id}-practice-test.csv'
        else:
            flashcards = pack.get('flashcards', [])
            if not flashcards:
                return jsonify({'error': 'No flashcards available'}), 400
            writer.writerow(['question', 'answer'])
            for card in flashcards:
                writer.writerow([card.get('front', ''), card.get('back', '')])
            filename = f'study-pack-{pack_id}-flashcards.csv'
        csv_bytes = io.BytesIO(output.getvalue().encode('utf-8'))
        csv_bytes.seek(0)
        return send_file(csv_bytes, mimetype='text/csv', as_attachment=True, download_name=filename)
    except Exception as e:
        print(f"Error exporting study pack flashcards CSV {pack_id}: {e}")
        return jsonify({'error': 'Could not export CSV'}), 500

@app.route('/api/study-packs/<pack_id>/export-notes', methods=['GET'])
def export_study_pack_notes(pack_id):
    decoded_token = verify_firebase_token(request)
    if not decoded_token:
        return jsonify({'error': 'Unauthorized'}), 401

    uid = decoded_token['uid']
    try:
        doc = db.collection('study_packs').document(pack_id).get()
        if not doc.exists:
            return jsonify({'error': 'Study pack not found'}), 404
        pack = doc.to_dict()
        if pack.get('uid', '') != uid:
            return jsonify({'error': 'Forbidden'}), 403

        notes_markdown = str(pack.get('notes_markdown', '') or '').strip()
        if not notes_markdown:
            return jsonify({'error': 'No integrated notes available'}), 400

        export_format = request.args.get('format', 'docx').strip().lower()
        base_name = f"study-pack-{pack_id}-notes"
        pack_title = str(pack.get('title', 'Lecture Notes') or 'Lecture Notes').strip()

        if export_format == 'md':
            md_bytes = io.BytesIO(notes_markdown.encode('utf-8'))
            md_bytes.seek(0)
            return send_file(
                md_bytes,
                mimetype='text/markdown',
                as_attachment=True,
                download_name=f"{base_name}.md"
            )

        if export_format == 'docx':
            docx = markdown_to_docx(notes_markdown, pack_title)
            docx_bytes = io.BytesIO()
            docx.save(docx_bytes)
            docx_bytes.seek(0)
            return send_file(
                docx_bytes,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f"{base_name}.docx"
            )

        return jsonify({'error': 'Invalid format'}), 400
    except Exception as e:
        print(f"Error exporting study pack notes {pack_id}: {e}")
        return jsonify({'error': 'Could not export notes'}), 500

@app.route('/api/study-packs/<pack_id>/export-pdf', methods=['GET'])
def export_study_pack_pdf(pack_id):
    decoded_token = verify_firebase_token(request)
    if not decoded_token:
        return jsonify({'error': 'Unauthorized'}), 401

    if not REPORTLAB_AVAILABLE:
        return jsonify({
            'error': "PDF export is currently unavailable on this server. Install dependency: pip install reportlab==4.2.5"
        }), 503

    uid = decoded_token['uid']
    try:
        doc = db.collection('study_packs').document(pack_id).get()
        if not doc.exists:
            return jsonify({'error': 'Study pack not found'}), 404

        pack = doc.to_dict()
        if pack.get('uid', '') != uid:
            return jsonify({'error': 'Forbidden'}), 403

        include_answers_raw = str(request.args.get('include_answers', '1')).strip().lower()
        include_answers = include_answers_raw in {'1', 'true', 'yes', 'on'}
        pdf_io = build_study_pack_pdf(pack, include_answers=include_answers)
        filename_suffix = '' if include_answers else '-no-answers'
        return send_file(
            pdf_io,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f"study-pack-{pack_id}{filename_suffix}.pdf"
        )
    except Exception as e:
        print(f"Error exporting study pack PDF {pack_id}: {e}")
        return jsonify({'error': 'Could not export PDF'}), 500

if __name__ == '__main__':
    app.run(debug=True, port=5000)
